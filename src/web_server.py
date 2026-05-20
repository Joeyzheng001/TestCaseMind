"""
ThesisMind local BS web server.

This intentionally uses the Python standard library so the first web prototype
can run without npm or extra backend dependencies. The API shape is kept close
to a future FastAPI implementation.
"""

from __future__ import annotations

import argparse
import datetime
import hashlib
import html
import io
import json
import logging
import mimetypes
import os
import re
import sqlite3
import sys
import tempfile
import threading
import time
import uuid
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from logging.handlers import RotatingFileHandler
from pathlib import Path
from typing import Any, Dict, List, Set, Tuple

from anthropic import Anthropic


PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.llm_config import load_llm_config
from src.method_registry import get_method_aliases, get_method_phases, get_method_alias_map, get_method_domains, get_method_keywords, get_research_tool_methods, get_registry
from src.tools import generate_outline, search_knowledge_base
from src.domain_templates import build_domain_templates, load_domain_template
from src.exporter import export_docx, export_pdf
from src.document_converter import extract_text
from src.vector_store import build_index
# asset_crypto removed — all assets use permanent paths directly
from src.license_manager import LicenseManager, TrialLicense
from src.api_registry import (
    CITATION_GENERATE_TASK_KIND,
    KB_INIT_TASK_KIND,
    OUTLINE_TASK_KIND,
    PAPER_PIPELINE_TASK_KIND,
    PROPOSAL_TASK_KIND,
    TASK_PERMISSION_BY_KIND,
)
from src.consistency_engine import (
    merge_commitments_to_memory,
    build_commitment_brief,
    build_unresolved_warning,
    verify_commitments,
    verify_citations,
)
from src.method_context import build_method_context
from src.risk_checker import run_risk_scan, format_risk_report, run_method_risk_scan
from src.kb_manager import (
    detect_category,
    categorize_output_dir,
    build_rich_metadata,
    build_outline_index,
    load_catalog,
    query_outlines_by_direction,
)

WEB_ROOT = PROJECT_ROOT / "web"
KB_ROOT = PROJECT_ROOT / "knowledge_base" / "references"
KB_CONVERTED_ROOT = KB_ROOT / "converted"
CARDS_DIR = PROJECT_ROOT / "cards" / "methods"
OUTPUT_ROOT = PROJECT_ROOT / "output"
ENV_PATH = PROJECT_ROOT / ".env"
WORKSPACE_DB = OUTPUT_ROOT / "workspace.sqlite3"
TASKS: Dict[str, Dict[str, Any]] = {}
TASK_LOCK = threading.Lock()
GLOBAL_WORKSPACE_KEYS = {"__projects", "__current_project_id"}
DEFAULT_PROJECT_ID = "default"
REFERENCE_SOURCE_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}
BP_PATH = PROJECT_ROOT / "knowledge_base" / "best_practices.json"
MAX_JSON_BODY_BYTES = 2 * 1024 * 1024
MAX_MULTIPART_BODY_BYTES = 50 * 1024 * 1024


# ── Logging ────────────────────────────────────────────────
LOG_DIR = PROJECT_ROOT / "logs"
LOG_DIR.mkdir(exist_ok=True)

logger = logging.getLogger("thesismind")
logger.setLevel(logging.INFO)

_fmt = logging.Formatter(
    "%(asctime)s  %(levelname)-7s  %(message)s", datefmt="%Y-%m-%d %H:%M:%S"
)
_file_handler = RotatingFileHandler(
    LOG_DIR / "server.log", maxBytes=5 * 1024 * 1024, backupCount=5, encoding="utf-8"
)
_file_handler.setFormatter(_fmt)
logger.addHandler(_file_handler)

_stdout_handler = logging.StreamHandler(sys.stdout)
_stdout_handler.setFormatter(_fmt)
logger.addHandler(_stdout_handler)


class PayloadTooLarge(ValueError):
    """Raised when an HTTP request body exceeds endpoint limits."""


def task_log(task_id: str, message: str) -> None:
    logger.info("[task:%s] %s", task_id[:8] if task_id else "-", message)
    with TASK_LOCK:
        task = TASKS.get(task_id)
        if not task:
            return
        logs = task.setdefault("logs", [])
        logs.append(
            {
                "time": time.strftime("%H:%M:%S"),
                "message": message,
            }
        )
        task["message"] = message


def workspace_connection() -> sqlite3.Connection:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    connection = sqlite3.connect(WORKSPACE_DB)
    connection.execute(
        """
        CREATE TABLE IF NOT EXISTS workspace_state (
            key TEXT PRIMARY KEY,
            value_json TEXT NOT NULL,
            updated_at REAL NOT NULL
        )
        """
    )
    connection.execute(
        """
        CREATE TABLE IF NOT EXISTS drafts (
            draft_key TEXT PRIMARY KEY,
            content TEXT NOT NULL,
            updated_at REAL NOT NULL
        )
        """
    )
    connection.commit()
    return connection


def _save_workspace_raw(key: str, value: Any) -> None:
    with workspace_connection() as connection:
        connection.execute(
            """
            INSERT INTO workspace_state(key, value_json, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(key) DO UPDATE SET
                value_json = excluded.value_json,
                updated_at = excluded.updated_at
            """,
            (key, json.dumps(value, ensure_ascii=False), time.time()),
        )


def _load_workspace_raw(key: str, default: Any = None) -> Any:
    with workspace_connection() as connection:
        row = connection.execute(
            "SELECT value_json FROM workspace_state WHERE key = ?", (key,)
        ).fetchone()
    return json.loads(row[0]) if row else default


def list_projects() -> List[Dict[str, Any]]:
    projects = _load_workspace_raw("__projects", None)
    if not projects:
        projects = [
            {
                "id": DEFAULT_PROJECT_ID,
                "topic": "默认论文项目",
                "created_at": time.time(),
                "updated_at": time.time(),
            }
        ]
        _save_workspace_raw("__projects", projects)
        _save_workspace_raw("__current_project_id", DEFAULT_PROJECT_ID)
    return projects


def current_project_id() -> str:
    project_id = _load_workspace_raw("__current_project_id", DEFAULT_PROJECT_ID)
    project_ids = {project["id"] for project in list_projects()}
    if project_id not in project_ids:
        project_id = DEFAULT_PROJECT_ID
        _save_workspace_raw("__current_project_id", project_id)
    return project_id


def set_current_project(project_id: str) -> None:
    if project_id not in {project["id"] for project in list_projects()}:
        raise ValueError("project not found")
    _save_workspace_raw("__current_project_id", project_id)


def touch_current_project(topic: str = "") -> None:
    project_id = current_project_id()
    projects = list_projects()
    for project in projects:
        if project["id"] == project_id:
            if topic:
                project["topic"] = topic
            project["updated_at"] = time.time()
            break
    _save_workspace_raw("__projects", projects)


def create_project(topic: str = "") -> Dict[str, Any]:
    project = {
        "id": uuid.uuid4().hex,
        "topic": topic.strip() or "未命名论文项目",
        "created_at": time.time(),
        "updated_at": time.time(),
    }
    projects = [project, *list_projects()]
    _save_workspace_raw("__projects", projects)
    _save_workspace_raw("__current_project_id", project["id"])
    return project


def scoped_key(key: str) -> str:
    if key in GLOBAL_WORKSPACE_KEYS:
        return key
    return f"project:{current_project_id()}:{key}"


def save_workspace_value(key: str, value: Any) -> None:
    _save_workspace_raw(scoped_key(key), value)


def load_workspace_value(key: str, default: Any = None) -> Any:
    if key in GLOBAL_WORKSPACE_KEYS:
        return _load_workspace_raw(key, default)
    value = _load_workspace_raw(scoped_key(key), None)
    if value is not None:
        return value
    if current_project_id() == DEFAULT_PROJECT_ID:
        return _load_workspace_raw(key, default)
    return default


def default_thesis_memory() -> Dict[str, Any]:
    return {
        "research_context": {
            "topic": "",
            "direction": "",
            "research_object": "",
            "research_boundary": "",
            "core_objective": "",
        },
        "problem_list": [],
        "method_usage": {},
        "solution_design": [],
        "evaluation_indicators": [],
        "terminology": {},
        "chapter_summaries": {},
        "section_summaries": {},
        "style_preferences": {
            "tone": "正式、克制、工程管理硕士论文风格",
            "avoid": ["口语化表达", "无依据夸大", "虚构文献"],
        },
        "outline_summary": [],
        "updated_at": time.time(),
    }


def merge_memory_schema(memory: Dict[str, Any]) -> Dict[str, Any]:
    base = default_thesis_memory()
    for key, value in (memory or {}).items():
        if isinstance(value, dict) and isinstance(base.get(key), dict):
            base[key].update(value)
        else:
            base[key] = value
    return base


def _unique_append(items: List[Any], value: Any, limit: int = 20) -> List[Any]:
    if not value:
        return items
    if value not in items:
        items.append(value)
    return items[-limit:]


def _compute_downstream_stale(
    memory: Dict[str, Any], chapter_key: str
) -> List[Dict[str, Any]]:
    """Find downstream chapters that have drafts and may be stale after an upstream edit."""
    outline = load_workspace_value("outline")
    if not outline:
        return []
    chapters = outline.get("chapters", [])
    try:
        current_idx = int(chapter_key) - 1
    except (ValueError, TypeError):
        return []
    drafts = load_drafts()
    stale: List[Dict[str, Any]] = []
    for ch in chapters[current_idx + 1:]:
        ch_num = str(ch.get("number", ""))
        # Check if any section of this chapter has a saved draft
        has_draft = any(
            dk.startswith(ch_num + ".") for dk in drafts
        )
        if has_draft:
            stale.append({
                "chapter": ch_num,
                "title": ch.get("title", ""),
                "reason": f"第{chapter_key}章已修改，本章可能需要重新生成以保持一致性",
            })
    return stale


def _mark_all_drafts_stale(reason: str) -> List[Dict[str, Any]]:
    """当框架或大纲变更时，标记所有已有草稿的章节为可能过时。"""
    drafts = load_drafts()
    if not drafts:
        return []
    chapters_with_drafts: set = set()
    for dk in drafts:
        ch = dk.split(".", 1)[0]
        chapters_with_drafts.add(ch)
    stale: List[Dict[str, Any]] = []
    for ch_num in sorted(chapters_with_drafts, key=lambda x: int(x) if x.isdigit() else 0):
        stale.append({
            "chapter": ch_num,
            "title": "",
            "reason": reason,
        })
    if stale:
        memory = merge_memory_schema(load_workspace_value("thesis_memory", {}) or {})
        memory["stale_chapters"] = stale
        save_workspace_value("thesis_memory", memory)
    return stale


def update_memory_from_draft(
    memory: Dict[str, Any], draft_key: str, content: str
) -> Dict[str, Any]:
    memory = merge_memory_schema(memory)
    snippet = re.sub(r"\s+", " ", content).strip()[:800]
    memory["section_summaries"][draft_key] = snippet
    chapter_key = draft_key.split(".", 1)[0]
    memory["chapter_summaries"][chapter_key] = (
        memory["chapter_summaries"].get(chapter_key, "") + " " + snippet
    )[:1200]

    sentences = re.split(r"[。；;]\s*", content)
    for sentence in sentences:
        compact = re.sub(r"\s+", "", sentence)
        if not compact:
            continue
        if any(word in compact for word in ["问题", "不足", "缺陷", "痛点", "原因"]):
            _unique_append(memory["problem_list"], compact[:120])
        if any(word in compact for word in ["方案", "优化", "改进", "措施", "机制"]):
            _unique_append(memory["solution_design"], compact[:120])
        if any(
            word in compact
            for word in ["指标", "评价", "缺陷率", "返工率", "满意度", "周期"]
        ):
            _unique_append(memory["evaluation_indicators"], compact[:120])

    # 从注册中心动态获取 distinctive 方法名作为术语检测词
    from src.method_registry import get_registry as _term_registry
    _reg = _term_registry()
    _term_set: Set[str] = set()
    for _name in _reg.get_all_names():
        # 缩略词（全大写+数字，2-6字符）或短中文名（≤6字）作为术语标志
        if re.match(r'^[A-Z0-9]{2,6}$', _name) or (len(_name) <= 6 and any('一' <= c <= '鿿' for c in _name)):
            _term_set.add(_name)
    for _term in sorted(_term_set, key=len, reverse=True):
        if _term.lower() in content.lower():
            memory["terminology"].setdefault(
                _term, f"{_term}：本文采用的关键理论或方法。"
            )

    memory = merge_commitments_to_memory(memory, chapter_key, draft_key, content)
    memory["updated_at"] = time.time()
    return memory


def save_draft(draft_key: str, content: str) -> Dict[str, Any]:
    content = clean_generated_content(content)
    stored_key = f"{current_project_id()}:{draft_key}"
    with workspace_connection() as connection:
        connection.execute(
            """
            INSERT INTO drafts(draft_key, content, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(draft_key) DO UPDATE SET
                content = excluded.content,
                updated_at = excluded.updated_at
            """,
            (stored_key, content, time.time()),
        )
    memory = update_memory_from_draft(
        load_workspace_value("thesis_memory", {}) or {}, draft_key, content
    )
    # Compute downstream stale chapters
    chapter_key = draft_key.split(".", 1)[0]
    stale_list = _compute_downstream_stale(memory, chapter_key)
    memory["stale_chapters"] = stale_list
    save_workspace_value("thesis_memory", memory)
    return {"status": "ok", "stale_chapters": stale_list}


def load_drafts() -> Dict[str, str]:
    prefix = f"{current_project_id()}:"
    with workspace_connection() as connection:
        rows = connection.execute("SELECT draft_key, content FROM drafts").fetchall()
    drafts = {}
    for key, content in rows:
        if key.startswith(prefix):
            drafts[key[len(prefix) :]] = content
        elif current_project_id() == DEFAULT_PROJECT_ID and ":" not in key:
            drafts[key] = content
    return drafts


def build_thesis_memory(
    payload: Dict[str, Any], outline: Dict[str, Any] = None
) -> Dict[str, Any]:
    existing = merge_memory_schema(load_workspace_value("thesis_memory", {}) or {})
    research_context = existing["research_context"]
    research_context.update(
        {
            "topic": payload.get("topic") or research_context.get("topic", ""),
            "direction": payload.get("direction_name")
            or payload.get("direction")
            or research_context.get("direction", ""),
            "project_context": payload.get("project_context")
            or research_context.get("project_context", ""),
        }
    )
    memory = {
        **existing,
        "research_context": research_context,
        "methods": payload.get("methods") or existing.get("methods", []),
        "phase_methods": payload.get("phase_methods")
        or existing.get("phase_methods", {}),
        "updated_at": time.time(),
    }
    if outline:
        memory["outline_summary"] = [
            {
                "chapter": chapter.get("title"),
                "sections": [
                    section.get("title") for section in chapter.get("sections", [])
                ],
            }
            for chapter in outline.get("chapters", [])
        ]
    save_workspace_value("thesis_memory", memory)
    return memory


def persist_project_context(payload: Dict[str, Any]) -> str:
    project_context = str(payload.get("project_context", "") or "").strip()
    if not project_context:
        return load_workspace_value("project_context", "") or ""
    save_workspace_value("project_context", project_context)
    memory = merge_memory_schema(load_workspace_value("thesis_memory", {}) or {})
    memory["research_context"]["project_context"] = project_context
    memory["research_context"]["topic"] = payload.get("topic") or memory[
        "research_context"
    ].get("topic", "")
    memory["research_context"]["direction"] = (
        payload.get("direction_name")
        or payload.get("direction")
        or memory["research_context"].get("direction", "")
    )
    memory["updated_at"] = time.time()
    save_workspace_value("thesis_memory", memory)
    return project_context


DIRECTIONS = [
    {"id": "quality_management", "name": "质量管理", "desc": "确保项目成果满足质量要求，包括质量规划、质量保证、质量控制、质量改进"},
    {"id": "risk_management", "name": "风险管理", "desc": "识别、分析和应对项目不确定性，包括风险识别、定性/定量分析、风险应对和监控"},
    {"id": "schedule_management", "name": "进度管理", "desc": "管理项目时间安排，包括活动定义、排序、工期估算、进度计划制定和进度控制"},
    {"id": "cost_management", "name": "成本管理", "desc": "管理项目预算和费用，包括成本估算、预算制定和成本控制"},
    {"id": "requirements_management", "name": "需求管理", "desc": "管理项目需求和范围，包括需求收集、需求分析、需求变更和范围控制"},
    {"id": "process_optimization", "name": "流程优化", "desc": "分析和改进项目流程，包括流程梳理、瓶颈分析、流程再造和持续优化"},
    {"id": "supply_chain_logistics", "name": "供应链与物流", "desc": "管理供应链和物流环节，包括供应商管理、库存优化、配送路径和物流成本"},
    {"id": "resource_management", "name": "资源管理", "desc": "管理项目人力、设备和物资资源，包括资源规划、资源分配、资源平衡和团队建设"},
    {"id": "communication_management", "name": "沟通管理", "desc": "管理项目信息传递和沟通机制，包括沟通规划、信息分发、绩效报告和相关方沟通"},
    {"id": "stakeholder_management", "name": "相关方管理", "desc": "管理项目相关方期望和参与，包括相关方识别、参与规划、期望管理和冲突协调"},
]

WORD_WEIGHTS = {
    "绪论": 0.14,
    "理论基础": 0.17,
    "文献综述": 0.17,
    "现状": 0.2,
    "问题": 0.2,
    "方案": 0.22,
    "实施": 0.17,
    "验证": 0.17,
    "评价": 0.17,
    "结论": 0.1,
}

SECTION_WORD_WEIGHTS = {
    "背景": 0.9,
    "意义": 0.8,
    "现状": 1.15,
    "文献": 1.1,
    "理论": 1.0,
    "基础": 1.0,
    "对象": 0.9,
    "数据": 1.15,
    "调研": 1.1,
    "问题": 1.35,
    "原因": 1.25,
    "成因": 1.25,
    "方案": 1.35,
    "模型": 1.2,
    "机制": 1.2,
    "实施": 1.15,
    "保障": 0.95,
    "评价": 1.2,
    "验证": 1.2,
    "指标": 1.1,
    "结论": 0.75,
    "展望": 0.55,
}

CITATION_INDEX_PATH = (
    PROJECT_ROOT / "knowledge_base" / "templates" / "citation_index.json"
)
METHODOLOGY_CATALOG_PATH = (
    PROJECT_ROOT / "knowledge_base" / "templates" / "methodology_catalog.json"
)
METHOD_SUMMARIES_CACHE = (
    PROJECT_ROOT / "knowledge_base" / "templates" / "method_summaries.json"
)


def _load_method_summaries() -> Dict[str, str]:
    """Load cached method summaries {method_name: description}."""
    if METHOD_SUMMARIES_CACHE.exists():
        try:
            return json.loads(METHOD_SUMMARIES_CACHE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def _save_method_summaries(cache: Dict[str, str]) -> None:
    METHOD_SUMMARIES_CACHE.parent.mkdir(parents=True, exist_ok=True)
    METHOD_SUMMARIES_CACHE.write_text(
        json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _generate_method_summaries_via_llm(method_names: List[str]) -> Dict[str, str]:
    """Call LLM to generate 1-2 sentence descriptions for each method name."""
    if not method_names:
        return {}
    config = load_llm_config()
    if not config.api_key:
        return {}

    names_list = "\n".join(f"- {n}" for n in method_names)
    prompt = f"""你是一位工程管理硕士论文方法论专家。请为以下每个方法论/工具写一句话的学术介绍（15-40字），说明它是什么、主要用于什么场景。

方法论列表：
{names_list}

请严格按照以下JSON格式输出，不要输出任何其它文字：
{{
  "方法名": "一句话介绍",
  ...
}}"""

    try:
        client, provider = _build_llm_client_and_provider(config)
        response = _llm_create_message(
            client, provider, config,
            system="",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=_token_budget(config, "tiny"),
        )
        text = _llm_response_text(response, provider)
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            return json.loads(match.group(0))
    except Exception:
        pass
    return {}


# 方向→搜索关键词（领域概念部分，方法名从注册中心动态补充）
_CITATION_DIRECTION_DOMAIN_KEYWORDS: Dict[str, List[str]] = {
    "quality_management": ["质量", "缺陷", "测试", "研发"],
    "risk_management": ["风险", "风控", "安全", "反欺诈"],
    "schedule_management": ["进度", "工期", "排产", "计划"],
    "requirements_management": ["需求"],
    "process_optimization": ["流程", "改善", "优化", "过程"],
    "cost_management": ["成本", "投资", "库存"],
    "supply_chain_logistics": ["供应链", "物流", "配送", "库存"],
    "resource_management": ["资源", "人力", "团队", "技能", "培训"],
    "communication_management": ["沟通", "会议", "信息传递", "报告"],
    "stakeholder_management": ["相关方", "利益相关", "冲突", "期望"],
}

def _build_citation_directions() -> Dict[str, List[str]]:
    """从注册中心动态构建方向→关键词映射（领域概念 + 方法名/别名）。"""
    from src.method_registry import get_registry as _cd_reg
    _reg = _cd_reg()
    _domains_map = _reg.get_domain_to_names()
    _aliases_map = _reg.get_aliases()
    _result: Dict[str, List[str]] = {}
    for _dir_id, _domain_kws in _CITATION_DIRECTION_DOMAIN_KEYWORDS.items():
        _kws = list(_domain_kws)
        for _name in _domains_map.get(_dir_id, []):
            if _name not in _kws:
                _kws.append(_name)
            for _alias in _aliases_map.get(_name, []):
                if _alias not in _kws:
                    _kws.append(_alias)
        _result[_dir_id] = _kws
    return _result

CITATION_DIRECTIONS: Dict[str, List[str]] = {}
CITATION_DIRECTIONS.update(_build_citation_directions())


def _get_method_keywords_map() -> Dict[str, List[str]]:
    """从注册中心动态获取方法名→搜索关键词映射。"""
    from src.method_registry import get_method_keywords
    return get_method_keywords()



def _json_response(
    handler: BaseHTTPRequestHandler, payload: Dict[str, Any], status: int = 200
) -> None:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


def _read_json(handler: BaseHTTPRequestHandler) -> Dict[str, Any]:
    length = int(handler.headers.get("Content-Length", "0") or 0)
    if length <= 0:
        return {}
    if length > MAX_JSON_BODY_BYTES:
        raise PayloadTooLarge("JSON 请求体过大")
    raw = handler.rfile.read(length).decode("utf-8")
    return json.loads(raw or "{}")


def _read_multipart(handler: BaseHTTPRequestHandler) -> Dict[str, Any]:
    """解析 multipart/form-data 请求体。返回 {field_name: value} 或 {file_name, file_data, content_type}。"""
    content_type = handler.headers.get("Content-Type", "")
    boundary_match = None
    for part in content_type.split(";"):
        part = part.strip()
        if part.startswith("boundary="):
            boundary_match = part[len("boundary="):].strip('"').strip("'")
            break
    if not boundary_match:
        return {}

    length = int(handler.headers.get("Content-Length", "0") or 0)
    if length > MAX_MULTIPART_BODY_BYTES:
        raise PayloadTooLarge("上传文件过大")
    body = handler.rfile.read(length)
    boundary_bytes = boundary_match.encode()
    parts = body.split(b"--" + boundary_bytes)

    result = {"file_data": None, "file_name": "", "content_type": ""}
    for part in parts:
        if not part or part == b"--\r\n" or part == b"--":
            continue
        # Split headers from body
        header_end = part.find(b"\r\n\r\n")
        if header_end < 0:
            continue
        headers_raw = part[:header_end].decode("utf-8", errors="replace")
        data = part[header_end + 4:]
        # Remove trailing \r\n and boundary artifacts
        if data.endswith(b"\r\n"):
            data = data[:-2]

        # Parse Content-Disposition
        name = ""
        filename = ""
        for disp_line in headers_raw.split("\r\n"):
            if disp_line.startswith("Content-Disposition"):
                for token in disp_line.split(";"):
                    token = token.strip()
                    if token.startswith("name="):
                        name = token[5:].strip('"').strip("'")
                    if token.startswith("filename="):
                        filename = token[9:].strip('"').strip("'")
            if disp_line.startswith("Content-Type:"):
                result["content_type"] = disp_line[len("Content-Type:"):].strip()

        if filename:
            result["file_name"] = filename
            result["file_data"] = data
            result["field_name"] = name
        elif name and data:
            result[name] = data.decode("utf-8", errors="replace")

    return result


def _handle_table_generate(
    handler: BaseHTTPRequestHandler, payload: Dict[str, Any] | None = None
) -> None:
    """处理 Excel 上传 + LLM 表格生成 (multipart/form-data)。"""
    data = payload if payload is not None else _read_multipart(handler)
    file_data = data.get("file_data")
    file_name = data.get("file_name", "")
    description = data.get("description", "").strip()

    if not description:
        _json_response(handler, {"status": "error", "message": "请填写表格描述"}, status=400)
        return

    excel_text = ""
    if file_data and file_name.lower().endswith((".xlsx", ".xls")):
        try:
            import openpyxl
            from io import BytesIO
            wb = openpyxl.load_workbook(BytesIO(file_data), data_only=True)
            parts = []
            for sn in wb.sheetnames:
                ws = wb[sn]
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue
                parts.append(f"=== Sheet: {sn} ===")
                for i, row in enumerate(rows[:50]):
                    cells = [str(c) if c is not None else "" for c in row[:20]]
                    parts.append("\t".join(cells))
            excel_text = "\n".join(parts)
        except Exception as e:
            _json_response(handler, {"status": "error", "message": f"Excel 解析失败: {e}"}, status=400)
            return
    elif file_data:
        _json_response(handler, {"status": "error", "message": "仅支持 .xlsx 或 .xls 文件"}, status=400)
        return

    if not excel_text:
        _json_response(handler, {"status": "error", "message": "请上传 Excel 文件"}, status=400)
        return

    config = load_llm_config()
    if not config.api_key:
        _json_response(handler, {"status": "error", "message": "请先配置 LLM API Key"}, status=400)
        return

    client, provider = _build_llm_client_and_provider(config)
    prompt = (
        "你是一个学术表格排版助手。用户上传了一个 Excel 文件，并描述了想要的表格效果。\n\n"
        f"用户描述：\n{description[:2000]}\n\n"
        f"Excel 数据内容：\n{excel_text[:8000]}\n\n"
        "请根据用户描述生成格式化的学术表格（Markdown 格式）。\n"
        "要求：\n"
        "1. 使用 Markdown pipe table 格式\n"
        "2. 使用中文表头\n"
        "3. 数字右对齐，文字左对齐\n"
        "4. 行数不超过30行，数据多时做适度汇总\n"
        "5. 只输出表格本身，附带简短标题（### 开头），不要输出无关解释"
    )

    try:
        response = _llm_create_message(
            client, provider, config,
            system="你是学术论文表格排版专家，擅长将原始数据转换为规范的学术表格。",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=config.max_tokens,
        )
        table_md = _llm_response_text(response, provider)
        _json_response(handler, {"status": "ok", "table": table_md, "format": "markdown"})
    except Exception as e:
        _json_response(handler, {"status": "error", "message": f"LLM 生成失败: {e}"}, status=500)


def _handle_table_generate_from_text(handler: BaseHTTPRequestHandler, payload: Dict[str, Any]) -> None:
    """从文本内容调用 LLM 生成表格（写作页调用）。"""
    text = (payload.get("text") or "").strip()
    description = (payload.get("description") or "").strip()
    topic = (payload.get("topic") or "").strip()

    if not description:
        _json_response(handler, {"status": "error", "message": "请填写表格描述"}, status=400)
        return

    config = load_llm_config()
    if not config.api_key:
        _json_response(handler, {"status": "error", "message": "请先配置 LLM API Key"}, status=400)
        return

    content_part = f"文本内容：\n{text[:6000]}" if text else "用户未提供文本，请根据描述自主构建表格数据。"
    prompt = (
        "你是一个学术表格排版助手。用户希望从以下文本中提取关键信息，生成格式化的学术表格（Markdown 格式）。\n\n"
        f"论文主题：{topic[:500]}\n\n"
        f"用户描述（想要的表格效果）：\n{description[:1500]}\n\n"
        f"{content_part}\n\n"
        "要求：\n"
        "1. 使用 Markdown pipe table 格式\n"
        "2. 使用中文表头\n"
        "3. 数字右对齐，文字左对齐\n"
        "4. 行数不超过25行，信息多时做适度汇总\n"
        "5. 只输出表格本身，附带简短标题（### 开头），不要输出无关解释"
    )

    try:
        client, provider = _build_llm_client_and_provider(config)
        response = _llm_create_message(
            client, provider, config,
            system="你是学术论文表格排版专家。",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=config.max_tokens,
        )
        table_md = _llm_response_text(response, provider)
        _json_response(handler, {"status": "ok", "table": table_md, "format": "markdown"})
    except Exception as e:
        _json_response(handler, {"status": "error", "message": f"LLM 生成失败: {e}"}, status=500)

def _handle_paper_upload(payload: Dict[str, Any]) -> Dict[str, Any]:
    """处理论文上传：保存文件并转换为 Markdown。"""
    file_data = payload.get("file_data")
    file_name = payload.get("file_name", "uploaded")

    if not file_data:
        return {"status": "error", "message": "未收到文件数据"}

    # Validate extension
    ext = Path(file_name).suffix.lower()
    if ext not in (".pdf", ".docx", ".doc"):
        return {"status": "error", "message": f"不支持的文件格式: {ext}，仅支持 PDF/DOCX"}

    # Save to papers/ directory
    papers_dir = Path(__file__).resolve().parent.parent / "papers"
    papers_dir.mkdir(parents=True, exist_ok=True)

    safe_name = Path(file_name).name
    dest = papers_dir / safe_name
    if dest.exists():
        safe_name = f"{dest.stem}_{int(time.time())}{ext}"
        dest = papers_dir / safe_name

    dest.write_bytes(file_data)

    # Convert to Markdown using existing document_converter
    md_path = None
    try:
        from src.document_converter import convert_document
        conv_result = convert_document(str(dest))
        md_path = conv_result.get("md_path") or conv_result.get("output_path") or ""
    except Exception as e:
        return {"status": "ok", "file_name": safe_name, "saved_path": str(dest),
                "message": f"文件已保存，但 MD 转换失败: {e}"}

    return {"status": "ok", "file_name": safe_name, "saved_path": str(dest),
            "md_path": str(md_path) if md_path else ""}


def _run_paper_pipeline_task(task_id: str, doc_id: str = "") -> None:
    """后台运行完整论文流水线。"""
    try:
        TASKS[task_id] = {"status": "running", "progress": 0, "message": "启动流水线..."}
        from src.paper_pipeline import run_full_pipeline

        for progress in run_full_pipeline(task_id=task_id, force_reindex=(not doc_id)):
            TASKS[task_id].update({
                "progress": min(95, TASKS[task_id].get("progress", 0) + 18),
                "message": progress.get("phase", ""),
                "logs": progress.get("logs", []),
            })

        TASKS[task_id].update({"status": "done", "progress": 100, "message": "流水线完成"})
    except Exception as e:
        TASKS[task_id].update({"status": "error", "progress": 0, "message": f"流水线失败: {e}"})


def start_paper_pipeline_task(doc_id: str = "") -> str:
    """创建论文流水线异步任务。"""
    with TASK_LOCK:
        task_id = uuid4().hex
        dedup = [k for k, v in TASKS.items()
                 if v.get("kind") == PAPER_PIPELINE_TASK_KIND and v.get("status") in ("running",)]
        if dedup:
            return dedup[0]

        TASKS[task_id] = {
            "task_id": task_id,
            "status": "starting",
            "progress": 0,
            "message": "正在启动论文流水线...",
            "kind": PAPER_PIPELINE_TASK_KIND,
            "permission_menu": TASK_PERMISSION_BY_KIND[PAPER_PIPELINE_TASK_KIND],
            "logs": [],
        }

    threading.Thread(
        target=_run_paper_pipeline_task,
        args=(task_id, doc_id),
        daemon=True,
    ).start()
    return task_id


def _safe_static_path(route: str) -> Path:
    relative = route.lstrip("/") or "index.html"
    path = (WEB_ROOT / relative).resolve()
    if not str(path).startswith(str(WEB_ROOT.resolve())):
        raise ValueError("invalid path")
    if path.is_dir():
        path = path / "index.html"
    return path


def _file_response(
    handler: BaseHTTPRequestHandler, path: Path, filename: str, content_type: str
) -> None:
    body = path.read_bytes()
    handler.send_response(200)
    handler.send_header("Content-Type", content_type)
    handler.send_header("Content-Length", str(len(body)))
    handler.send_header("Content-Disposition", f'attachment; filename="{filename}"')
    handler.end_headers()
    handler.wfile.write(body)


def _check_license_api(handler, api_path: str) -> bool:
    """检查许可证是否允许访问指定 API。不允许时自动返回 403。"""
    manager = LicenseManager()
    allowed, reason = manager.can_access_api(api_path, method=getattr(handler, "command", "GET"))
    if not allowed:
        _json_response(handler, {"error": reason, "code": "LICENSE_REQUIRED"}, status=403)
        return False
    return True


def _check_license_menu(handler, menu_id: str) -> bool:
    """检查当前许可证是否允许访问指定功能菜单。"""
    manager = LicenseManager()
    allowed, reason = manager.can_access_menu(menu_id)
    if not allowed:
        _json_response(handler, {"error": reason, "code": "LICENSE_REQUIRED"}, status=403)
        return False
    return True


def _mask_secret(value: str) -> str:
    if not value:
        return ""
    if len(value) <= 10:
        return "*" * len(value)
    return f"{value[:4]}...{value[-4:]}"


def _read_env_lines() -> List[str]:
    if not ENV_PATH.exists():
        return []
    return ENV_PATH.read_text(encoding="utf-8").splitlines()


def _write_env_values(values: Dict[str, str]) -> None:
    lines = _read_env_lines()
    seen = set()
    output = []

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in line:
            output.append(line)
            continue
        key = line.split("=", 1)[0].strip()
        if key in values:
            output.append(f"{key}={values[key]}")
            seen.add(key)
        else:
            output.append(line)

    for key, value in values.items():
        if key not in seen:
            output.append(f"{key}={value}")

    ENV_PATH.write_text("\n".join(output).rstrip() + "\n", encoding="utf-8")


def _enrich_items_from_cards(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """将方法卡数据注入扫描结果，并补全未在 KB 中出现但属于平台标准方法卡的条目。"""
    cards_db = PROJECT_ROOT / "knowledge_base" / "cards.sqlite3"
    card_list: List[Dict[str, Any]] = []
    if cards_db.exists():
        try:
            conn = sqlite3.connect(str(cards_db))
            conn.row_factory = sqlite3.Row
            rows = conn.execute(
                "SELECT id, name, short_name, aliases, category, phase, difficulty, data_type, "
                "outputs, pairs_with, conflicts_with, domains, source_type, body_summary FROM cards "
                "WHERE type='method_card' AND scope='platform'"
            ).fetchall()
            conn.close()
            for r in rows:
                card = dict(r)
                for field in ["domains", "data_type", "outputs", "pairs_with", "conflicts_with", "phase", "aliases"]:
                    val = card.get(field)
                    if isinstance(val, str):
                        try:
                            card[field] = json.loads(val)
                        except (json.JSONDecodeError, TypeError):
                            card[field] = [] if field != "phase" else []
                card_list.append(card)
        except Exception:
            pass

    # 纯领域概念关键词 → domains（非方法名，而是论文内容中出现的领域信号词）
    _DOMAIN_CONCEPT_KEYWORDS: Dict[str, List[str]] = {
        "返工率": ["quality_management"],
        "缺陷率": ["quality_management"],
        "控制图": ["quality_management"],
        "工期偏差": ["schedule_management"],
        "成本偏差": ["cost_management"],
        "信息化": ["process_optimization", "requirements_management"],
        "库存": ["supply_chain_logistics", "cost_management"],
        "物流": ["supply_chain_logistics"],
        "供应链": ["supply_chain_logistics"],
        "配送": ["supply_chain_logistics"],
        "供应商": ["supply_chain_logistics"],
        "绩效考核": ["cost_management", "quality_management", "requirements_management"],
        "全生命": ["cost_management", "quality_management", "schedule_management"],
        "度量": ["quality_management", "cost_management"],
    }

    def _fallback_domains(method_name: str) -> List[str]:
        # 先查领域概念关键词
        for keyword, domains in _DOMAIN_CONCEPT_KEYWORDS.items():
            if keyword in method_name:
                return domains
        # 再从方法注册中心动态匹配
        from src.method_registry import get_registry as _get_registry
        reg = _get_registry()
        alias_map = reg.get_alias_map()
        domains_map = reg.get_domains()
        # 检查是否匹配任何已知方法名/别名
        for alias_key, canon in alias_map.items():
            if alias_key in method_name.lower():
                return domains_map.get(canon, [])
        return []

    if not card_list:
        cached = _load_method_summaries()
        for item in items:
            domains = _fallback_domains(item.get("name", ""))
            name = item.get("name", "")
            item.update({"domains": domains, "difficulty": "", "data_type": [], "outputs": [], "pairs_with": [], "conflicts_with": [], "summary": cached.get(name, "")})
        return items

    def _normalize(s: str) -> str:
        return re.sub(r"[（）()\s]", "", s).lower()

    # 为每个 API item 建立标准化名字
    norm_items: Dict[str, int] = {}  # norm_name → index in items
    for i, item in enumerate(items):
        norm_items[_normalize(item["name"])] = i

    # 为每个 card 建立标准化名字列表
    card_norms: Dict[str, int] = {}  # norm_name → card_list index
    for ci, card in enumerate(card_list):
        names = [card["name"]]
        if card.get("short_name"):
            names.append(card["short_name"])
        # 也把别名纳入匹配，解决 "变更控制流程优化法" → "流程再造BPR" 这类映射
        for alias in card.get("aliases", []) or []:
            if alias and alias not in names:
                names.append(alias)
        for n in names:
            norm = _normalize(n)
            if norm:
                card_norms[norm] = ci

    def _apply_card(item: Dict[str, Any], card: Dict[str, Any]) -> None:
        # 用卡片 DB 中的规范 ID 和名称覆盖别名单生成的合成 ID/名称
        card_db_id = card.get("id", "")
        if card_db_id:
            item["id"] = card_db_id
        card_name = card.get("name", "")
        if card_name:
            item["name"] = card_name
        item["domains"] = card.get("domains", [])
        item["difficulty"] = card.get("difficulty", "")
        item["data_type"] = card.get("data_type", [])
        item["outputs"] = card.get("outputs", [])
        item["pairs_with"] = card.get("pairs_with", [])
        item["conflicts_with"] = card.get("conflicts_with", [])
        item["source_type"] = card.get("source_type", "builtin")
        summary = (card.get("body_summary") or "").strip()
        if not summary:
            summary = _load_method_summaries().get(item.get("name", ""), "")
        item["summary"] = summary
        # 用卡片的 phase 覆盖扫描结果（卡片更权威）
        card_phases = card.get("phase", [])
        if card_phases:
            api_phases = ["validate" if p == "verify" else p for p in card_phases]
            item["phases"] = api_phases

    matched_card_indices: set = set()

    # 第一轮：精确匹配（标准化后）
    for norm_name, item_idx in norm_items.items():
        if norm_name in card_norms:
            ci = card_norms[norm_name]
            _apply_card(items[item_idx], card_list[ci])
            matched_card_indices.add(ci)
            continue
        # 模糊匹配：子串
        for cnorm, ci in card_norms.items():
            if cnorm in norm_name or norm_name in cnorm:
                _apply_card(items[item_idx], card_list[ci])
                matched_card_indices.add(ci)
                break
        else:
            items[item_idx].update({"domains": [], "difficulty": "", "data_type": [], "outputs": [], "pairs_with": [], "conflicts_with": [], "summary": ""})

    # 第二轮：把没匹配上的卡片作为新条目注入
    existing_norms = set(norm_items.keys())
    for ci, card in enumerate(card_list):
        if ci in matched_card_indices:
            continue
        # 检查是否已有近似条目
        card_norm = _normalize(card["name"])
        if card_norm in existing_norms:
            continue
        # 用 phase 映射到 API phase（card 用 verify，API 用 validate）
        card_phases = card.get("phase", [])
        api_phases = ["validate" if p == "verify" else p for p in card_phases]
        card_db_id = card.get("id", "")
        new_item: Dict[str, Any] = {
            "id": card_db_id if card_db_id else re.sub(r"[^a-zA-Z0-9一-鿿]+", "_", card["name"]).strip("_").lower(),
            "name": card["name"],
            "type": "method",
            "source_count": 0,
            "sources": ["来自方法知识库"],
            "phases": api_phases or ["discover", "solve", "validate"],
            "detected_by": "method_card",
            "domains": card.get("domains", []),
            "difficulty": card.get("difficulty", ""),
            "data_type": card.get("data_type", []),
            "outputs": card.get("outputs", []),
            "pairs_with": card.get("pairs_with", []),
            "conflicts_with": card.get("conflicts_with", []),
            "source_type": card.get("source_type", "builtin"),
            "summary": (card.get("body_summary") or "").strip(),
        }
        items.append(new_item)

    # 过滤掉没有匹配卡片的条目（别名单中有但卡片已删除的）
    items = [it for it in items if it.get("id", "").startswith("method_")]
    # 按 ID 去重，保留第一个（通常是名字更规范的条目）
    seen_ids: set = set()
    deduped: List[Dict[str, Any]] = []
    for it in items:
        cid = it.get("id", "")
        if cid not in seen_ids:
            seen_ids.add(cid)
            deduped.append(it)
    items = deduped

    return items


def scan_methodologies(force: bool = False) -> List[Dict[str, Any]]:
    """从平台方法卡片（METHODOLOGY_ALIASES）构建方法目录，KB 扫描仅用于统计本地论文支撑数量。"""
    if not force and METHODOLOGY_CATALOG_PATH.exists():
        try:
            cached = json.loads(METHODOLOGY_CATALOG_PATH.read_text(encoding="utf-8"))
            if cached.get("items"):
                return _enrich_items_from_cards(cached["items"])
        except Exception:
            pass

    kb_files = (
        [p for p in KB_ROOT.rglob("*") if p.is_file()]
        if KB_ROOT.exists()
        else []
    )
    corpus = "\n".join(str(p.relative_to(KB_ROOT)) for p in kb_files).lower()

    def count_sources(aliases: List[str]) -> List[str]:
        return [
            str(p.relative_to(KB_ROOT))
            for p in kb_files
            if any(a.lower() in str(p).lower() for a in aliases)
        ][:8]

    items: List[Dict[str, Any]] = []
    for method_name, aliases in get_method_aliases().items():
        sources = count_sources(aliases)
        phases = get_method_phases().get(method_name, ["discover", "solve", "validate"])
        items.append({
            "id": re.sub(r"[^a-zA-Z0-9一-鿿]+", "_", method_name).strip("_").lower(),
            "name": method_name,
            "type": "method",
            "source_count": len(sources),
            "sources": sources,
            "phases": phases,
            "detected_by": "platform_card",
        })

    items.sort(key=lambda item: (-item["source_count"], item["name"]))
    items = _enrich_items_from_cards(items)

    METHODOLOGY_CATALOG_PATH.parent.mkdir(parents=True, exist_ok=True)
    METHODOLOGY_CATALOG_PATH.write_text(
        json.dumps(
            {"built_at": time.strftime("%Y-%m-%d %H:%M:%S"), "items": items},
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return items


def classify_methodologies_with_llm(
    candidates: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    if not candidates:
        return []
    config = load_llm_config()
    if not config.api_key:
        return [item for item in candidates if item.get("phases")]

    client, provider = _build_llm_client_and_provider(config)
    decisions: Dict[str, Dict[str, Any]] = {}

    for offset in range(0, min(len(candidates), 90), 25):
        compact_candidates = [
            {
                "id": item["id"],
                "name": item["name"],
                "type": item.get("type", "method"),
                "source_count": item.get("source_count", 0),
                "sources": item.get("sources", [])[:2],
                "rule_phases": item.get("phases", []),
            }
            for item in candidates[offset : offset + 25]
        ]

        prompt = f"""请判断下列研究方法、研究理论或模型是否适用于工程管理硕士论文研究框架，并分配到研究阶段。

阶段只能从以下三类选择：
discover = 发现问题
solve = 解决问题
validate = 验证问题

判断规则：
1. 只保留能明确用于论文研究设计的方法、理论或模型。
2. 如果不适用于任何阶段、名称太泛、无法判断、只是行业对象或资料名，直接丢弃。
3. 一个条目可以属于多个阶段。
4. 研究理论如果主要提供分析框架或改进框架，也可以归入 solve；如果主要支撑评价或检验，归入 validate；如果主要支撑诊断或问题识别，归入 discover。
5. 只输出 JSON，不要 Markdown。

候选项：
{json.dumps(compact_candidates, ensure_ascii=False)}

JSON 结构：
{{"items":[{{"id":"候选id","phases":["discover"],"reason":"一句话说明适配原因"}}]}}
"""
        try:
            response = _llm_create_message(
                client, provider, config,
                system="你是工程管理硕士论文研究方法与理论框架评审专家，擅长判断方法适用阶段。",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=_token_budget(config, "small"),
            )
            raw = _extract_json_object(_llm_response_text(response, provider))
        except Exception:
            continue
        for item in raw.get("items", []):
            phases = [
                phase
                for phase in item.get("phases", [])
                if phase in {"discover", "solve", "validate"}
            ]
            if phases:
                decisions[item.get("id")] = {
                    "phases": phases,
                    "reason": strip_markdown(item.get("reason", "")),
                }

    if not decisions:
        # LLM调用失败或未返回有效决策时，使用规则阶段作为后备
        for item in candidates:
            if item.get("phases"):
                item["llm_reason"] = "基于规则阶段分配（LLM未返回有效结果）"
                item["detected_by"] = f"{item.get('detected_by', 'local')}+rule_fallback"
        return sorted(
            candidates,
            key=lambda item: (-item["source_count"], item["type"], item["name"]),
        )

    filtered = []
    for item in candidates:
        decision = decisions.get(item["id"])
        if not decision or not decision["phases"]:
            continue
        item["phases"] = decision["phases"]
        item["llm_reason"] = (
            decision["reason"] or "大模型确认该条目适用于所列研究阶段。"
        )
        item["detected_by"] = f"{item.get('detected_by', 'local')}+llm"
        filtered.append(item)
    if filtered:
        return sorted(
            filtered,
            key=lambda item: (-item["source_count"], item["type"], item["name"]),
        )
    return []


def _matches_keywords(text: str, keywords: List[str]) -> bool:
    lowered = text.lower()
    return any(keyword.lower() in lowered for keyword in keywords)


def _iter_reference_source_files() -> List[Path]:
    if not KB_ROOT.exists():
        return []
    files = []
    for path in KB_ROOT.rglob("*"):
        if not path.is_file():
            continue
        try:
            path.relative_to(KB_CONVERTED_ROOT)
            continue
        except ValueError:
            pass
        if path.name.startswith("."):
            continue
        if path.suffix.lower() not in REFERENCE_SOURCE_EXTENSIONS:
            continue
        files.append(path)
    return sorted(files)


def _converted_markdown_path(source_path: Path) -> Path:
    relative = source_path.relative_to(KB_ROOT)
    return KB_CONVERTED_ROOT / relative.with_suffix(".md")


def _file_digest(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _markdown_front_matter(source_path: Path, digest: str, metadata: Optional[Dict[str, Any]] = None) -> str:
    relative = str(source_path.relative_to(KB_ROOT))
    lines = [
        "---",
        f"source_path: {relative}",
        f"source_type: {source_path.suffix.lower().lstrip('.')}",
        f"source_sha256: {digest}",
        f"source_size: {source_path.stat().st_size}",
        f"converted_at: {time.strftime('%Y-%m-%d %H:%M:%S')}",
    ]
    if metadata:
        lines.append(f"category: {metadata.get('category', 'unknown')}")
        if metadata.get("research_direction"):
            lines.append(f"research_direction: {metadata['research_direction']}")
        if metadata.get("methodologies"):
            lines.append(f"methodologies: {', '.join(metadata['methodologies'])}")
        if metadata.get("title"):
            lines.append(f"title: {metadata['title']}")
        if metadata.get("tags"):
            lines.append(f"tags: {', '.join(metadata['tags'])}")
    lines.append("---")
    lines.append("")
    return "\n".join(lines)


def _converted_file_is_current(output_path: Path, digest: str) -> bool:
    if not output_path.exists():
        return False
    try:
        head = output_path.read_text(encoding="utf-8", errors="replace")[:800]
    except Exception:
        return False
    return f"source_sha256: {digest}" in head


def _source_path_from_markdown(text: str, fallback: str) -> str:
    match = re.search(r"(?m)^source_path:\s*(.+?)\s*$", text[:1000])
    return match.group(1).strip() if match else fallback


def convert_references_to_markdown(task_id: str = "") -> Dict[str, Any]:
    files = _iter_reference_source_files()
    total = len(files)
    KB_CONVERTED_ROOT.mkdir(parents=True, exist_ok=True)
    converted = 0
    reused = 0
    skipped = 0
    errors = []
    category_stats: Dict[str, Dict[str, int]] = {}

    if task_id:
        task_log(task_id, f"开始转换知识库文件为 Markdown（分类存储），共 {total} 个")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 5

    for index, source_path in enumerate(files, 1):
        try:
            digest = _file_digest(source_path)

            category_info = detect_category(source_path)
            output_dir = categorize_output_dir(category_info)
            output_path = output_dir / f"{source_path.stem}.md"

            cat = category_info.get("category", "other")
            category_stats.setdefault(cat, {"total": 0, "converted": 0, "reused": 0})

            if _converted_file_is_current(output_path, digest):
                reused += 1
                category_stats[cat]["total"] += 1
                category_stats[cat]["reused"] += 1
            else:
                text = extract_text(source_path)
                metadata = build_rich_metadata(source_path, text)
                output_path.parent.mkdir(parents=True, exist_ok=True)
                output_path.write_text(
                    _markdown_front_matter(source_path, digest, metadata) + text,
                    encoding="utf-8",
                )
                converted += 1
                category_stats[cat]["total"] += 1
                category_stats[cat]["converted"] += 1
        except Exception as exc:
            skipped += 1
            errors.append(
                {
                    "path": str(source_path.relative_to(KB_ROOT)),
                    "error": str(exc),
                }
            )

        if task_id and (index % 10 == 0 or index == total):
            task_log(
                task_id,
                f"Markdown 转换 {index}/{total}：新增/更新 {converted}，复用 {reused}，跳过 {skipped}",
            )
            with TASK_LOCK:
                TASKS[task_id]["progress"] = min(29, 5 + int(index / max(total, 1) * 24))

    return {
        "source_files": total,
        "converted_files": converted,
        "reused_files": reused,
        "skipped_files": skipped,
        "converted_root": str(KB_CONVERTED_ROOT.relative_to(PROJECT_ROOT)),
        "category_stats": category_stats,
        "errors": errors[:20],
    }


def _citation_language(text: str) -> str:
    ascii_chars = sum(1 for char in text if ord(char) < 128)
    chinese_chars = sum(1 for char in text if "\u4e00" <= char <= "\u9fff")
    return "en" if ascii_chars > chinese_chars * 1.5 else "zh"


def _reference_section(text: str) -> str:
    if not text:
        return ""

    lines = text.splitlines()
    for index, line in enumerate(lines):
        stripped = line.strip()
        if re.fullmatch(r"(参考文献|References|REFERENCES|Bibliography)", stripped):
            return "\n".join(lines[index + 1 :]).strip()

    match = re.search(
        r"(?mi)^\s*(参考文献|References|REFERENCES|Bibliography)\s*$", text
    )
    if match:
        return text[match.end() :].strip()
    return ""


def _clean_reference_candidate(entry: str) -> str:
    cleaned = re.sub(r"^\s*(?:\[\d+\]|\[\d+\)|\d+[.)]|0\d+)\s*", "", entry)
    cleaned = re.sub(r"##\s*Page\s+\d+", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    cleaned = re.sub(r"\s+([,.;:，。；：])", r"\1", cleaned)
    cleaned = re.sub(r"([（(])\s+", r"\1", cleaned)
    cleaned = re.sub(r"\s+([）)])", r"\1", cleaned)
    return cleaned


def _looks_like_reference(entry: str) -> bool:
    if not entry or len(entry) < 18 or len(entry) > 700:
        return False
    if re.search(r"(第\s*\d+\s*章|^表\s*\d|Table\s+\d|续表|Page\s+\d)", entry):
        return False
    if re.search(
        r"(步骤|理由如下|综上|本文将|指标体系|风险因素|主要风险因素|图片截图|数据来自|资料来源)",
        entry,
    ):
        return False
    # AHP 问卷/调查表：含填空下划线的条目不是参考文献
    if re.search(r"_{5,}", entry):
        return False
    # 问卷题干：重要性比较、您的xxx 等
    if re.search(r"(经验年限|重要程度比较|相对于.*重要性|相对于.*重要)", entry):
        return False
    # OCR 乱码：同一短语重复 5 次以上
    if re.search(r"(.{6,}?)\1{4,}", entry):
        return False
    if not re.search(r"(19|20)\d{2}", entry):
        return False
    if re.search(r"\[[JMDCS]\]|\[EB/OL\]|\[J/OL\]|\[N\]", entry, re.IGNORECASE):
        return True
    if re.search(r"\b(journal|review|management|science|press|routledge)\b", entry, re.I):
        return True
    if re.search(r"(出版社|学报|研究|管理|科学|大学|标准|报)", entry):
        return True
    return False


def _extract_reference_entries(text: str, max_entries: int = 10) -> List[str]:
    section = _reference_section(text)
    if not section:
        return []

    section = re.sub(r"(?<!^)(\[\d+\])", r"\n\1", section)
    section = re.sub(r"(?<!^)(?<!\d)(\d+[.)]\s*[\u4e00-\u9fffA-Za-z])", r"\n\1", section)
    lines = [line.strip() for line in section.splitlines() if line.strip()]
    entries: List[str] = []
    buffer = ""

    for line in lines:
        if re.match(
            r"^(?:\[\d+\]|\[\d+\)|\d+\.\s*|\d+\)\s*|第[一二三四五六七八九十]+章?\b)",
            line,
        ):
            if buffer:
                entries.append(buffer.strip())
            buffer = line
        else:
            if buffer:
                buffer += " " + line
            else:
                buffer = line

    if buffer:
        entries.append(buffer.strip())

    cleaned_entries = [
        cleaned
        for cleaned in (_clean_reference_candidate(entry) for entry in entries)
        if _looks_like_reference(cleaned)
    ]

    return cleaned_entries[:max_entries]


def build_citation_index(task_id: str = "") -> Dict[str, Any]:
    source_root = KB_CONVERTED_ROOT if KB_CONVERTED_ROOT.exists() else KB_ROOT
    files = [
        path
        for path in (source_root.rglob("*") if source_root.exists() else [])
        if path.is_file()
        and path.suffix.lower() in {".md", ".txt"}
        and not path.name.startswith(".")
    ]
    total_files = len(files)
    if task_id:
        task_log(task_id, f"=== 开始重建引用索引，共 {total_files} 个文件 ===")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 5

    # ==== Phase 0: 向量搜索为每篇论文分配研究方向 ====
    if task_id:
        task_log(task_id, "Phase 0: 向量搜索分配论文研究方向...")
    paper_directions: Dict[str, Set[str]] = {}
    dir_name_map = {
        "quality_management": "质量管理",
        "risk_management": "风险管理",
        "schedule_management": "进度管理",
        "requirements_management": "需求管理",
        "process_optimization": "流程优化",
        "cost_management": "成本管理",
        "supply_chain_logistics": "供应链与物流",
    }

    for dir_id, keywords in CITATION_DIRECTIONS.items():
        dir_label = dir_name_map.get(dir_id, dir_id)
        query = f"{dir_label} {' '.join(keywords[:4])} 工程管理 硕士论文"
        try:
            results = search_knowledge_base(query, limit=60).get("results", [])
        except Exception:
            results = []
        for item in results:
            paper_path = item.get("path", "")
            if paper_path:
                if paper_path not in paper_directions:
                    paper_directions[paper_path] = set()
                paper_directions[paper_path].add(dir_id)
        if task_id:
            task_log(task_id, f"  {dir_label}: 向量匹配 {len(results)} 篇论文")

    total_matched = len(paper_directions)
    if task_id:
        task_log(task_id, f"向量搜索完成：{total_matched} 篇论文获得方向归属")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 30

    # ==== Phase 1: 逐论文提取引用 ====
    if task_id:
        task_log(task_id, "Phase 1: 逐论文提取参考文献...")
    entries = []
    processed = 0
    direction_counts: Dict[str, int] = {d: 0 for d in CITATION_DIRECTIONS}
    direction_counts["general"] = 0

    for path in files:
        processed += 1
        relative = str(path.relative_to(source_root))
        kb_relative = str(path.relative_to(KB_ROOT)) if path.is_relative_to(KB_ROOT) else relative

        # 向量库中 path 格式可能是相对于 PROJECT_ROOT，尝试多种匹配
        paper_dirs: Set[str] = set()
        for candidate in (
            kb_relative,
            f"knowledge_base/references/{kb_relative}",
            f"knowledge_base/references/converted/{relative}",
        ):
            if candidate in paper_directions:
                paper_dirs = paper_directions[candidate]
                break
        # 回退：向量没匹配到时用文件名关键词
        if not paper_dirs:
            haystack = f"{relative} {path.stem}"
            paper_dirs = {
                d for d, kw in CITATION_DIRECTIONS.items() if _matches_keywords(haystack, kw)
            }
        directions = list(paper_dirs) if paper_dirs else ["general"]
        for d in directions:
            direction_counts[d] = direction_counts.get(d, 0) + 1

        # 读取论文内容
        try:
            text = extract_text(path)
        except Exception:
            text = ""

        # 方法标记：从论文正文匹配（前 8000 字符），不是文件名
        content_sample = text[:8000] if text else ""
        _method_keywords_map = _get_method_keywords_map()
        paper_methods = [
            method
            for method, keywords in _method_keywords_map.items()
            if _matches_keywords(content_sample, keywords)
        ]

        # 提取参考文献
        reference_lines: List[str] = []
        try:
            reference_lines = _extract_reference_entries(text, max_entries=10)
        except Exception:
            reference_lines = []
        source_path = _source_path_from_markdown(text, relative)

        if reference_lines:
            for index, formatted in enumerate(reference_lines, 1):
                # 引用级别方法标记：从引用文本自身匹配
                ref_methods = [
                    method
                    for method, keywords in _method_keywords_map.items()
                    if _matches_keywords(formatted, keywords)
                ]
                entries.append(
                    {
                        "id": uuid.uuid5(uuid.NAMESPACE_URL, f"{source_path}-{index}").hex,
                        "title": formatted[:120],
                        "formatted": formatted,
                        "source_path": source_path,
                        "converted_path": kb_relative,
                        "language": _citation_language(formatted),
                        "type": "参考文献",
                        "directions": directions,
                        "methods": ref_methods or paper_methods,
                        "source": "local_reference",
                        "verify_status": "本地知识库",
                    }
                )

        if task_id and (processed % 15 == 0 or processed == total_files):
            task_log(task_id, f"  已处理 {processed}/{total_files} 篇，累计 {len(entries)} 条引用")
            with TASK_LOCK:
                TASKS[task_id]["progress"] = min(
                    95, 30 + int(processed / max(total_files, 1) * 65)
                )

    payload = {
        "built_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "entry_count": len(entries),
        "entries": entries,
        "direction_stats": direction_counts,
    }
    CITATION_INDEX_PATH.parent.mkdir(parents=True, exist_ok=True)
    CITATION_INDEX_PATH.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    if task_id:
        task_log(task_id, f"=== 索引重建完成：{len(entries)} 条引用，分布：{json.dumps(direction_counts, ensure_ascii=False)} ===")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 100
    return payload


def initialize_local_knowledge_base(task_id: str = "") -> Dict[str, Any]:
    def _log_vector_progress(file_index: int, total_files: int, file_path: Path) -> None:
        if not task_id or total_files <= 0:
            return
        if file_index % 10 != 0 and file_index != total_files:
            return
        task_log(
            task_id,
            f"向量库已处理 {file_index}/{total_files} 个文件：{file_path.name}",
        )
        with TASK_LOCK:
            TASKS[task_id]["progress"] = min(54, 30 + int(file_index / total_files * 24))

    conversion = convert_references_to_markdown(task_id=task_id)
    if task_id:
        task_log(
            task_id,
            f"Markdown 转换完成：新增/更新 {conversion.get('converted_files', 0)}，"
            f"复用 {conversion.get('reused_files', 0)}，跳过 {conversion.get('skipped_files', 0)}",
        )
        category_stats = conversion.get("category_stats", {})
        if category_stats:
            parts = [f"{cat}({stats['total']})" for cat, stats in sorted(category_stats.items())]
            task_log(task_id, f"分类统计: {', '.join(parts)}")
        task_log(task_id, "开始基于分类 Markdown 层重建本地向量库")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 30
    vector_index = build_index(
        source_dirs=[
            "knowledge_base/references/converted/papers",
            "knowledge_base/references/converted/methodologies",
            "knowledge_base/references/converted/templates",
            "skills",
        ],
        reset=True,
        progress_callback=_log_vector_progress if task_id else None,
    )
    if task_id:
        task_log(
            task_id,
            f"向量库构建完成：{vector_index.get('documents', 0)} 个文档，"
            f"{vector_index.get('chunks', 0)} 个文本块，跳过 {vector_index.get('skipped_files', 0)} 个文件",
        )
        task_log(task_id, "开始构建论文大纲索引（按研究方向分类）")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 55
    outline_index = build_outline_index()
    if task_id:
        task_log(
            task_id,
            f"大纲索引构建完成：{outline_index.get('total', 0)} 篇论文，"
            f"{len(outline_index.get('directions', {}))} 个研究方向",
        )
    if task_id:
        task_log(task_id, "开始构建研究方向模板")
    templates = build_domain_templates()
    if task_id:
        task_log(
            task_id,
            f"方向模板构建完成：{len(templates.get('templates', []))} 个模板",
        )
        task_log(task_id, "开始构建引用索引")
    citation_index = build_citation_index(task_id=task_id)
    if task_id:
        task_log(task_id, "本地知识库初始化完成")
    return {
        "conversion": conversion,
        "vector_index": vector_index,
        "outline_index": outline_index,
        "templates": templates,
        "citation_index": citation_index,
        "initialized_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }


def _run_knowledge_base_init_task(task_id: str) -> None:
    with TASK_LOCK:
        TASKS[task_id].update(
            {"status": "running", "message": "正在初始化本地知识库...", "progress": 0}
        )
    try:
        task_log(task_id, "任务开始：本地知识库初始化")
        result = initialize_local_knowledge_base(task_id=task_id)
        with TASK_LOCK:
            TASKS[task_id].update(
                {
                    "status": "done",
                    "message": "初始化完成",
                    "result": result,
                    "finished_at": time.time(),
                    "progress": 100,
                }
            )
    except Exception as exc:
        with TASK_LOCK:
            TASKS[task_id].update(
                {
                    "status": "error",
                    "message": str(exc),
                    "finished_at": time.time(),
                }
            )


def start_knowledge_base_init_task() -> str:
    with TASK_LOCK:
        for existing_id, task in TASKS.items():
            if (
                task.get("kind") == KB_INIT_TASK_KIND
                and task.get("status") in {"queued", "running"}
            ):
                return existing_id

        task_id = uuid.uuid4().hex
        TASKS[task_id] = {
            "kind": KB_INIT_TASK_KIND,
            "permission_menu": TASK_PERMISSION_BY_KIND[KB_INIT_TASK_KIND],
            "status": "queued",
            "message": "初始化任务已创建",
            "logs": [
                {
                    "time": time.strftime("%H:%M:%S"),
                    "message": "初始化任务已创建，等待后台执行",
                }
            ],
            "created_at": time.time(),
            "progress": 0,
        }
    thread = threading.Thread(
        target=_run_knowledge_base_init_task,
        args=(task_id,),
        daemon=True,
    )
    thread.start()
    return task_id


def load_citation_index() -> Dict[str, Any]:
    if not CITATION_INDEX_PATH.exists():
        return build_citation_index()
    try:
        return json.loads(CITATION_INDEX_PATH.read_text(encoding="utf-8"))
    except Exception:
        return build_citation_index()


def filter_local_citations(
    direction: str, methods: List[str], limit: int = 12
) -> List[Dict[str, Any]]:
    index = load_citation_index()
    method_set = set(methods or [])
    scored: List[Tuple[int, Dict[str, Any]]] = []
    for entry in index.get("entries", []):
        score = 0
        if direction and direction in entry.get("directions", []):
            score += 4
        score += len(method_set.intersection(set(entry.get("methods", [])))) * 3
        if entry.get("language") == "zh":
            score += 1
        if score > 0:
            scored.append((score, entry))
    scored.sort(key=lambda item: (-item[0], item[1].get("title", "")))
    return [entry for _, entry in scored[:limit]]


def _normalize_citations(raw_items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    normalized = []
    seen = set()
    for index, item in enumerate(raw_items, 1):
        title = strip_markdown(str(item.get("title", ""))).strip()
        if not title or title in seen:
            continue
        seen.add(title)
        normalized.append(
            {
                "id": item.get("id")
                or uuid.uuid5(uuid.NAMESPACE_URL, f"{title}-{index}").hex,
                "card_id": item.get("card_id", ""),
                "title": title,
                "authors": strip_markdown(str(item.get("authors", ""))).strip(),
                "year": str(item.get("year", "")).strip(),
                "language": item.get("language")
                if item.get("language") in {"zh", "en"}
                else _citation_language(title),
                "type": strip_markdown(str(item.get("type") or item.get("ref_type", "期刊/图书"))).strip(),
                "formatted": strip_markdown(str(item.get("formatted", ""))).strip()
                or strip_markdown(str(item.get("gbt7714", ""))).strip()
                or title,
                "directions": item.get("directions", []),
                "methods": item.get("methods", []),
                "reason": strip_markdown(str(item.get("reason", ""))).strip(),
                "source_path": item.get("source_path", ""),
                "source": item.get("source", "llm_supplement"),
                "verify_status": item.get("verify_status", "LLM补充，需核验"),
            }
        )
    return normalized


def _extract_refs_from_kb_papers(
    query: str, limit: int = 3, topic: str = "", paper_limit: int = 3
) -> List[Dict[str, Any]]:
    """搜索向量知识库，读取匹配论文并提取参考文献条目。"""
    all_refs: List[Dict[str, Any]] = []
    seen_papers: Set[str] = set()
    try:
        results = search_knowledge_base(query, limit=paper_limit).get("results", [])
    except Exception:
        return all_refs

    for item in results:
        paper_path_str = item.get("path", "")
        if not paper_path_str or paper_path_str in seen_papers:
            continue
        seen_papers.add(paper_path_str)

        # 找到对应的 markdown 文件
        # 向量库中的 path 是相对于 PROJECT_ROOT 的路径
        paper_path = PROJECT_ROOT / paper_path_str
        if not paper_path.exists():
            # 尝试去掉可能重复的前缀
            if paper_path_str.startswith("knowledge_base/"):
                paper_path = PROJECT_ROOT / paper_path_str
            else:
                paper_path = KB_ROOT / paper_path_str
        if not paper_path.exists():
            # 尝试 converted 目录中同名 .md 文件
            stem = Path(paper_path_str).stem
            candidates = list(KB_CONVERTED_ROOT.rglob(f"{stem}.md"))
            if candidates:
                paper_path = candidates[0]
            else:
                continue

        try:
            text = extract_text(paper_path)
        except Exception:
            continue

        entries = _extract_reference_entries(text, max_entries=8)
        for entry in entries:
            all_refs.append({
                "title": entry[:120],
                "formatted": entry,
                "source_path": paper_path_str,
                "source_title": item.get("title", ""),
                "source_score": round(item.get("score", 0), 4),
                "language": _citation_language(entry),
                "type": "参考文献",
                "source": "local_kb",
                "verify_status": "本地知识库提取",
            })

    return all_refs


# 研究工具类方法——从注册中心动态加载，cards DB 中 method_role=research_tool 的方法
def _get_research_tool_methods() -> Set[str]:
    from src.method_registry import get_research_tool_methods
    tools = get_research_tool_methods()
    if tools:
        return tools
    # 注册中心无 tool 标记时（cards DB 尚未更新 method_role）回退到最常见的工具方法
    return {"问卷调查法", "案例研究法", "文献研究法", "访谈法"}

# 方法关键词到搜索词的映射——从注册中心动态加载
def _get_method_content_keywords() -> Dict[str, List[str]]:
    from src.method_registry import get_method_keywords
    return get_method_keywords()


def _content_filter_cards(
    cards: List[Dict[str, Any]], methods: List[str]
) -> List[Dict[str, Any]]:
    """内容级过滤：引文内容必须包含方法关键词才算真正匹配。

    研究工具类方法（问卷、AHP、访谈等）跳过内容过滤——引文标题讨论的是研究主题而非
    研究工具本身，只要来源论文使用了该方法即可。
    """
    if not methods:
        return cards
    tool_methods = _get_research_tool_methods()
    subject_methods = [m for m in methods if m not in tool_methods]
    if not subject_methods:
        # 全是工具类方法，不做内容过滤
        return cards
    content_keywords = _get_method_content_keywords()
    search_texts = set()
    for m in subject_methods:
        for kw in content_keywords.get(m, [m]):
            search_texts.add(kw.lower())
    filtered = []
    for c in cards:
        haystack = (c.get("formatted", "") + " " + c.get("title", "")).lower()
        if any(kw in haystack for kw in search_texts):
            filtered.append(c)
    return filtered


def _normalize_title(title: str) -> str:
    """Normalize title for dedup comparison."""
    s = title.lower().strip()
    s = re.sub(r'[，,、､]\s*', ', ', s)   # All commas → ', '
    s = re.sub(r'\.\s*', '. ', s)         # Periods → '. '
    s = re.sub(r'([一-鿿])([a-zA-Z])', r'\1 \2', s)  # CJK-Latin boundary
    s = re.sub(r'([a-zA-Z])([一-鿿])', r'\1 \2', s)  # Latin-CJK boundary
    s = re.sub(r'\s+', ' ', s)            # Collapse whitespace
    s = re.sub(r'[\[\]【】（）()]', '', s)
    return s.strip().rstrip(',.;;.')


def _cards_to_citations(cards: List[Dict[str, Any]], source: str = "local") -> List[Dict[str, Any]]:
    """将 paper_store 引用卡片转为 API 引用格式。"""
    out = []
    seen = set()
    for c in cards:
        title = c.get("title", "")
        if not title:
            continue
        norm = _normalize_title(title)
        if norm in seen:
            continue
        seen.add(norm)
        seen.add(title)  # Also track raw title for belt-and-suspenders
        out.append({
            "id": c.get("card_id", ""),
            "title": title,
            "formatted": c.get("formatted", ""),
            "authors": c.get("authors", ""),
            "year": c.get("year", ""),
            "language": c.get("language", "zh"),
            "type": c.get("ref_type", "期刊文章"),
            "methods": c.get("methods", []),
            "directions": [c.get("direction_id", "")],
            "source": source,
            "source_path": c.get("source_paper_title", ""),
            "verify_status": "本地知识库提取" if source == "local" else "LLM补充",
            "quality_score": c.get("quality_score", 0.0),
        })
    return out


def generate_citations(payload: Dict[str, Any], task_id: str = "") -> Dict[str, Any]:
    from src.paper_store import get_cards_by_direction_and_methods, search_cards, get_paper_stats, _diversify_cards

    topic = payload.get("topic", "论文主题")
    project_context = persist_project_context(payload)
    direction = payload.get("direction") or ""
    direction_name = payload.get("direction_name") or direction or "工程管理"
    methods = payload.get("methods") or []
    expected_count = max(10, min(150, int(payload.get("expected_count", 100) or 100)))

    # 配额分配：有方法时方向占 20%；纯方向检索时方向占 100%
    if methods:
        direction_quota = max(2, round(expected_count * 0.2))
        methods_quota_total = expected_count - direction_quota
    else:
        direction_quota = expected_count
        methods_quota_total = 0

    stats = get_paper_stats()
    if task_id:
        task_log(task_id, f"开始生成引用：方向={direction_name}，方法={len(methods)}个，期望 {expected_count} 条")
        task_log(task_id, f"配额：方向 {direction_quota} + 方法汇总 {methods_quota_total} 条")
        task_log(task_id, f"论文库：{stats['papers']} 篇论文，{stats['cards']} 张引用卡片")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 5

    # === 第一阶段：方向检索（先内容过滤，再多样性，保证每条引用都匹配方法） ===
    dir_cards_raw = get_cards_by_direction_and_methods(
        direction, methods, limit=300, verified_only=False,
    )
    dir_cards_filtered = _content_filter_cards(dir_cards_raw, methods)
    dir_cards = _diversify_cards(dir_cards_filtered, direction_quota * 2, max_per_source=2)
    dir_citations = _cards_to_citations(dir_cards, source="local_direction")
    if task_id:
        task_log(task_id, f"方向检索（内容级匹配+多样性）：{len(dir_citations)} 条（配额 {direction_quota}）")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 20

    # === 第二阶段：方法汇总检索（多方法合并查询，非逐方法单独查） ===
    method_citations_map: Dict[str, List[Dict[str, Any]]] = {}
    if methods:
        m_cards_raw = get_cards_by_direction_and_methods(
            direction, methods, limit=500,
        )
        m_cards_filtered = _content_filter_cards(m_cards_raw, methods)
        m_cards = _diversify_cards(m_cards_filtered, methods_quota_total * 2, max_per_source=4)
        m_cits = _cards_to_citations(m_cards, source="local_method")
        # 按方法拆分统计（仅用于日志，不影响合并检索）
        for method in methods:
            method_citations_map[method] = [
                c for c in m_cits if method in (c.get("methods") or [])
            ]
        if task_id:
            task_log(task_id, f"  方法汇总检索：{len(m_cits)} 条（配额 {methods_quota_total}）")
            for method in methods:
                task_log(task_id, f"    {method}: {len(method_citations_map.get(method, []))} 条")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 55
    else:
        m_cits = []

    # 合并去重：方向优先，方法补充
    dir_final = dir_citations[:direction_quota]
    used_titles = {c["title"] for c in dir_final}

    method_final = [c for c in m_cits if c["title"] not in used_titles][:methods_quota_total]
    for c in method_final:
        used_titles.add(c["title"])

    local_citations = dir_final + method_final

    if task_id:
        task_log(
            task_id,
            f"本地检索完成：方向引用 {len(dir_final)} 条 + 方法引用 {len(method_final)} 条 = {len(local_citations)} 条",
        )
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 55

    # === 第三阶段：LLM 补充 ===
    config = load_llm_config()
    llm_citations: List[Dict[str, Any]] = []
    llm_status = "未配置 API Key，仅返回本地引用候选"
    gap = expected_count - len(local_citations)

    if config.api_key and gap > 0:
        if task_id:
            task_log(task_id, f"本地缺口 {gap} 条，调用大模型补充：{config.model}")
            with TASK_LOCK:
                TASKS[task_id]["progress"] = 60
        client, provider = _build_llm_client_and_provider(config)

        methods_quota_desc = "、".join(
            f"{m}" for m in methods[:8]
        ) if methods else "无"

        prompt = f"""请为工程管理硕士论文补充参考文献候选清单。

论文主题：{topic}
项目背景与论文思路：
{project_context[:2400] or "用户未填写。"}

研究方向：{direction_name}（配额约 {direction_quota} 篇，含该领域经典理论与近期进展）
研究方法与配额：{methods_quota_desc}
期望总引用数量：约 {expected_count} 篇，当前本地已检索到 {len(local_citations)} 篇，需要补充约 {gap} 篇

本地已检索候选（不要重复这些）：
{json.dumps(local_citations[: min(20, len(local_citations))], ensure_ascii=False)}

要求：
1. 研究方向「{direction_name}」相关：提供该领域的中外经典理论书籍、综述文章、高被引用刊论文。
2. 每个研究方法：提供与该方法的理论基础、应用案例、近期（近3年）高水平期刊文章。
3. 不要重复本地已有的候选条目。
4. 严禁捏造。无法确认作者、题名、年份或来源真实性的条目不要输出。
5. formatted 字段必须使用 GB/T 7714—2015 参考文献格式。
6. 中英文比例约 5:5。
7. 只输出 JSON，不要 Markdown，不要解释。

JSON 结构：
{{"citations":[{{"formatted":"GB/T 7714—2015格式参考文献","title":"文献题名","authors":"作者","year":"年份","language":"zh或en","type":"图书或期刊文章","methods":["相关方法"],"reason":"研究方向/方法对应说明","verify_status":"需人工核验"}}]}}
"""
        try:
            response = _llm_create_message(
                client, provider, config,
                system="你是工程管理硕士论文参考文献规划助手，擅长补充真实可核验的学术参考文献。只输出你能确认存在的文献，不确定的不要输出。",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max(config.max_tokens, 6000),
            )
            raw = _extract_json_object(_llm_response_text(response, provider))
            llm_citations = _normalize_citations(raw.get("citations", []))
            for c in llm_citations:
                c["source"] = "llm_supplement"
                c["verify_status"] = "LLM补充，需核验"
            llm_status = f"LLM 补充返回 {len(llm_citations)} 条引用候选；建议导出前人工核验"
            if task_id:
                task_log(task_id, llm_status)
                with TASK_LOCK:
                    TASKS[task_id]["progress"] = 85
        except Exception as exc:
            llm_status = f"LLM 引用补充失败：{exc}"
            if task_id:
                task_log(task_id, llm_status)
                with TASK_LOCK:
                    TASKS[task_id]["progress"] = 80
    elif task_id:
        if gap <= 0:
            task_log(task_id, "本地检索已满足配额，跳过大模型补充")
        else:
            task_log(task_id, "未配置 API Key，跳过 LLM 补充")
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 80

    # === 第四阶段：LLM 去重 + 中英文平衡 ===
    if task_id:
        task_log(task_id, "合并去重，按配额裁剪...")

    llm_final = [c for c in llm_citations if c["title"] not in used_titles]
    for c in llm_final:
        used_titles.add(c["title"])

    merged_local = dir_final + method_final
    merged = merged_local + llm_final

    # 中英文平衡
    zh = [c for c in merged if c.get("language") == "zh"]
    en = [c for c in merged if c.get("language") == "en"]
    half = max(1, expected_count // 2)
    balanced = (zh[:half] + en[:half]) if en else merged[:expected_count]
    if len(balanced) < min(expected_count, len(merged)):
        used = {c["id"] for c in balanced}
        balanced.extend(
            [c for c in merged if c["id"] not in used][: expected_count - len(balanced)]
        )

    result = {
        "status": "ok",
        "message": llm_status,
        "citations": balanced[:expected_count],
        "local_citations": merged_local,
        "llm_citations": llm_final,
        "local_count": len(merged_local),
        "llm_count": len(llm_final),
        "direction_count": len(dir_final),
        "method_counts": {m: len(method_citations_map.get(m, [])) for m in methods},
        "index_entry_count": stats.get("cards", 0),
        "allocation": {
            "direction_quota": direction_quota,
            "methods_quota_total": methods_quota_total,
        },
    }
    save_workspace_value("citations", result["citations"])
    if task_id:
        task_log(
            task_id,
            f"引用生成完成：输出 {len(result['citations'])} 条（方向 {len(dir_final)} + 方法 {len(method_final)} + LLM {len(llm_final)}）",
        )
        with TASK_LOCK:
            TASKS[task_id]["progress"] = 100
    return result


def _classify_citations_batch(
    citations: List[Dict[str, Any]], batch_size: int = 20
) -> int:
    """LLM 逐条分类引用涉及的研究方法，使用方向限定的候选方法词表。

    每条引用提供：标题 + 来源论文章节 + 摘要上下文。
    候选方法词表限定为该研究方向下的方法，减少误分类。

    Returns: 成功分类的引用数量。
    """
    from src.paper_store import (
        get_method_candidates_for_direction,
        update_card_methods,
    )

    config = load_llm_config()
    config.model = "deepseek-v4-flash"
    if not config.api_key:
        raise RuntimeError("未配置 API Key")
    client, provider = _build_llm_client_and_provider(config)
    total_updated = 0

    for i in range(0, len(citations), batch_size):
        batch = citations[i : i + batch_size]

        # Group by direction to provide focused method candidates
        direction_groups: Dict[str, List[Dict[str, Any]]] = {}
        for c in batch:
            d = c.get("direction_label") or "未分类"
            if d not in direction_groups:
                direction_groups[d] = []
            direction_groups[d].append(c)

        for direction_label, group in direction_groups.items():
            method_candidates = get_method_candidates_for_direction(direction_label)
            # Build a compact candidate list for the prompt
            candidate_names = sorted(set(
                m["name"] for m in method_candidates
            ))
            candidate_hint = (
                f"候选方法词表（{direction_label}方向，共{len(candidate_names)}个）：\n"
                + "、".join(candidate_names[:120])
            )
            if len(candidate_names) > 120:
                candidate_hint += f"\n... 还有 {len(candidate_names) - 120} 个方法，如需可自行补充"

            # Build input data with richer context
            input_items = []
            for c in group:
                title = (c.get("title") or "")[:250]
                source_section = (c.get("source_section") or "")[:100]
                abstract = (c.get("paper_abstract") or "")[:300]
                input_items.append({
                    "card_id": c["card_id"],
                    "title": title,
                    "source_section": source_section,
                    "abstract_excerpt": abstract,
                })

            prompt = (
                f"你是工程管理研究方法分类专家。请分析以下参考文献，判断每条文献使用的研究方法。\n\n"
                f"研究方向：{direction_label}\n\n"
                f"{candidate_hint}\n\n"
                f"分类规则：\n"
                f"1. 优先从候选方法词表中选择，如果标题明确包含某方法的关键词，标注该方法\n"
                f"2. 标题中有明确的非词表方法（如某个专业领域特有方法），也可以标注，但方法名要规范\n"
                f"3. 标题中没有明确方法论关键词的，methods 返回空数组\n"
                f"4. source_section 和 abstract_excerpt 是辅助上下文——它们说明该引用出现在论文的哪个章节和什么研究背景中，可以帮助你判断\n"
                f"5. 每条引用可以标注多个方法（最多5个），按确定性排序\n\n"
                f"参考文献列表：\n{json.dumps(input_items, ensure_ascii=False, indent=2)}\n\n"
                f"只输出JSON数组（不要markdown代码块）：\n"
                f'[{{"card_id": "...", "methods": ["方法1"], "theories": []}}]'
            )

            try:
                response = _llm_create_message(
                    client, provider, config,
                    system="你是工程管理硕士论文研究方法分类专家。根据文献标题+上下文判断该文献涉及的研究方法，优先从给定的候选词表中选择。只标注明确的方法，不猜测。",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=min(config.max_tokens, 400 + len(group) * 100),
                    disable_thinking=True,
                )
                raw_text = _llm_response_text(response, provider).strip()
                raw_text = re.sub(r"^```(?:json)?\s*", "", raw_text)
                raw_text = re.sub(r"\s*```$", "", raw_text)
                if raw_text.startswith("["):
                    results = json.loads(raw_text)
                elif raw_text.startswith("{"):
                    obj = json.loads(raw_text)
                    results = obj if isinstance(obj, list) else obj.get("results", obj.get("data", []))
                else:
                    results = []
            except Exception:
                results = []

            llm_classified_ids: set = set()
            for item in results:
                card_id = item.get("card_id", "")
                methods = item.get("methods", []) or []
                theories = item.get("theories", []) or []
                if card_id and methods:
                    update_card_methods(card_id, methods, theories)
                    total_updated += 1
                    llm_classified_ids.add(card_id)

            # Mark cards that LLM checked but found no methods for,
            # so they don't get re-queried in the next batch (infinite loop guard).
            batch_ids = {c["card_id"] for c in group}
            for cid in batch_ids - llm_classified_ids:
                update_card_methods(cid, ["__none__"], [])

    return total_updated


def _run_classify_citations_task(task_id: str, payload: Dict[str, Any]) -> None:
    from src.paper_store import (
        count_citations_without_methods,
        get_citations_without_methods_enriched,
        keyword_prefilter_citations,
    )

    with TASK_LOCK:
        TASKS[task_id].update(
            {"status": "running", "message": "正在统计待分类的引用...", "progress": 0}
        )

    try:
        total = count_citations_without_methods()
        max_cards = payload.get("max_cards")
        target = min(total, max_cards) if max_cards else total
        task_log(task_id, f"共 {total} 条引用未分类" + (f"，本次限制 {max_cards} 条" if max_cards else ""))

        # ── Step 1: Keyword pre-filtering (fast, no LLM cost) ──
        with TASK_LOCK:
            TASKS[task_id].update({"message": "Step 1/2: 关键词预筛选（标题匹配方法名）...", "progress": 5})
        task_log(task_id, "Step 1/2: 关键词预筛选...")
        kw_classified = keyword_prefilter_citations(max_cards=target)
        task_log(task_id, f"  关键词预筛选完成：{kw_classified} 条通过标题关键词匹配到方法")

        remaining = count_citations_without_methods()
        task_log(task_id, f"  剩余 {remaining} 条引用需要 LLM 分类")

        if remaining == 0:
            task_log(task_id, "所有引用已分类完成！")
            with TASK_LOCK:
                TASKS[task_id].update({
                    "status": "done",
                    "message": f"分类完成：关键词匹配 {kw_classified} 条，LLM 分类 0 条",
                    "result": {"keyword_classified": kw_classified, "llm_classified": 0, "total": total},
                    "finished_at": time.time(),
                    "progress": 100,
                })
            return

        # ── Step 2: LLM classification for remaining ──
        with TASK_LOCK:
            TASKS[task_id].update({"message": "Step 2/2: LLM 分类（方向限定候选词表）...", "progress": 10})
        task_log(task_id, "Step 2/2: LLM 分类（使用方向限定候选词表 + 富文本上下文）...")

        batch_size = 20
        processed = 0
        updated = 0

        while processed < remaining:
            batch = get_citations_without_methods_enriched(limit=batch_size, offset=0)
            if not batch:
                break
            n = _classify_citations_batch(batch, batch_size=batch_size)
            updated += n
            processed += len(batch)

            pct = min(95, 10 + int(processed / max(remaining, 1) * 85))
            task_log(task_id, f"  批次 {processed}/{remaining}: 本批 {len(batch)} 条，LLM 标注 {n} 条有方法")
            with TASK_LOCK:
                TASKS[task_id].update({
                    "progress": pct,
                    "message": f"LLM分类 {processed}/{remaining}，标注 {updated} 条有方法",
                })

        final_remaining = count_citations_without_methods()
        total_classified = kw_classified + updated
        task_log(task_id, f"分类完成：关键词 {kw_classified} + LLM {updated} = {total_classified} 条，剩余 {final_remaining} 条无明确方法")
        with TASK_LOCK:
            TASKS[task_id].update({
                "status": "done",
                "message": f"分类完成：关键词 {kw_classified} 条 + LLM {updated} 条 = 共 {total_classified} 条已标注方法",
                "result": {
                    "keyword_classified": kw_classified,
                    "llm_classified": updated,
                    "remaining": final_remaining,
                    "total": total,
                    "processed": processed,
                },
                "finished_at": time.time(),
                "progress": 100,
            })
    except Exception as exc:
        with TASK_LOCK:
            TASKS[task_id].update({
                "status": "error",
                "message": str(exc),
                "finished_at": time.time(),
            })


_CLASSIFY_CITATIONS_TASK_KIND = "classify_citations"


def start_classify_citations_task(payload: Dict[str, Any]) -> str:
    task_id = uuid.uuid4().hex
    with TASK_LOCK:
        TASKS[task_id] = {
            "kind": _CLASSIFY_CITATIONS_TASK_KIND,
            "permission_menu": TASK_PERMISSION_BY_KIND.setdefault(
                _CLASSIFY_CITATIONS_TASK_KIND, "methods"
            ),
            "status": "queued",
            "message": "引用分类任务已创建",
            "logs": [{"time": time.strftime("%H:%M:%S"), "message": "引用分类任务已创建，等待后台执行"}],
            "created_at": time.time(),
            "progress": 0,
        }
    thread = threading.Thread(
        target=_run_classify_citations_task, args=(task_id, payload), daemon=True
    )
    thread.start()
    return task_id


# ── Scan-verify citations task ──

_SCAN_VERIFY_TASK_KIND = "scan_verify_citations"


def _start_scan_verify_task(payload: Dict[str, Any]) -> str:
    task_id = uuid.uuid4().hex
    with TASK_LOCK:
        TASKS[task_id] = {
            "kind": _SCAN_VERIFY_TASK_KIND,
            "permission_menu": TASK_PERMISSION_BY_KIND.setdefault(
                _SCAN_VERIFY_TASK_KIND, "paper_manager"
            ),
            "status": "queued",
            "message": "扫描校验任务已创建",
            "logs": [{"time": time.strftime("%H:%M:%S"), "message": "扫描校验任务已创建"}],
            "created_at": time.time(),
            "progress": 0,
        }
    thread = threading.Thread(
        target=_run_scan_verify_task, args=(task_id, payload), daemon=True
    )
    thread.start()
    return task_id


def _run_scan_verify_task(task_id: str, payload: Dict[str, Any]) -> None:
    from src.paper_store import (
        count_unverified_citations,
        get_unverified_citations_rich,
        get_method_candidates_for_direction,
        update_card_verified,
    )

    with TASK_LOCK:
        TASKS[task_id].update(
            {"status": "running", "message": "正在统计待校验引用...", "progress": 0}
        )

    try:
        total = count_unverified_citations()
        max_cards = payload.get("max_cards")
        target = min(total, max_cards) if max_cards else total
        task_log(task_id, f"共 {total} 条引用未校验" + (f"，本次限制 {max_cards} 条" if max_cards else ""))

        if target == 0:
            task_log(task_id, "没有需要校验的引用")
            with TASK_LOCK:
                TASKS[task_id].update({
                    "status": "done",
                    "message": "没有需要校验的引用",
                    "result": {"confirmed": 0, "fake": 0, "format_error": 0},
                    "finished_at": time.time(),
                    "progress": 100,
                })
            return

        config = load_llm_config()
        config.model = "deepseek-v4-flash"
        if not config.api_key:
            raise RuntimeError("未配置 API Key")
        client, provider = _build_llm_client_and_provider(config)

        # Direction candidates for classification
        direction_entries = [
            {"id": k, "label": v}
            for k, v in {"quality_management": "质量管理", "risk_management": "风险管理",
                          "schedule_management": "进度管理", "requirements_management": "需求管理",
                          "process_optimization": "流程优化", "cost_management": "成本管理",
                          "supply_chain_logistics": "供应链与物流"}.items()
        ]

        batch_size = 15
        processed = 0
        confirmed = 0
        fake = 0
        fmt_error = 0
        method_classified = 0
        dir_classified = 0

        while processed < target:
            batch = get_unverified_citations_rich(limit=batch_size, offset=0)
            if not batch:
                break

            # Group by existing direction_label for method candidates
            direction_groups: Dict[str, List[Dict[str, Any]]] = {}
            for c in batch:
                d = c.get("direction_label") or ""
                if d not in direction_groups:
                    direction_groups[d] = []
                direction_groups[d].append(c)

            for dir_label, group in direction_groups.items():
                method_candidates = get_method_candidates_for_direction(dir_label) if dir_label else get_method_candidates_for_direction("质量管理")
                candidate_names = sorted(set(m["name"] for m in method_candidates))

                # Truncate long fields for prompt
                input_items = []
                for c in group:
                    input_items.append({
                        "card_id": c["card_id"],
                        "formatted": (c.get("formatted") or "")[:500],
                        "title": (c.get("title") or "")[:250],
                        "authors": (c.get("authors") or "")[:100],
                        "year": c.get("year", ""),
                        "ref_type": c.get("ref_type", ""),
                        "language": c.get("language", "zh"),
                        "existing_direction": c.get("direction_label", ""),
                    })

                direction_hint = "\n".join(
                    f"- {d['id']}: {d['label']}" for d in direction_entries
                )

                prompt = (
                    f"你是工程管理硕士论文引用校验专家。对每条引用完成以下任务：\n\n"
                    f"1. **真实性校验**：判断该引用是否为真实存在的学术文献。\n"
                    f"   - verified=1：确认真实（标题/作者/期刊信息匹配）\n"
                    f"   - verified=-1：疑似虚假（信息无法对应、来源可疑）\n"
                    f"   - verified=-2：格式问题（引用存在但格式不符合 GB/T 7714-2015）\n"
                    f"2. **格式检查**：如 verified=-2，在 verification_note 中说明具体格式问题\n"
                    f"3. **方法分类**：判断文献涉及的研究方法，从候选词表选择\n"
                    f"4. **方向分类**：判断文献所属研究方向\n\n"
                    f"候选研究方向：\n{direction_hint}\n\n"
                    f"候选方法词表（共{len(candidate_names)}个）：\n"
                    + "、".join(candidate_names[:120]) + "\n\n"
                    f"参考文献列表：\n{json.dumps(input_items, ensure_ascii=False, indent=2)}\n\n"
                    f"只输出JSON数组（不要markdown代码块）：\n"
                    f'[{{"card_id": "...", "verified": 1, "verification_note": "", '
                    f'"methods": ["方法1"], "theories": [], '
                    f'"direction_id": "quality_management", "direction_label": "质量管理"}}]'
                )

                try:
                    response = _llm_create_message(
                        client, provider, config,
                        system="你是工程管理硕士论文引用校验与分类专家。根据引用文本校验真实性、格式规范、研究方法、研究方向。",
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=min(config.max_tokens, 500 + len(group) * 200),
                        disable_thinking=True,
                    )
                    raw_text = _llm_response_text(response, provider).strip()
                    raw_text = re.sub(r"^```(?:json)?\s*", "", raw_text)
                    raw_text = re.sub(r"\s*```$", "", raw_text)
                    if raw_text.startswith("["):
                        results = json.loads(raw_text)
                    elif raw_text.startswith("{"):
                        obj = json.loads(raw_text)
                        results = obj if isinstance(obj, list) else obj.get("results", obj.get("data", []))
                    else:
                        results = []
                except Exception:
                    results = []

                batch_ids = {c["card_id"] for c in group}
                for item in results:
                    cid = item.get("card_id", "")
                    if not cid or cid not in batch_ids:
                        continue
                    v = item.get("verified", 0)
                    note = item.get("verification_note", "") or ""
                    methods = item.get("methods", []) or []
                    theories = item.get("theories", []) or []
                    did = item.get("direction_id", "") or ""
                    dl = item.get("direction_label", "") or ""

                    update_card_verified(cid, v, note, methods, theories, did, dl)

                    if v == -1:
                        fake += 1
                    elif v == -2:
                        fmt_error += 1
                    else:
                        confirmed += 1
                    if methods:
                        method_classified += 1
                    if did:
                        dir_classified += 1
                    batch_ids.discard(cid)

                # Mark unprocessed cards as checked (no useful info found)
                for cid in batch_ids:
                    update_card_verified(cid, 1, "", [], [], "", "")

            processed += len(batch)
            pct = min(95, 10 + int(processed / max(target, 1) * 85))
            task_log(task_id, f"  批次 {processed}/{target}: 确认{confirmed} 虚假{fake} 格式问题{fmt_error}")
            with TASK_LOCK:
                TASKS[task_id].update({
                    "progress": pct,
                    "message": f"扫描校验 {processed}/{target}，确认{confirmed} 虚假{fake} 格式{fmt_error}",
                })

        task_log(task_id, f"校验完成：确认 {confirmed} 条，虚假 {fake} 条，格式问题 {fmt_error} 条，方法归类 {method_classified}，方向归类 {dir_classified}")
        with TASK_LOCK:
            TASKS[task_id].update({
                "status": "done",
                "message": f"校验完成：确认{confirmed}、虚假{fake}、格式问题{fmt_error}",
                "result": {
                    "confirmed": confirmed,
                    "fake": fake,
                    "format_error": fmt_error,
                    "method_classified": method_classified,
                    "direction_classified": dir_classified,
                    "total": total,
                    "processed": processed,
                },
                "finished_at": time.time(),
                "progress": 100,
            })
    except Exception as exc:
        with TASK_LOCK:
            TASKS[task_id].update({
                "status": "error",
                "message": str(exc),
                "finished_at": time.time(),
            })


def _run_citation_generate_task(task_id: str, payload: Dict[str, Any]) -> None:
    with TASK_LOCK:
        TASKS[task_id].update(
            {"status": "running", "message": "正在生成引用...", "progress": 0}
        )
    try:
        task_log(task_id, "任务开始：引用生成")
        result = generate_citations(payload, task_id=task_id)
        with TASK_LOCK:
            TASKS[task_id].update(
                {
                    "status": "done",
                    "message": "引用生成完成",
                    "result": result,
                    "finished_at": time.time(),
                    "progress": 100,
                }
            )
    except Exception as exc:
        with TASK_LOCK:
            TASKS[task_id].update(
                {
                    "status": "error",
                    "message": str(exc),
                    "finished_at": time.time(),
                }
            )


def start_citation_generate_task(payload: Dict[str, Any]) -> str:
    task_id = uuid.uuid4().hex
    with TASK_LOCK:
        TASKS[task_id] = {
            "kind": CITATION_GENERATE_TASK_KIND,
            "permission_menu": TASK_PERMISSION_BY_KIND[CITATION_GENERATE_TASK_KIND],
            "status": "queued",
            "message": "引用生成任务已创建",
            "logs": [
                {
                    "time": time.strftime("%H:%M:%S"),
                    "message": "引用生成任务已创建，等待后台执行",
                }
            ],
            "created_at": time.time(),
            "progress": 0,
        }
    thread = threading.Thread(
        target=_run_citation_generate_task,
        args=(task_id, payload),
        daemon=True,
    )
    thread.start()
    return task_id


def _load_valid_method_map() -> Dict[str, str]:
    """Build a map of every known identifier (id, name, short_name) → canonical card id.

    This lets us normalize mixed-format phase_methods data back to clean IDs.
    """
    import src.card_builder as card_builder
    mapping: Dict[str, str] = {}
    db_path = card_builder.CARDS_DB
    if not db_path.exists():
        return mapping
    conn = None
    try:
        conn = sqlite3.connect(str(db_path))
        rows = conn.execute(
            "SELECT id, name, short_name FROM cards WHERE type='method_card' AND scope='platform'"
        ).fetchall()
        for card_id, name, short_name in rows:
            canonical = str(card_id).strip()
            if not canonical:
                continue
            mapping[canonical] = canonical
            for v in (name, short_name):
                if v:
                    key = str(v).strip()
                    if key:
                        mapping[key] = canonical
            # Also map lowercase variants for case-insensitive matching
            for v in (card_id, name, short_name):
                if v:
                    key_l = str(v).strip().lower()
                    if key_l and key_l not in mapping:
                        mapping[key_l] = canonical
    except Exception:
        pass
    finally:
        if conn:
            conn.close()
    return mapping


def _normalize_method_identifier(value: str, mapping: Dict[str, str]) -> Optional[str]:
    """Convert a single method identifier (could be an ID, name, short_name, or mismatch) to a canonical card ID.

    Returns None if no matching card is found.
    """
    if not value or not isinstance(value, str):
        return None
    v = value.strip()
    if not v:
        return None
    # Direct match (exact string match on ID, name, or short_name)
    if v in mapping:
        return mapping[v]
    # Case-insensitive match
    v_lower = v.lower()
    if v_lower in mapping:
        return mapping[v_lower]
    # Strip "prompt_" prefix and try matching
    if v.startswith("prompt_"):
        base = v[len("prompt_"):]
        if base in mapping:
            return mapping[base]
        base_lower = base.lower()
        if base_lower in mapping:
            return mapping[base_lower]
    # Separator-normalized match: normalize non-alphanumeric chars to single space.
    # Catches e.g. "层次分析法_ahp" vs "层次分析法 ahp" but NOT "method_wbs" vs
    # "method_wbs_analysis" (different words → different IDs).
    import re
    v_norm = re.sub(r'[^a-z0-9一-鿿]+', ' ', v_lower).strip()
    for key, canonical in mapping.items():
        key_norm = re.sub(r'[^a-z0-9一-鿿]+', ' ', key.lower()).strip()
        if v_norm == key_norm:
            return canonical
    return None


def _normalize_phase_methods(phase_methods: Dict[str, List[str]]) -> Dict[str, List[str]]:
    """Convert all values in phase_methods to canonical card IDs, dropping stale entries."""
    mapping = _load_valid_method_map()
    if not mapping:
        return phase_methods  # DB not available, pass through
    result: Dict[str, List[str]] = {}
    for phase, methods in phase_methods.items():
        normalized: List[str] = []
        seen: Set[str] = set()
        for m in methods:
            cid = _normalize_method_identifier(m, mapping)
            if cid and cid not in seen:
                normalized.append(cid)
                seen.add(cid)
        result[phase] = normalized
    return result


def _normalize_method_pool(method_pool: List[str]) -> List[str]:
    """Convert all values in method_pool to canonical card IDs, dropping stale entries."""
    mapping = _load_valid_method_map()
    if not mapping:
        return method_pool
    result: List[str] = []
    for m in method_pool:
        cid = _normalize_method_identifier(m, mapping)
        if cid and cid not in result:
            result.append(cid)
    return result


def _load_valid_method_names() -> Set[str]:
    """Build a set of valid method names from the cards DB (respects patched path)."""
    mapping = _load_valid_method_map()
    return set(mapping.keys())


def _filter_stale_methods(phase_methods: Dict[str, List[str]]) -> Dict[str, List[str]]:
    """Remove method names from phase_methods that no longer exist in the cards DB."""
    return _normalize_phase_methods(phase_methods)


def _svg_method_boxes(methods: List[str], col_x: int, col_w: int, start_y: int, box_color: str) -> Tuple[str, int]:
    """Render each method as a vertical-text box, arranged horizontally in rows.

    Returns (svg_string, total_height_including_rows).
    """
    if not methods:
        methods = ["待配置方法"]

    BOX_W = 30
    BOX_GAP = 6
    ROW_GAP = 10
    CHAR_H = 22
    PAD_TOP = 12
    PAD_BOTTOM = 10
    CORNER = 6

    # Max name length determines box height (uniform per phase)
    max_chars = max(len(m) for m in methods)
    BOX_H = max_chars * CHAR_H + PAD_TOP + PAD_BOTTOM

    # How many boxes fit per row
    per_row = max(1, (col_w + BOX_GAP) // (BOX_W + BOX_GAP))

    lines = []
    for idx, method in enumerate(methods):
        row = idx // per_row
        col = idx % per_row
        # Center the row of boxes within the column width
        row_count = min(per_row, len(methods) - row * per_row)
        row_w = row_count * BOX_W + (row_count - 1) * BOX_GAP
        offset_x = (col_w - row_w) / 2
        bx = col_x + offset_x + col * (BOX_W + BOX_GAP)
        by = start_y + row * (BOX_H + ROW_GAP)

        safe = html.escape(method)
        # Box rect
        lines.append(
            f'<rect x="{bx}" y="{by}" width="{BOX_W}" height="{BOX_H}" '
            f'rx="{CORNER}" fill="#ffffff" stroke="{box_color}" stroke-width="1.5"/>'
        )
        # Vertical text: one char per line, centered in box
        chars = list(method)
        text_x = bx + BOX_W // 2
        for ci, ch in enumerate(chars):
            ty = by + PAD_TOP + ci * CHAR_H + CHAR_H // 2 + 4
            lines.append(
                f'<text x="{text_x}" y="{ty}" text-anchor="middle" '
                f'font-size="14" fill="#26384d">{html.escape(ch)}</text>'
            )

    total_rows = (len(methods) + per_row - 1) // per_row
    total_h = total_rows * BOX_H + max(0, total_rows - 1) * ROW_GAP
    return "\n  ".join(lines), total_h


def make_framework_svg(
    topic: str, direction: str, phase_methods: Dict[str, List[str]]
) -> str:
    safe_topic = html.escape(topic or "未命名论文主题")
    safe_direction = html.escape(direction or "工程管理")
    discover = phase_methods.get("discover", []) or ["待配置方法"]
    solve = phase_methods.get("solve", []) or ["待配置方法"]
    validate = phase_methods.get("validate", []) or ["待配置方法"]

    # ── layout constants ──
    COL_X: Dict[str, int] = {"discover": 350, "solve": 740, "validate": 1130}
    COL_COLOR: Dict[str, str] = {"discover": "#6f9ed8", "solve": "#6fb889", "validate": "#9d7ce2"}
    RECT_X: Dict[str, int] = {"discover": 185, "solve": 575, "validate": 965}
    RECT_W = 330

    PHASE_HEADERS_Y = 140
    PHASE_HEADER_H = 78

    CONTENT_Y = 280
    CONTENT_H = 300

    # Method area — vertical boxes, horizontal layout, wrapping rows
    METHOD_COL_LABEL_Y = 636
    METHOD_FIRST_BOX_Y = 662
    BOX_W = 30
    BOX_GAP = 6
    ROW_GAP = 10
    CHAR_H = 22
    PAD_TOP = 12
    PAD_BOTTOM = 10
    PER_ROW = max(1, (RECT_W + BOX_GAP) // (BOX_W + BOX_GAP))  # ≈9 per row

    def _box_h(methods_list):
        max_chars = max((len(m) for m in methods_list), default=4)
        return max_chars * CHAR_H + PAD_TOP + PAD_BOTTOM

    disc_box_h = _box_h(discover)
    solve_box_h = _box_h(solve)
    valid_box_h = _box_h(validate)

    def _rows(methods_list):
        return (len(methods_list) + PER_ROW - 1) // PER_ROW

    disc_rows = _rows(discover)
    solve_rows = _rows(solve)
    valid_rows = _rows(validate)

    disc_methods_h = disc_rows * disc_box_h + max(0, disc_rows - 1) * ROW_GAP
    solve_methods_h = solve_rows * solve_box_h + max(0, solve_rows - 1) * ROW_GAP
    valid_methods_h = valid_rows * valid_box_h + max(0, valid_rows - 1) * ROW_GAP

    max_methods_h = max(disc_methods_h, solve_methods_h, valid_methods_h, 60)

    BIG_BOX_Y = 624
    BIG_BOX_X = 120
    BIG_BOX_W = 1200
    BIG_BOX_PAD = 24
    BIG_BOX_H = max_methods_h + (METHOD_COL_LABEL_Y - BIG_BOX_Y) + BIG_BOX_PAD

    OUTPUTS_Y = BIG_BOX_Y + BIG_BOX_H + 20
    OUTPUT_H = 42
    TOTAL_H = OUTPUTS_Y + OUTPUT_H + 40
    ARROW_Y = 430
    CONTENT_BOTTOM_Y = CONTENT_Y + CONTENT_H

    # ── build method boxes ──
    discover_boxes, _ = _svg_method_boxes(discover, RECT_X["discover"], RECT_W, METHOD_FIRST_BOX_Y, COL_COLOR["discover"])
    solve_boxes, _ = _svg_method_boxes(solve, RECT_X["solve"], RECT_W, METHOD_FIRST_BOX_Y, COL_COLOR["solve"])
    validate_boxes, _ = _svg_method_boxes(validate, RECT_X["validate"], RECT_W, METHOD_FIRST_BOX_Y, COL_COLOR["validate"])

    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="1440" height="{TOTAL_H}" viewBox="0 0 1440 {TOTAL_H}">
  <defs>
    <marker id="arrow" markerWidth="12" markerHeight="10" refX="10" refY="5" orient="auto">
      <path d="M0,0 L12,5 L0,10 z" fill="#617083"/>
    </marker>
    <filter id="shadow" x="-10%" y="-10%" width="120%" height="120%">
      <feDropShadow dx="0" dy="8" stdDeviation="10" flood-color="#c6d7ea" flood-opacity="0.45"/>
    </filter>
  </defs>
  <rect width="1440" height="{TOTAL_H}" fill="#edf7ff"/>
  <rect x="38" y="34" width="1364" height="{TOTAL_H - 68}" rx="24" fill="#fafdff" stroke="#cfe0f2"/>
  <text x="740" y="88" text-anchor="middle" font-family="Arial, PingFang SC, Microsoft YaHei, sans-serif" font-size="30" font-weight="700" fill="#17324d">{safe_topic}研究框架</text>

  <!-- 研究思路 -->
  <text x="92" y="183" text-anchor="middle" font-size="18" font-weight="700" fill="#4a6178">研究思路</text>
  <rect x="{RECT_X["discover"]}" y="{PHASE_HEADERS_Y}" width="{RECT_W}" height="{PHASE_HEADER_H}" rx="14" fill="#e8f2ff" stroke="#6f9ed8"/>
  <rect x="{RECT_X["solve"]}" y="{PHASE_HEADERS_Y}" width="{RECT_W}" height="{PHASE_HEADER_H}" rx="14" fill="#eaf8f0" stroke="#6fb889"/>
  <rect x="{RECT_X["validate"]}" y="{PHASE_HEADERS_Y}" width="{RECT_W}" height="{PHASE_HEADER_H}" rx="14" fill="#f2ecff" stroke="#9d7ce2"/>
  <text x="{COL_X["discover"]}" y="188" text-anchor="middle" font-size="22" font-weight="700" fill="#18324d">发现问题</text>
  <text x="{COL_X["solve"]}" y="188" text-anchor="middle" font-size="22" font-weight="700" fill="#18324d">解决问题</text>
  <text x="{COL_X["validate"]}" y="188" text-anchor="middle" font-size="22" font-weight="700" fill="#18324d">验证问题</text>

  <!-- 研究内容 -->
  <text x="92" y="470" text-anchor="middle" font-size="18" font-weight="700" fill="#4a6178">研究内容</text>
  <g filter="url(#shadow)">
    <rect x="{RECT_X["discover"]}" y="{CONTENT_Y}" width="{RECT_W}" height="{CONTENT_H}" rx="16" fill="#ffffff" stroke="#6f9ed8"/>
    <rect x="{RECT_X["solve"]}" y="{CONTENT_Y}" width="{RECT_W}" height="{CONTENT_H}" rx="16" fill="#ffffff" stroke="#6fb889"/>
    <rect x="{RECT_X["validate"]}" y="{CONTENT_Y}" width="{RECT_W}" height="{CONTENT_H}" rx="16" fill="#ffffff" stroke="#9d7ce2"/>
  </g>
  <text x="{COL_X["discover"]}" y="340" text-anchor="middle" font-size="21" font-weight="700" fill="#1f3650">现状诊断与问题识别</text>
  <text x="{COL_X["discover"]}" y="390" text-anchor="middle" font-size="18" fill="#26384d">研究对象界定</text>
  <text x="{COL_X["discover"]}" y="430" text-anchor="middle" font-size="18" fill="#26384d">流程与质量数据分析</text>
  <text x="{COL_X["discover"]}" y="470" text-anchor="middle" font-size="18" fill="#26384d">核心问题归类与排序</text>
  <text x="{COL_X["solve"]}" y="340" text-anchor="middle" font-size="21" font-weight="700" fill="#1f3650">优化方案构建</text>
  <text x="{COL_X["solve"]}" y="390" text-anchor="middle" font-size="18" fill="#26384d">{safe_direction}理论适配</text>
  <text x="{COL_X["solve"]}" y="430" text-anchor="middle" font-size="18" fill="#26384d">过程改进与闭环机制</text>
  <text x="{COL_X["solve"]}" y="470" text-anchor="middle" font-size="18" fill="#26384d">实施保障与工具支撑</text>
  <text x="{COL_X["validate"]}" y="340" text-anchor="middle" font-size="21" font-weight="700" fill="#1f3650">效果评价与趋势判断</text>
  <text x="{COL_X["validate"]}" y="390" text-anchor="middle" font-size="18" fill="#26384d">实施前后对比</text>
  <text x="{COL_X["validate"]}" y="430" text-anchor="middle" font-size="18" fill="#26384d">关键指标跟踪</text>
  <text x="{COL_X["validate"]}" y="470" text-anchor="middle" font-size="18" fill="#26384d">改进结论输出</text>
  <line x1="515" y1="{ARROW_Y}" x2="575" y2="{ARROW_Y}" stroke="#617083" stroke-width="4" marker-end="url(#arrow)"/>
  <line x1="905" y1="{ARROW_Y}" x2="965" y2="{ARROW_Y}" stroke="#617083" stroke-width="4" marker-end="url(#arrow)"/>

  <!-- 研究方法：大框包含所有方法 -->
  <text x="92" y="{BIG_BOX_Y + 28}" text-anchor="middle" font-size="18" font-weight="700" fill="#4a6178">研究方法</text>
  <rect x="{BIG_BOX_X}" y="{BIG_BOX_Y}" width="{BIG_BOX_W}" height="{BIG_BOX_H}" rx="20" fill="#f8fcff" stroke="#8199b3" stroke-width="2.5"/>
  <!-- 列标题 -->
  <text x="{COL_X["discover"]}" y="{METHOD_COL_LABEL_Y}" text-anchor="middle" font-size="15" font-weight="700" fill="#31526f">发现问题方法</text>
  <text x="{COL_X["solve"]}" y="{METHOD_COL_LABEL_Y}" text-anchor="middle" font-size="15" font-weight="700" fill="#31526f">解决问题方法</text>
  <text x="{COL_X["validate"]}" y="{METHOD_COL_LABEL_Y}" text-anchor="middle" font-size="15" font-weight="700" fill="#31526f">验证问题方法</text>
  {discover_boxes}
  {solve_boxes}
  {validate_boxes}

  <!-- 研究方法 → 研究内容箭头 -->
  <line x1="{COL_X["discover"]}" y1="{BIG_BOX_Y}" x2="{COL_X["discover"]}" y2="{CONTENT_BOTTOM_Y}" stroke="#617083" stroke-width="3" stroke-dasharray="8 6" marker-end="url(#arrow)"/>
  <line x1="{COL_X["solve"]}" y1="{BIG_BOX_Y}" x2="{COL_X["solve"]}" y2="{CONTENT_BOTTOM_Y}" stroke="#617083" stroke-width="3" stroke-dasharray="8 6" marker-end="url(#arrow)"/>
  <line x1="{COL_X["validate"]}" y1="{BIG_BOX_Y}" x2="{COL_X["validate"]}" y2="{CONTENT_BOTTOM_Y}" stroke="#617083" stroke-width="3" stroke-dasharray="8 6" marker-end="url(#arrow)"/>

  <!-- 输出 -->
  <rect x="{RECT_X["discover"]}" y="{OUTPUTS_Y}" width="{RECT_W}" height="{OUTPUT_H}" rx="12" fill="#e8f2ff" stroke="#6f9ed8"/>
  <rect x="{RECT_X["solve"]}" y="{OUTPUTS_Y}" width="{RECT_W}" height="{OUTPUT_H}" rx="12" fill="#eaf8f0" stroke="#6fb889"/>
  <rect x="{RECT_X["validate"]}" y="{OUTPUTS_Y}" width="{RECT_W}" height="{OUTPUT_H}" rx="12" fill="#f2ecff" stroke="#9d7ce2"/>
  <text x="{COL_X["discover"]}" y="{OUTPUTS_Y + 28}" text-anchor="middle" font-size="16" font-weight="700" fill="#24384f">输出：问题清单与诊断依据</text>
  <text x="{COL_X["solve"]}" y="{OUTPUTS_Y + 28}" text-anchor="middle" font-size="16" font-weight="700" fill="#24384f">输出：管理优化方案</text>
  <text x="{COL_X["validate"]}" y="{OUTPUTS_Y + 28}" text-anchor="middle" font-size="16" font-weight="700" fill="#24384f">输出：效果评价与改进结论</text>
</svg>"""


_INDUSTRY_FILTER_KEYWORDS: List[str] = [
    "新能源", "新能源汽车", "光伏", "风电", "储能",
    "半导体", "芯片", "集成电路",
    "电池", "锂电池", "动力电池",
    "汽车", "车载", "自动驾驶", "智能驾驶",
    "建筑", "房地产", "桥梁", "隧道",
    "钢铁", "冶金", "化工",
    "医药", "制药", "医疗器械", "医疗设备",
    "通信", "5G", "互联网", "软件开发",
    "金融", "保险", "银行",
    "物流", "供应链",
    "食品", "纺织", "石油",
    "人工智能", "大数据", "物联网",
]


def _filter_references_by_topic(
    topic: str, references: List[Dict[str, Any]], limit: int = 3
) -> List[Dict[str, Any]]:
    """Filter out knowledge base references whose industry keywords don't match the user's topic."""
    filtered = []
    for ref in references:
        title = str(ref.get("title", ""))
        content = str(ref.get("content", ""))
        combined = title + content
        # Filter out references that mention industries NOT in the user's topic
        mismatched = [
            kw for kw in _INDUSTRY_FILTER_KEYWORDS
            if kw in combined and kw not in topic
        ]
        if mismatched:
            continue
        filtered.append(ref)
        if len(filtered) >= limit:
            break

    # 如果过滤后太少，回退到原始结果
    if len(filtered) < 2:
        return references[:limit]
    return filtered


def allocate_words(outline: Dict[str, Any], total_words: int) -> Dict[str, Any]:
    def title_weight(
        title: str, weights: Dict[str, float], default: float = 1.0
    ) -> float:
        value = default
        for keyword, weight in weights.items():
            if keyword in str(title):
                value = max(value, weight)
        return value

    def distribute(
        total: int, items: List[Dict[str, Any]], weights: List[float]
    ) -> List[int]:
        if not items:
            return []
        weight_sum = sum(weights) or len(items)
        raw = [max(1, int(total * weight / weight_sum)) for weight in weights]
        diff = total - sum(raw)
        order = sorted(range(len(items)), key=lambda idx: weights[idx], reverse=True)
        cursor = 0
        while diff != 0 and order:
            idx = order[cursor % len(order)]
            if diff > 0:
                raw[idx] += 1
                diff -= 1
            elif raw[idx] > 1:
                raw[idx] -= 1
                diff += 1
            cursor += 1
            if cursor > len(order) * max(abs(diff), 1) + 20:
                break
        return raw

    chapters = outline.get("chapters", [])
    raw_weights = [
        title_weight(chapter.get("title", ""), WORD_WEIGHTS, 0.15)
        for chapter in chapters
    ]
    chapter_allocations = distribute(total_words, chapters, raw_weights)

    for chapter, chapter_words in zip(chapters, chapter_allocations):
        chapter["estimated_words"] = chapter_words
        sections = chapter.get("sections", [])
        section_weights = [
            title_weight(section.get("title", ""), SECTION_WORD_WEIGHTS, 1.0)
            for section in sections
        ]
        for section, section_words in zip(
            sections, distribute(chapter_words, sections, section_weights)
        ):
            section["estimated_words"] = section_words
            subsections = section.get("subsections", [])
            if subsections:
                subsection_weights = [
                    title_weight(subsection.get("title", ""), SECTION_WORD_WEIGHTS, 1.0)
                    for subsection in subsections
                ]
                for subsection, subsection_words in zip(
                    subsections,
                    distribute(section_words, subsections, subsection_weights),
                ):
                    subsection["estimated_words"] = subsection_words
    outline["metadata"]["estimated_words"] = total_words
    return outline


def outline_to_markdown(outline: Dict[str, Any]) -> str:
    def clean(title: str, number: str = "") -> str:
        value = re.sub(r"^#{1,6}\s*", "", str(title or "")).strip()
        if number:
            value = re.sub(rf"^{re.escape(number)}\s+", "", value)
        value = re.sub(r"^\d+(?:\.\d+){{0,3}}\s+", "", value)
        return value.strip()

    lines = [f"# {outline.get('title', '论文大纲')}", ""]
    for chapter in outline.get("chapters", []):
        words = chapter.get("estimated_words", 0)
        ch_num = chapter.get("number", "")
        lines.append(f"## 第{ch_num}章 {clean(chapter['title'])}（约{words}字）")
        for section in chapter.get("sections", []):
            lines.append(
                f"### {section['number']} {clean(section['title'], section['number'])}（约{section.get('estimated_words', 0)}字）"
            )
            for subsection in section.get("subsections", []):
                lines.append(
                    f"#### {subsection['number']} {clean(subsection['title'], subsection['number'])}（约{subsection.get('estimated_words', 0)}字）"
                )
        lines.append("")
    return "\n".join(lines)


def _extract_json_object(text: str) -> Dict[str, Any]:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("model did not return a JSON object")
    return json.loads(text[start : end + 1])


def _normalize_llm_outline(raw: Dict[str, Any], topic: str) -> Dict[str, Any]:
    chapters = []
    for chapter_index, chapter in enumerate(raw.get("chapters", []), 1):
        sections = []
        for section_index, section in enumerate(chapter.get("sections", []), 1):
            subsections = []
            for subsection_index, subsection in enumerate(
                section.get("subsections", []), 1
            ):
                title = (
                    subsection.get("title")
                    if isinstance(subsection, dict)
                    else str(subsection)
                )
                subsections.append(
                    {
                        "level": 3,
                        "number": f"{chapter_index}.{section_index}.{subsection_index}",
                        "title": strip_markdown(title),
                    }
                )
            sections.append(
                {
                    "level": 2,
                    "number": f"{chapter_index}.{section_index}",
                    "title": strip_markdown(
                        section.get("title", f"小节{section_index}")
                    ),
                    "subsections": subsections,
                }
            )
        chapters.append(
            {
                "level": 1,
                "number": chapter_index,
                "title": f"第{chapter_index}章 {strip_markdown(chapter.get('title', f'章节{chapter_index}').replace(f'第{chapter_index}章', '').strip())}",
                "sections": sections,
            }
        )
    if not chapters:
        raise ValueError("outline has no chapters")
    return {
        "title": raw.get("title") or f"论文: {topic}",
        "depth": 3,
        "chapters": chapters,
        "chapter_count": len(chapters),
        "metadata": {
            "total_chapters": len(chapters),
            "generated_by": "llm_rag",
            "created_at": raw.get("created_at", ""),
        },
    }


def generate_llm_outline(
    payload: Dict[str, Any], total_words: int, task_id: str = ""
) -> Dict[str, Any]:
    started_at = time.monotonic()
    topic = payload.get("topic", "未命名")
    project_context = persist_project_context(payload)
    direction = payload.get("direction_name") or payload.get("direction", "工程管理")
    domain_template = load_domain_template(payload.get("direction", ""))
    methods = payload.get("methods", [])
    phase_methods = payload.get("phase_methods", {})
    if task_id:
        task_log(task_id, "正在检索研究方向相关论文...")
    # Step 1: 研究方向优先 — 检索与研究方向相关的论文
    dir_query = f"{direction} 工程管理 硕士论文 研究"
    dir_refs = search_knowledge_base(dir_query, limit=4).get("results", [])
    dir_refs = _filter_references_by_topic(topic, dir_refs, limit=3)

    # Step 2: 方法相关 — 为每个已选方法检索相关论文
    method_refs: List[Dict[str, Any]] = []
    seen_paths = {item.get("path", "") for item in dir_refs}
    if task_id:
        task_log(task_id, f"研究方向检索命中 {len(dir_refs)} 篇，继续检索方法相关论文...")
    for method in methods[:5]:
        m_query = f"{method} 工程管理 论文 研究"
        m_results = search_knowledge_base(m_query, limit=2).get("results", [])
        for item in m_results:
            if item.get("path", "") not in seen_paths:
                method_refs.append(item)
                seen_paths.add(item.get("path", ""))
        if len(method_refs) >= 6:
            break

    # 合并：方向论文在前，方法论文在后
    references = dir_refs + method_refs[:6]
    if task_id:
        titles = "、".join(item.get("title", "未命名资料") for item in references[:4])
        task_log(
            task_id,
            f"知识库检索完成，用时 {time.monotonic() - started_at:.1f}s；方向论文 {len(dir_refs)} 篇 + 方法论文 {len(method_refs[:6])} 篇：{titles or '无高相关资料'}",
        )
    reference_text = "\n\n".join(
        f"[资料{idx}] {item.get('title')} / {item.get('path')}\n{item.get('content', '')[:260]}"
        for idx, item in enumerate(references, 1)
    )
    template_text = json.dumps(
        {
            "方向": domain_template.get("name", direction),
            "样本数量": domain_template.get("source_count", 0),
            "常见章节命名": domain_template.get("chapter_naming_patterns", [])[:12],
            "常见小节命名": domain_template.get("section_naming_patterns", [])[:18],
            "常用方法组合": domain_template.get("method_patterns", [])[:12],
            "典型来源": domain_template.get("sources", [])[:6],
        },
        ensure_ascii=False,
    )

    # 大纲索引：同方向论文的实际章节结构
    outline_index_text = ""
    try:
        direction_outlines = query_outlines_by_direction(direction)
        if direction_outlines:
            samples = direction_outlines[:5]
            outline_index_text = "同方向论文大纲参考（章节结构样本）：\n" + "\n".join(
                f"- {o.get('title', '?')[:60]}\n  {json.dumps(o.get('outline', []), ensure_ascii=False)[:300]}"
                for o in samples
            )
    except Exception:
        pass

    config = load_llm_config()
    if not config.api_key:
        raise RuntimeError("未配置 API Key")
    if task_id:
        task_log(
            task_id,
            f"正在调用大模型：{config.provider}/{config.model}；已压缩知识库上下文",
        )

    client, provider = _build_llm_client_and_provider(config)

    prompt = f"""请基于选题、研究方向、阶段方法和本地知识库资料，生成工程管理硕士论文大纲骨架。

论文主题：{topic}
研究方向：{direction}
已选方法：{json.dumps(methods, ensure_ascii=False)}
阶段方法配置：{json.dumps(phase_methods, ensure_ascii=False)}
目标总字数：{total_words}

方向模板（仅供参考章节命名习惯，不要照搬案例公司的行业）：
{template_text}

{outline_index_text}

要求：
1. 论文主题必须严格围绕「{topic}」展开，章节标题和内容必须与这个具体题目匹配。
2. 严禁引入用户未提及的行业、公司类型或技术领域（如新能源汽车、半导体、电池等），不要根据知识库资料中的案例行业来替换用户的选题。
3. 必须优先吸收下方「项目背景与论文思路」中的章节结构、核心章节安排、研究方法和写作要求。
4. 优先参考方向模板中的章节命名规则和方法组合，再结合用户选定方法微调。
5. 章节逻辑体现"发现问题-解决问题-验证问题"。
6. 生成 6 章；每章只给 3-4 个二级小节。
7. 不要生成三级小节，subsections 固定为空数组。
8. 只输出 JSON，不要 Markdown，不要解释。

项目背景与论文思路：
{project_context[:3200] or "用户未填写。"}

JSON 结构：
{{
  "title": "论文: ...",
  "chapters": [
    {{
      "title": "绪论",
      "sections": [
        {{"title": "研究背景", "subsections": []}}
      ]
    }}
  ]
}}

本地知识库参考（仅供研究方法论和论文结构参考，不要照搬其行业/公司）：
{reference_text or "未检索到高相关资料。"}
"""
    response = _llm_create_message(
        client, provider, config,
        system="你是工程管理硕士论文大纲设计专家，擅长结合本地案例论文和研究方法生成可落地的章节结构。",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=_token_budget(config, "small"),
    )
    if task_id:
        task_log(
            task_id,
            f"大模型已返回，总用时 {time.monotonic() - started_at:.1f}s；正在解析 JSON 大纲...",
        )
    raw = _extract_json_object(_llm_response_text(response, provider))
    outline = _normalize_llm_outline(raw, topic)
    if task_id:
        task_log(
            task_id,
            f"已解析 {len(outline.get('chapters', []))} 个章节，正在分配字数...",
        )
    outline["metadata"]["references"] = [
        {
            "title": item.get("title"),
            "path": item.get("path"),
            "score": item.get("score"),
        }
        for item in references
    ]
    return allocate_words(outline, total_words)


def generate_outline_result(
    payload: Dict[str, Any], task_id: str = ""
) -> Dict[str, Any]:
    total_words = int(payload.get("total_words", 50000))
    generation_mode = "llm_rag"
    try:
        outline = generate_llm_outline(payload, total_words, task_id=task_id)
    except Exception as exc:
        generation_mode = f"fallback_rule: {exc}"
        if task_id:
            task_log(task_id, f"大模型生成未完成，启用本地规则兜底：{exc}")
        framework = {
            "topic": payload.get("topic", "未命名"),
            "discipline": payload.get("direction", "general"),
            "research_phases": ["发现问题", "解决问题", "验证问题"],
            "key_components": payload.get("methods", []),
        }
        outline = allocate_words(generate_outline(framework, depth=3), total_words)
        outline.setdefault("metadata", {})["generated_by"] = generation_mode

    markdown = outline_to_markdown(outline)
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    (OUTPUT_ROOT / "web_thesis_outline.md").write_text(markdown, encoding="utf-8")
    save_workspace_value("outline", outline)
    save_workspace_value("markdown", markdown)
    build_thesis_memory(payload, outline)
    if task_id:
        task_log(task_id, "已生成 Markdown，并写入 output/web_thesis_outline.md")
    return {
        "outline": outline,
        "markdown": markdown,
        "generation_mode": outline.get("metadata", {}).get(
            "generated_by", generation_mode
        ),
    }


def _run_outline_task(task_id: str, payload: Dict[str, Any]) -> None:
    with TASK_LOCK:
        TASKS[task_id].update(
            {"status": "running", "message": "正在检索知识库并调用大模型..."}
        )
    try:
        task_log(task_id, "任务开始：准备生成章节大纲")
        result = generate_outline_result(payload, task_id=task_id)
        with TASK_LOCK:
            TASKS[task_id].update(
                {
                    "status": "done",
                    "message": "生成完成",
                    "result": result,
                    "finished_at": time.time(),
                }
            )
    except Exception as exc:
        with TASK_LOCK:
            TASKS[task_id].update(
                {
                    "status": "error",
                    "message": str(exc),
                    "finished_at": time.time(),
                }
            )


def start_outline_task(payload: Dict[str, Any]) -> str:
    task_id = uuid.uuid4().hex
    with TASK_LOCK:
        TASKS[task_id] = {
            "kind": OUTLINE_TASK_KIND,
            "permission_menu": TASK_PERMISSION_BY_KIND[OUTLINE_TASK_KIND],
            "status": "queued",
            "message": "任务已创建",
            "logs": [
                {
                    "time": time.strftime("%H:%M:%S"),
                    "message": "任务已创建，等待后台执行",
                }
            ],
            "created_at": time.time(),
        }
    thread = threading.Thread(
        target=_run_outline_task, args=(task_id, payload), daemon=True
    )
    thread.start()
    return task_id


_DSML_RE = re.compile(r'</?[|\s]*DSML[^>]*>', re.IGNORECASE)
_DSML_FRAG_RE = re.compile(r'<[|]{1,2}\s*>|</[|]{1,2}\s*>', re.IGNORECASE)
_DSML_INVOKE_RE = re.compile(
    r'<[|\s]*DSML[|\s]*invokename="([^"]+)"[^>]*>', re.IGNORECASE
)
_DSML_PARAM_RE = re.compile(
    r'<[|\s]*DSML[|\s]*parameter\s+name="([^"]+)"\s*string="(true|false)"[^>]*>(.*?)(?=</?[|\s]*DSML|$)',
    re.DOTALL | re.IGNORECASE,
)


def _strip_dsml(text: str) -> str:
    """Remove DSML tool-call tokens that DeepSeek may emit as raw text."""
    if not text or "DSML" not in text:
        return text
    text = _DSML_RE.sub("", text)
    text = _DSML_FRAG_RE.sub("", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_dsml_tools(text: str) -> list:
    """Parse DSML tool calls from LLM text into a list of (name, params_dict) tuples.
    Returns empty list if no DSML tool calls found."""
    if not text or "DSML" not in text or "invokename" not in text:
        return []
    tools = []
    # Find all tool invocations and parameter blocks independently
    invocations = list(_DSML_INVOKE_RE.finditer(text))
    param_matches = list(_DSML_PARAM_RE.finditer(text))
    for i, inv in enumerate(invocations):
        tool_name = inv.group(1)
        # Determine the span of this tool's content:
        # from this invocation to the next invocation (or end of text)
        start = inv.end()
        end = invocations[i + 1].start() if i + 1 < len(invocations) else len(text)
        params = {}
        for pm in param_matches:
            if start <= pm.start() < end:
                pname = pm.group(1)
                is_string = pm.group(2) == "true"
                pvalue = pm.group(3).strip()
                if not is_string:
                    try:
                        pvalue = int(pvalue)
                    except (ValueError, TypeError):
                        pass
                params[pname] = pvalue
        tools.append((tool_name, params))
    return tools


def _response_text(response: Any) -> str:
    parts = []
    for block in response.content:
        if hasattr(block, "text"):
            parts.append(_strip_dsml(block.text))
    return "\n".join(parts).strip()


def _llm_response_text(response: Any, provider: str) -> str:
    if provider == "openai":
        return response.choices[0].message.content or ""
    return _response_text(response)


def _token_budget(config: Any = None, preset: str = "default") -> int:
    """Return per-call token budget derived from the global config max_tokens."""
    if config is None:
        config = load_llm_config()
    base = int(getattr(config, "max_tokens", 4000) or 4000)
    ratios = {
        "test": 0.025,       # ~100 tokens (just "连接成功")
        "tiny": 0.15,        # ~600 tokens (classification labels)
        "small": 0.35,       # ~1400 tokens (subsections, outlines)
        "default": 0.5,      # ~2000 tokens (expand, rewrite, scan)
        "large": 0.8,        # ~3200 tokens (card content, full chapters)
        "max": 1.0,          # full config value
    }
    ratio = ratios.get(preset, 0.5)
    return max(50, int(base * ratio))


def _build_llm_client_and_provider(config: Any = None) -> tuple:
    if config is None:
        config = load_llm_config()
    if config.provider == "openai":
        import openai
        kwargs: Dict[str, Any] = {"api_key": config.api_key, "timeout": 60}
        if config.base_url:
            kwargs["base_url"] = config.base_url
        return openai.OpenAI(**kwargs), "openai"
    client_kwargs: Dict[str, Any] = {"api_key": config.api_key, "timeout": 60}
    if config.auth_mode == "auth_token" and config.auth_token:
        client_kwargs["auth_token"] = config.auth_token
    else:
        os.environ.pop("ANTHROPIC_AUTH_TOKEN", None)
    if config.base_url:
        client_kwargs["base_url"] = config.base_url
    return Anthropic(**client_kwargs), config.provider


def _llm_create_message(
    client: Any,
    provider: str,
    config: Any,
    system: str,
    messages: List[Dict[str, str]],
    max_tokens: int,
    tools: Any = None,
    disable_thinking: bool = False,
) -> Any:
    # 记录完整提示词到日志
    logger.info(
        "LLM call → %s/%s | max_tokens=%s | system=%.120s",
        provider, config.model, max_tokens, (system or "").replace("\n", " ")
    )
    for i, msg in enumerate(messages):
        content = msg.get("content", "")
        logger.info("LLM prompt [msg %s/%s]:\n%s", i + 1, len(messages), content)

    if provider == "openai":
        api_messages: List[Dict[str, str]] = []
        if system:
            api_messages.append({"role": "system", "content": system})
        api_messages.extend(messages)
        return client.chat.completions.create(
            model=config.model,
            max_tokens=max_tokens,
            messages=api_messages,
        )
    kwargs: Dict[str, Any] = dict(
        model=config.model,
        max_tokens=max_tokens,
        system=system,
        messages=messages,
    )
    if disable_thinking:
        kwargs["thinking"] = {"type": "disabled"}
    if tools:
        kwargs["tools"] = tools
    return client.messages.create(**kwargs)


def strip_markdown(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"```[\s\S]*?```", lambda m: m.group(0).strip("`"), value)
    value = re.sub(r"^#{1,6}\s*", "", value, flags=re.MULTILINE)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"__([^_]+)__", r"\1", value)
    value = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", value)
    value = re.sub(r"(?<!_)_([^_\n]+)_(?!_)", r"\1", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"^\s*[-*+]\s+", "", value, flags=re.MULTILINE)
    return value.strip()


CHAT_SYSTEM_PROMPT = """你是 ThesisMind 论文辅助工作台的智能引导助手，专门服务工程管理硕士论文写作。

## 工作台流程（7 步）
01 基本配置 — 配置大模型连接（provider / model / base_url / api_key）
02 论文信息 — 设置论文主题、研究方向、项目背景，初始化知识库
03 方法论选择 — 从知识库扫描可用研究方法，分配到"发现问题/解决问题/验证问题"三个阶段
04 研究框架 — 生成研究框架 SVG 图（方法论与论文主题的映射关系）
05 章节大纲 — 生成三级目录结构，分配字数，支持后续手动编辑
06 引用生成 — 从本地知识库匹配真实文献，生成 GB/T 7714 格式引用
07 章节写作 — 按小节扩写/重写，支持一键串行完成所有小节

## 你的职责
1. 根据用户当前所处步骤，给出针对性的引导和建议。
2. 帮助用户理解每一步的目的和操作方法。
3. 当用户不确定选什么方法论、怎么设计框架、怎么分配字数时，给出具体可行的建议。
4. 始终围绕工程管理论文写作（质量管理、风险管理、进度管理、流程优化等方向）提供专业建议。
5. 回答简洁实用，避免空泛的学术说教。

## 当前状态
用户可以通过左侧边栏的 7 个步骤导航切换页面。页面上的表单和按钮可以完成每一步的操作。
你的任务是用自然对话的方式解释每一步该做什么、为什么这么做、怎么做更好。
当用户询问具体论文问题时，结合工程管理学科特点给出建议。"""




def _test_llm_connection() -> Dict[str, Any]:
    config = load_llm_config()
    if not config.api_key:
        return {"status": "error", "message": "未配置 API Key，请先在基本配置页面填写 API Key 并保存。"}
    try:
        client, provider = _build_llm_client_and_provider(config)
        response = _llm_create_message(
            client, provider, config,
            system="你是一个测试助手。",
            messages=[{"role": "user", "content": "请回复'连接成功'。"}],
            max_tokens=_token_budget(config, "test"),
        )
        text = _llm_response_text(response, provider)
        return {
            "status": "ok",
            "message": f"连接成功！模型 {config.model} 已就绪。",
            "model": config.model,
            "provider": config.provider,
            "reply": text[:100],
        }
    except Exception as exc:
        return {"status": "error", "message": f"连接失败：{exc}"}


CHAT_TOOLS_SCHEMA = [
    {
        "name": "search_knowledge_base",
        "description": "搜索本地知识库，查找与关键词相关的论文、方法论、模板等资料。适用于用户询问示例、参考、写作方法等场景。",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "搜索关键词，如'项目背景示例''质量管理论文思路''PDCA论文结构'"},
                "limit": {"type": "integer", "description": "返回结果数，默认5，最多10"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "query_outlines",
        "description": "按研究方向查询已建索引的论文大纲目录，了解类似论文的章节结构设计。",
        "input_schema": {
            "type": "object",
            "properties": {
                "direction": {"type": "string", "description": "研究方向，如'质量管理''风险管理''进度管理'"},
            },
            "required": ["direction"],
        },
    },
]

PAGE_SKILL_MAP = {
    "setup": "PAGE_SETUP",
    "paper_info": "PAGE_PAPER_INFO",
    "methods": "PAGE_METHODOLOGY",
    "framework": "PAGE_FRAMEWORK",
    "outline": "PAGE_OUTLINE",
    "citations": "PAGE_CITATIONS",
    "writing": "PAGE_WRITING",
}

SKILLS_DIR = PROJECT_ROOT / "skills"


def _load_page_skill(step: str) -> str:
    skill_name = PAGE_SKILL_MAP.get(step)
    if not skill_name:
        return ""
    skill_path = SKILLS_DIR / f"{skill_name}.md"
    if not skill_path.exists():
        return ""
    try:
        raw = skill_path.read_text(encoding="utf-8")
        # Strip YAML frontmatter
        content = re.sub(r"^\s*---\s*\n.*?\n\s*---\s*\n", "", raw, flags=re.DOTALL)
        return f"\n\n# 当前页面专属引导（{step}）\n{content.strip()}"
    except Exception:
        return ""


def _chat_tools_schema() -> List[Dict[str, Any]]:
    return CHAT_TOOLS_SCHEMA


def _execute_chat_tool(name: str, tool_input: Dict[str, Any]) -> str:
    if name == "search_knowledge_base":
        query = str(tool_input.get("query", ""))
        limit = min(int(tool_input.get("limit", 5)), 10)
        result = search_knowledge_base(query=query, limit=limit)
        results = result.get("results", [])
        if not results:
            return "未找到匹配的资料。"
        lines = []
        for item in results[:limit]:
            path = item.get("path", item.get("file", ""))
            title = item.get("title", "")
            content = str(item.get("content", item.get("text", "")))[:400]
            score = item.get("score", item.get("similarity", 0))
            lines.append(f"### {title or path}\n路径: {path}\n相关度: {score:.2f}\n内容片段: {content}")
        return "\n\n---\n\n".join(lines)
    elif name == "query_outlines":
        direction = str(tool_input.get("direction", ""))
        outlines = query_outlines_by_direction(direction)
        if not outlines:
            return f"未找到「{direction}」方向的大纲索引。可用方向：质量管理、风险管理、进度管理、成本管理、流程优化、绩效评价、需求管理、供应链物流、通用项目管理。"
        lines = []
        for entry in outlines[:5]:
            title = entry.get("title", "")
            methods = entry.get("methodologies", [])
            outline = entry.get("outline", {})
            chapters = outline.get("chapters", [])
            chapter_titles = [c.get("title", "") for c in chapters[:8]]
            lines.append(f"- {title}\n  方法论: {', '.join(methods[:5]) or '未标注'}\n  章: {' → '.join(chapter_titles)}")
        return "\n".join(lines)
    return f"未知工具: {name}"


def _chat_with_llm(payload: Dict[str, Any]) -> Dict[str, Any]:
    config = load_llm_config()
    if not config.api_key:
        return {"status": "error", "message": "未配置 API Key，无法使用聊天功能。"}

    user_messages = payload.get("messages", [])
    if not user_messages:
        return {"status": "error", "message": "消息列表为空。"}

    context = payload.get("context", {})
    current_step = context.get("current_step", "setup")
    topic = context.get("topic", "")
    direction = context.get("direction", "")
    methods = context.get("methods", [])
    project_bg = context.get("project_bg", "")
    project_approach = context.get("project_approach", "")

    context_note = f"用户当前在「{current_step}」步骤。"
    if topic:
        context_note += f" 论文主题：「{topic}」。"
    if direction:
        context_note += f" 研究方向：「{direction}」。"
    if methods:
        context_note += f" 已选方法论：{'、'.join(methods[:8])}。"
    if project_bg:
        context_note += f"\n用户已填写的项目背景：{project_bg[:600]}"
    if project_approach:
        context_note += f"\n用户已填写的论文思路：{project_approach[:600]}"

    system = CHAT_SYSTEM_PROMPT + f"\n\n## 用户当前状态\n{context_note}"
    system += _best_practices_summary()
    system += _load_page_skill(current_step)

    api_messages: List[Dict[str, Any]] = []
    for msg in user_messages[-30:]:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        if role in ("user", "assistant"):
            api_messages.append({"role": role, "content": content})

    try:
        client, provider = _build_llm_client_and_provider(config)
        # OpenAI 路径暂不支持工具调用，降级为纯文本对话
        use_tools = provider != "openai"
        response = _llm_create_message(
            client, provider, config,
            system=system,
            messages=api_messages,
            max_tokens=_token_budget(config, "default"),
            tools=_chat_tools_schema() if use_tools else None,
        )

        if not use_tools:
            text = _llm_response_text(response, provider)
            return {"status": "ok", "message": {"role": "assistant", "content": text}}

        tool_blocks = []
        text_blocks = []
        for block in response.content:
            if hasattr(block, "type") and block.type == "tool_use":
                tool_blocks.append(block)
            elif hasattr(block, "text"):
                text_blocks.append(_strip_dsml(block.text))

        if tool_blocks:
            preserved = []
            for block in response.content:
                if hasattr(block, "type") and block.type == "thinking":
                    preserved.append({"type": "thinking", "thinking": getattr(block, "thinking", "")})
                elif hasattr(block, "type") and block.type == "tool_use":
                    preserved.append({"type": "tool_use", "id": block.id, "name": block.name, "input": block.input})
                elif hasattr(block, "text") and not hasattr(block, "type"):
                    preserved.append({"type": "text", "text": _strip_dsml(block.text)})
            api_messages.append({"role": "assistant", "content": preserved})
            tool_results = []
            for tb in tool_blocks:
                result_text = _execute_chat_tool(tb.name, dict(tb.input) if hasattr(tb.input, "items") else tb.input)
                tool_results.append({"type": "tool_result", "tool_use_id": tb.id, "content": result_text})
            api_messages.append({"role": "user", "content": tool_results})

            response2 = _llm_create_message(
                client, provider, config,
                system=system,
                messages=api_messages,
                max_tokens=_token_budget(config, "default"),
            )
            text = _llm_response_text(response2, provider)
        else:
            raw_text = "\n".join(text_blocks).strip()
            # DeepSeek may emit tool calls as DSML text instead of tool_use blocks
            dsml_tools = _parse_dsml_tools(raw_text)
            if dsml_tools:
                preserved_tool = []
                for tname, tparams in dsml_tools:
                    preserved_tool.append({"type": "tool_use", "name": tname, "input": tparams})
                api_messages.append({"role": "assistant", "content": preserved_tool})
                tool_results = []
                for tname, tparams in dsml_tools:
                    result_text = _execute_chat_tool(tname, tparams)
                    tool_results.append({"type": "tool_result", "tool_use_id": tname, "content": result_text})
                api_messages.append({"role": "user", "content": tool_results})
                response2 = _llm_create_message(
                    client, provider, config,
                    system=system,
                    messages=api_messages,
                    max_tokens=_token_budget(config, "default"),
                )
                text = _llm_response_text(response2, provider)
            else:
                text = _strip_dsml(raw_text)

        return {"status": "ok", "message": {"role": "assistant", "content": text}}
    except Exception as exc:
        return {"status": "error", "message": f"聊天请求失败：{exc}"}


def _generate_proposal_section(
    client: Any, provider: str, config: Any,
    section_num: int, section_title: str, section_guide: str,
    context: Dict[str, Any], prev_text: str = "",
) -> str:
    """生成开题报告单个部分。"""
    topic = context["topic"]
    methods_str = context["methods_str"]
    project_context = context["project_context"]
    project_approach = context["project_approach"]
    ch1_content = context["ch1_content"]
    ch2_content = context["ch2_content"]
    lit_text = context["lit_text"]
    citations = context["citations"]
    citation_count = len(citations)

    citation_text = "\n".join(
        f"[{i+1}] {c.get('formatted','')}" for i, c in enumerate(citations)
    )

    if citations:
        if section_num == 3:
            citation_rule = f"引用下方引用库中的全部 {citation_count} 篇文献，每篇至少出现一次。使用右上角标 [N] 格式标注，序号对应引用库编号。最后放置完整参考文献列表。"
            citation_section = f"\n引用库：\n{citation_text[:6000]}\n"
        else:
            citation_rule = "使用右上角标 [N] 格式标注引用（如 [1][2]），序号对应引用库编号。"
            citation_section = ""
    else:
        citation_rule = ""
        citation_section = ""

    prev_block = f"\n前一部分已生成内容（承接上文）：\n{prev_text[-500:]}" if prev_text else ""

    prompt = f"""为工程管理硕士（MEM）论文开题报告生成以下部分。

论文题目：{topic}
研究方向：{context['direction']}
选用方法论：{methods_str}

项目背景：{project_context[:1500] or "未填写。"}
论文思路：{project_approach[:1500] or "未填写。"}
第一章内容摘要：{ch1_content[:1500] or "尚未撰写。"}
第二章内容摘要：{ch2_content[:1500] or "尚未撰写。"}
知识库参考：{lit_text[:1000] or "无。"}

---
当前部分：{section_title}
内容要求：
{section_guide}

写作要求：
1. 正式学术中文，避免口语化
2. 基于项目背景和论文思路，不编造
3. {citation_rule}
4. 严禁 Markdown 语法
5. 直奔主题，不要寒暄
6. 严禁输出标题行（如【一、...】），标题由系统添加{citation_section}{prev_block}"""

    system = "你是工程管理硕士（MEM）学位论文指导教师，擅长撰写开题报告。"

    try:
        response = _llm_create_message(
            client, provider, config,
            system=system,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=config.max_tokens,
        )
        text = _llm_response_text(response, provider)
        # Strip section title if LLM still outputs it
        title_prefixes = [section_title, section_title.lstrip("【").rstrip("】"), section_title.replace("【", "").replace("】", "")]
        for prefix in title_prefixes:
            if prefix and text.startswith(prefix):
                text = text[len(prefix):].lstrip("\n")
                break
        return text
    except Exception as exc:
        raise RuntimeError(f"生成{section_title}失败：{exc}")


def _run_proposal_task(task_id: str, payload: Dict[str, Any]) -> None:
    """后台任务：分三部分生成开题报告。"""
    with TASK_LOCK:
        TASKS[task_id].update({"status": "running", "message": "正在准备..."})

    try:
        config = load_llm_config()
        if not config.api_key:
            raise RuntimeError("未配置 API Key")

        topic = str(payload.get("topic", "") or "工程管理硕士论文").strip()
        direction = str(payload.get("direction", "") or "质量管理").strip()
        methods = payload.get("methods", [])
        methods_str = "、".join(methods[:10]) if methods else "未指定"
        project_context = str(payload.get("project_context", "") or "").strip()
        project_approach = str(payload.get("project_approach", "") or "").strip()
        chapters = payload.get("chapters", [])

        ch1_content = ""
        ch2_content = ""
        for ch in chapters[:2]:
            num = ch.get("number", "")
            content = str(ch.get("content", "")).strip()
            if str(num) == "1" or "一" in str(num):
                ch1_content = content
            else:
                ch2_content = content

        lit_results = search_knowledge_base(f"{topic} {direction} 研究综述 文献", limit=6).get("results", [])
        lit_text = "\n".join(
            f"- {item.get('title','')}: {str(item.get('content',''))[:200]}"
            for item in lit_results
        )

        citations = load_workspace_value("citations", []) or []

        client, provider = _build_llm_client_and_provider(config)

        ctx = {
            "topic": topic,
            "direction": direction,
            "methods_str": methods_str,
            "project_context": project_context,
            "project_approach": project_approach,
            "ch1_content": ch1_content,
            "ch2_content": ch2_content,
            "lit_text": lit_text,
            "citations": citations,
        }

        sections = [
            (
                1,
                "【一、选题有关的国内外研究综述】",
                "- 围绕论文题目梳理国内外研究现状，区分国内和国外研究脉络\n- 指出现有研究的不足或空白\n- 明确本研究的切入点和学术定位\n- 仅综述，不列参考文献列表",
            ),
            (
                2,
                "【二、选题的理论意义、实际意义、创新点】",
                "- 理论意义：本研究对学科理论体系的贡献\n- 实际意义：对工程管理实践、行业或企业的应用价值\n- 创新点：方法创新、视角创新或应用场景创新（2-3点）",
            ),
            (
                3,
                "【三、所要解决的主要问题及研究途径与方法】",
                "- 明确列出要解决的核心问题（3-5个）\n- 阐述研究途径和主要方法，说明方法选择理由\n- 给出全文六章的章节安排（每章1-2句话概括，从第一章到第六章全部描述，不得只写其中几章）\n- 最后放置完整参考文献列表（所有文献按编号排列）",
            ),
        ]

        parts = []
        total_sections = len(sections)
        for i, (num, title, guide) in enumerate(sections):
            with TASK_LOCK:
                TASKS[task_id].update({
                    "status": "running",
                    "message": f"正在生成第{num}部分：{title}...",
                    "progress": int((i / total_sections) * 90),
                })

            text = _generate_proposal_section(
                client, provider, config, num, title, guide, ctx,
                prev_text=parts[-1] if parts else "",
            )
            parts.append(f"{title}\n{text}")

        full_content = "\n\n".join(parts)

        save_workspace_value("proposal_content", full_content)

        with TASK_LOCK:
            TASKS[task_id].update({
                "status": "done",
                "message": "生成完成",
                "progress": 100,
                "result": {"content": full_content},
                "finished_at": time.time(),
            })

    except Exception as exc:
        with TASK_LOCK:
            TASKS[task_id].update({
                "status": "error",
                "message": str(exc),
                "finished_at": time.time(),
            })


def start_proposal_task(payload: Dict[str, Any]) -> str:
    task_id = uuid.uuid4().hex
    with TASK_LOCK:
        TASKS[task_id] = {
            "kind": PROPOSAL_TASK_KIND,
            "permission_menu": TASK_PERMISSION_BY_KIND[PROPOSAL_TASK_KIND],
            "status": "queued",
            "message": "任务已创建",
            "logs": [
                {
                    "time": time.strftime("%H:%M:%S"),
                    "message": "开题报告生成任务已创建，分三部分生成",
                }
            ],
            "created_at": time.time(),
        }
    thread = threading.Thread(
        target=_run_proposal_task, args=(task_id, payload), daemon=True
    )
    thread.start()
    return task_id


def _set_run_font(run, font_name: str, size=None, bold=False):
    """设置 run 的拉丁和东亚字体。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        font_name,
    )
    if size:
        run.font.size = size
    run.bold = bold


def _build_proposal_docx(markdown: str) -> bytes:
    """将开题报告 markdown 转换为 DOCX，宋体正文/黑体标题。"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    lines = markdown.strip().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            heading = doc.add_heading(line.lstrip("# ").strip(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                _set_run_font(run, "黑体", Pt(16), bold=True)
        elif line.startswith("## "):
            heading = doc.add_heading(line.lstrip("# ").strip(), level=2)
            for run in heading.runs:
                _set_run_font(run, "黑体", Pt(14), bold=True)
        elif line.startswith("### "):
            heading = doc.add_heading(line.lstrip("# ").strip(), level=3)
            for run in heading.runs:
                _set_run_font(run, "黑体", Pt(13), bold=True)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line.lstrip("-* ").strip(), style="List Bullet")
            for run in p.runs:
                _set_run_font(run, "宋体", Pt(12))
        elif re.match(r"^\d+[\.\)]\s", line):
            p = doc.add_paragraph(re.sub(r"^\d+[\.\)]\s*", "", line), style="List Number")
            for run in p.runs:
                _set_run_font(run, "宋体", Pt(12))
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Cm(0.74)
            for run in p.runs:
                _set_run_font(run, "宋体", Pt(12))
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _generate_ppt(ppt_type: str, payload: Dict[str, Any]) -> bytes:
    """Generate a thesis PPTX file using the ppt_skill module and LLM for content planning."""
    import json
    from skills.ppt_skill import create_pptx

    topic = str(payload.get("topic", "")).strip()
    direction = str(payload.get("direction", "")).strip()
    methods = payload.get("methods", []) or []
    project_context = str(payload.get("project_context", "")).strip()
    chapters = payload.get("chapters", []) or []
    proposal_content = str(payload.get("proposal_content", "")).strip()

    ppt_type_labels = {
        "proposal": "开题答辩",
        "midterm": "中期答辩",
        "defense": "毕业答辩",
    }
    label = ppt_type_labels.get(ppt_type, "论文")

    # Build content context
    chapter_summaries = []
    for ch in chapters:
        content = (ch.get("content", "") or "")[:300]
        chapter_summaries.append(f"第{ch.get('number', '?')}章 {ch.get('title', '')}: {content[:200]}")
    chapter_text = "\n".join(chapter_summaries) if chapter_summaries else "暂无章节内容"

    # Determine which content to use
    if ppt_type == "proposal" and proposal_content:
        base_content = f"开题报告内容：\n{proposal_content[:3000]}"
    else:
        base_content = f"论文章节内容摘要：\n{chapter_text}"

    # Try LLM to generate slide outline if API key is configured
    config = load_llm_config()
    slide_outline = None
    if config.api_key:
        try:
            prompt = f"""请为{label}PPT生成幻灯片大纲，以JSON格式输出。

论文题目：{topic}
研究方向：{direction}
选用方法：{", ".join(methods) if methods else "未指定"}
项目背景：{project_context[:500] or "未填写"}

{base_content}

请输出如下JSON格式（只输出JSON，不要其它文字）：
{{
  "slides": [
    {{"layout": "title", "title": "{topic or '论文题目'}", "subtitle": "{direction} — {label}"}},
    {{"layout": "toc", "title": "目录", "items": ["1. 研究背景", "2. 研究方法", "3. 主要工作", "4. 总结展望"]}},
    {{"layout": "content", "title": "研究背景", "items": ["要点一", "要点二"]}},
    {{"layout": "thanks", "title": "谢谢！", "subtitle": "恳请各位老师批评指正"}}
  ]
}}

要求：8-15页幻灯片，content页3-5个要点，学术风格，中文。"""
            client, provider = _build_llm_client_and_provider(config)
            response = _llm_create_message(
                client, provider, config,
                system="",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=config.max_tokens,
            )
            text = _llm_response_text(response, provider)
            json_match = re.search(r'\{[\s\S]*\}', text)
            if json_match:
                outline = json.loads(json_match.group(0))
                slide_outline = outline.get("slides", [])
        except Exception:
            slide_outline = None

    if not slide_outline:
        # Fallback: build slides from chapters without LLM
        slide_outline = [
            {"layout": "title", "title": topic or "论文题目", "subtitle": f"{direction} — {label}"},
            {"layout": "toc", "title": "目录", "items": [ch.get("title", f"第{ch.get('number', '?')}章") for ch in chapters[:8]]},
        ]
        for ch in chapters[:10]:
            title = f"第{ch.get('number', '?')}章 {ch.get('title', '')}"
            content = (ch.get("content", "") or "")[:500]
            items = [line.strip("- ").strip() for line in content.split("\n") if line.strip()][:5]
            if not items:
                items = ["（本章内容待完善）"]
            slide_outline.append({"layout": "content", "title": title, "items": items})
        slide_outline.append({"layout": "thanks", "title": "谢谢！", "subtitle": "恳请各位老师批评指正"})

    # Generate PPTX
    buf = io.BytesIO()
    tmp_path = os.path.join(tempfile.gettempdir(), f"thesismind_ppt_{ppt_type}.pptx")
    try:
        create_pptx(slide_outline, tmp_path, title=topic or "论文", subtitle=direction)
        with open(tmp_path, "rb") as f:
            buf.write(f.read())
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    buf.seek(0)
    return buf.read()


def _create_method_card(
    name: str, phase: str, direction: str
) -> Dict[str, Any]:
    """将用户自定义方法写入方法卡片库（cards/methods/ + cards.sqlite3）。"""
    import yaml
    safe_id = re.sub(r"[^a-zA-Z0-9_一-鿿]", "", name.lower().replace(" ", "_"))[:40] or "custom"
    card_id = f"method_custom_{safe_id}"
    now = datetime.date.today().isoformat()
    slug = card_id

    frontmatter = {
        "id": slug,
        "type": "method_card",
        "name": name,
        "short_name": name[:12],
        "aliases": [name],
        "category": "hybrid",
        "phase": [phase],
        "disciplines": ["mem"],
        "domains": [direction] if direction else [],
        "applicable_sections": [],
        "pairs_with": [],
        "requires": [],
        "conflicts_with": [],
        "difficulty": "intermediate",
        "data_type": [],
        "outputs": [],
        "risk_tags": [],
        "inject_policy": "auto",
        "embedding_fields": ["name", "body"],
        "source_type": "user_created",
        "scope": "platform",
        "status": "draft",
        "version": "1.0.0",
        "created_at": now,
        "updated_at": now,
    }
    body = f"## 方法定位\n{name}是用户自定义的研究方法，适用于{phase}阶段。\n\n## 定义\n待补充。\n\n## 适用场景\n待补充。\n"

    yaml_block = yaml.dump(frontmatter, allow_unicode=True, default_flow_style=False, sort_keys=False)
    md_content = f"---\n{yaml_block}---\n{body}"

    methods_dir = PROJECT_ROOT / "cards" / "methods"
    methods_dir.mkdir(parents=True, exist_ok=True)
    card_path = methods_dir / f"{slug}.md"
    card_path.write_text(md_content, encoding="utf-8")

    # Insert into SQLite — try build_cards first, fall back to direct insert
    rebuild_ok = False
    try:
        from src.card_builder import build_cards
        result = build_cards()
        rebuild_ok = result.get("status") == "ok"
    except Exception as exc:
        logger.error("[_create_method_card] build_cards failed: %s", exc)

    if not rebuild_ok:
        # Direct insert fallback
        try:
            import json as _json_mod
            db_path = PROJECT_ROOT / "knowledge_base" / "cards.sqlite3"
            conn = sqlite3.connect(str(db_path))
            conn.execute(
                "INSERT OR REPLACE INTO cards (id, type, name, short_name, category, phase, disciplines, domains, applicable_sections, aliases, pairs_with, requires, conflicts_with, difficulty, data_type, outputs, risk_tags, source_type, scope, status, version, frontmatter_json, body_markdown, body_summary) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (
                    frontmatter["id"], frontmatter["type"], frontmatter["name"], frontmatter["short_name"],
                    frontmatter["category"], _json_mod.dumps(frontmatter["phase"], ensure_ascii=False),
                    _json_mod.dumps(frontmatter["disciplines"], ensure_ascii=False),
                    _json_mod.dumps(frontmatter["domains"], ensure_ascii=False),
                    _json_mod.dumps(frontmatter["applicable_sections"], ensure_ascii=False),
                    _json_mod.dumps(frontmatter["aliases"], ensure_ascii=False),
                    _json_mod.dumps(frontmatter["pairs_with"], ensure_ascii=False),
                    _json_mod.dumps(frontmatter["requires"], ensure_ascii=False),
                    _json_mod.dumps(frontmatter["conflicts_with"], ensure_ascii=False),
                    frontmatter["difficulty"],
                    _json_mod.dumps(frontmatter["data_type"], ensure_ascii=False),
                    _json_mod.dumps(frontmatter["outputs"], ensure_ascii=False),
                    _json_mod.dumps(frontmatter["risk_tags"], ensure_ascii=False),
                    frontmatter["source_type"], frontmatter["scope"], frontmatter["status"], frontmatter["version"],
                    _json_mod.dumps(frontmatter, ensure_ascii=False, sort_keys=True),
                    body, ""
                )
            )
            conn.commit()
            conn.close()
            rebuild_ok = True
        except Exception as exc2:
            logger.error("[_create_method_card] direct insert also failed: %s", exc2)

    if rebuild_ok:
        # Invalidate consistency engine cache
        try:
            from src.consistency_engine import _invalidate_methods_pattern_cache
            _invalidate_methods_pattern_cache()
        except Exception:
            pass
        # Clear methodology catalog cache
        catalog_path = PROJECT_ROOT / "knowledge_base" / "templates" / "methodology_catalog.json"
        if catalog_path.exists():
            catalog_path.unlink()

    return {
        "status": "ok",
        "id": slug,
        "name": name,
        "phase": [phase],
        "source_type": "user_created",
        "card_path": str(card_path),
    }


def _supplement_method(payload: Dict[str, Any]) -> Dict[str, Any]:
    """为自定义方法论爬取权威资料并生成方法卡。"""
    config = load_llm_config()
    if not config.api_key:
        return {"status": "error", "message": "未配置 API Key。"}

    method_name = str(payload.get("name", "")).strip()
    direction = str(payload.get("direction", "")).strip()
    if not method_name:
        return {"status": "error", "message": "方法名称为空。"}

    # Search knowledge base for related references
    kb_results = search_knowledge_base(f"{method_name} 方法论 应用", limit=4).get("results", [])
    kb_text = "\n".join(
        f"- {r.get('title','')}: {str(r.get('content',''))[:400]}"
        for r in kb_results
    )

    # Attempt web search for authoritative sources
    web_snippets = ""
    try:
        import urllib.request
        import urllib.parse
        query = urllib.parse.quote(f"{method_name} 研究方法 工程管理")
        req = urllib.request.Request(
            f"https://www.google.com/search?q={query}",
            headers={"User-Agent": "Mozilla/5.0"}
        )
        # Fallback: use LLM's own knowledge for authority sources
        web_snippets = ""
    except Exception:
        web_snippets = ""

    aliases_str = str(payload.get("aliases", "")).strip()
    phase = str(payload.get("phase", "discover")).strip()

    client, provider = _build_llm_client_and_provider(config)

    prompt = f"""你是一位工程管理研究方法专家。请为「{method_name}」生成一份完整的研究方法卡片。

该方法用于工程管理硕士（MEM）论文，研究方向是「{direction}」。
{f'该方法也被称为：{aliases_str}' if aliases_str else ''}

知识库相关参考资料：
{kb_text[:2000] or "无相关知识库资料。"}

请严格按照以下14个部分输出方法卡内容：

1. 方法定位：一句话说明该方法在论文中承担的角色
2. 定义：该方法的学术定义（注明提出者/机构和提出年份）
3. 适用场景：列举3-5个工程管理论文中的典型使用场景
4. 操作步骤：详细列出6-8个操作步骤
5. 常见错误：列举4-6个论文中使用该方法时常见的错误
6. 示例表达：提供一段200字左右的论文中应用该方法的写作示例
7. 不适用场景：说明该方法不适合的情况
8. 输入数据：该方法需要的输入数据类型
9. 输出结果：该方法产出的结果类型
10. 优缺点：各列2-3条
11. 适用章节：在论文哪些章节使用
12. 可搭配方法：与哪些方法可以组合使用
13. 生成约束：论文中使用该方法必须遵守的规则
14. 参考文献：提供2-4条真实的学术参考文献

对于每个部分，请提供实质性的、可操作的内容，避免空泛。格式：每部分用"## 部分名"开头，然后是具体内容。"""

    try:
        response = _llm_create_message(
            client, provider, config,
            system="你是研究方法论专家，输出结构化、学术化、可操作的方法卡内容。",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=_token_budget(config, "large"),
        )
        card_content = _llm_response_text(response, provider)

        # Generate method card file
        card_id = f"method_{_safe_method_id(method_name)}"
        yaml_frontmatter = f"""---
id: {card_id}
type: method_card
name: {method_name}
short_name: {method_name[:8]}
aliases:
  - {method_name}
category: qualitative
phase:
  - {phase}
disciplines:
  - mem
  - mba
domains:
  - {direction.replace(' ', '_')}
applicable_sections:
  - chapter_3_problem_analysis
pairs_with: []
requires: []
conflicts_with: []
difficulty: intermediate
data_type: []
sample_requirement: {{}}
evidence_level: medium
outputs: []
risk_tags: []
source_type: supplemented
scope: platform
status: draft
version: 1.0.0
created_at: {time.strftime('%Y-%m-%d')}
updated_at: {time.strftime('%Y-%m-%d')}
---
"""

        card_path = CARDS_DIR / f"{card_id}.md"
        card_path.write_text(yaml_frontmatter + "\n" + card_content, encoding="utf-8")

        # Rebuild cards database
        try:
            from src.card_builder import build_cards
            result = build_cards()
            if result.get("status") == "ok":
                # Reload methods into the running state
                pass
        except Exception:
            pass  # Card file saved, rebuild can be done manually

        return {
            "status": "ok",
            "card_id": card_id,
            "card_path": str(card_path.relative_to(PROJECT_ROOT)),
            "content_preview": card_content[:500],
        }
    except Exception as exc:
        return {"status": "error", "message": f"方法论补充失败：{exc}"}


def _safe_method_id(name: str) -> str:
    """从中文方法名提取安全的英文 ID。"""
    # Try to extract English acronyms first
    import re as _re
    eng = _re.findall(r'[A-Za-z0-9]+', name)
    if eng:
        return "_".join(e.lower() for e in eng[:2])
    # Fallback: use pinyin-like slug from common chars
    slug_map = {
        "层": "ceng", "次": "ci", "分": "fen", "析": "xi", "法": "fa",
        "模": "mo", "糊": "hu", "综": "zong", "合": "he", "评": "ping",
        "价": "jia", "问": "wen", "卷": "juan", "调": "diao", "查": "cha",
        "德": "de", "尔": "er", "菲": "fei", "实": "shi", "验": "yan",
        "研": "yan", "究": "jiu", "管": "guan", "理": "li", "质": "zhi",
        "量": "liang", "风": "feng", "险": "xian", "进": "jin", "度": "du",
        "成": "cheng", "本": "ben", "流": "liu", "程": "cheng", "优": "you",
        "化": "hua", "绩": "ji", "效": "xiao", "需": "xu", "求": "qiu",
        "网": "wang", "络": "luo", "计": "ji", "划": "hua", "技": "ji",
        "术": "shu", "数": "shu", "据": "ju", "统": "tong", "计": "ji",
    }
    slug = "".join(slug_map.get(c, "") for c in name[:6])
    return slug or f"method_{uuid.uuid4().hex[:8]}"


def _load_best_practices() -> List[Dict[str, Any]]:
    if not BP_PATH.exists():
        return []
    try:
        data = json.loads(BP_PATH.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _save_best_practice(entry: Dict[str, Any]) -> Dict[str, Any]:
    practices = _load_best_practices()
    entry["id"] = uuid.uuid4().hex[:12]
    entry["created_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
    if "tags" not in entry:
        entry["tags"] = []
    practices.insert(0, entry)
    BP_PATH.parent.mkdir(parents=True, exist_ok=True)
    BP_PATH.write_text(
        json.dumps(practices, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return entry


def _delete_best_practice(practice_id: str) -> bool:
    practices = _load_best_practices()
    filtered = [p for p in practices if p.get("id") != practice_id]
    if len(filtered) == len(practices):
        return False
    BP_PATH.write_text(
        json.dumps(filtered, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return True


def _best_practices_summary(limit: int = 20) -> str:
    practices = _load_best_practices()
    if not practices:
        return ""
    lines = ["", "## 历史最佳实践经验", ""]
    for p in practices[:limit]:
        title = p.get("title", "未命名")
        tags = " · ".join(p.get("tags", [])[:3])
        tag_str = f" [{tags}]" if tags else ""
        lines.append(f"- **{title}**{tag_str}")
    return "\n".join(lines)


def clean_generated_content(text: str, section: Dict[str, Any] | None = None) -> str:
    value = strip_markdown(text)
    number = str((section or {}).get("number", "")).strip()
    title = str((section or {}).get("title", "")).strip()
    if number:
        title = re.sub(rf"^{re.escape(number)}\s+", "", title).strip()
    title = re.sub(r"^\d+(?:\.\d+){1,3}\s+", "", title).strip()
    lines = [line.strip() for line in value.splitlines()]
    output = []
    for index, line in enumerate(lines):
        bare = re.sub(r"^\d+(?:\.\d+){1,3}\s+", "", line).strip()
        looks_like_current_heading = (
            bool(number and line.startswith(number))
            or bool(title and bare == title)
        )
        looks_like_other_heading = (
            bool(re.match(r"^第[一二三四五六七八九十\d]+章", line))
            or bool(re.match(r"^\d+(?:\.\d+){1,3}\s+\S{2,80}$", line))
        )
        # Only strip the first few lines matching the current section heading;
        # don't strip content headings like "第N章" (important for chapter arrangement)
        if index < 3 and looks_like_current_heading:
            continue
        if index < 3 and looks_like_other_heading and not looks_like_current_heading:
            # Only strip if it's the very first non-empty line (LLM echo of prompt title)
            if index == 0 or (index == 1 and not output):
                continue
        if line and output and output[-1] == line:
            continue
        output.append(line)
    return "\n".join(output).strip()


def local_expand(payload: Dict[str, Any], rewrite: bool = False) -> Dict[str, Any]:
    section = payload.get("section", {})
    chapter = payload.get("chapter", {})
    topic = payload.get("topic", "论文主题")
    title = section.get("title", "本节")
    words = section.get("estimated_words", 800)
    methods = "、".join(payload.get("methods", [])[:8]) or "未指定"
    section_prompt = clean_generated_content(payload.get("section_prompt", ""))[:1800]
    memory = merge_memory_schema(load_workspace_value("thesis_memory", {}) or {})
    project_context = (
        str(payload.get("project_context", "") or "").strip()
        or memory.get("research_context", {}).get("project_context", "")
        or load_workspace_value("project_context", "")
        or ""
    )
    if project_context:
        persist_project_context({**payload, "project_context": project_context})
    memory_brief = {
        "research_context": memory.get("research_context", {}),
        "project_context": project_context[:2200],
        "problem_list": memory.get("problem_list", [])[-8:],
        "method_usage": memory.get("method_usage", {}),
        "solution_design": memory.get("solution_design", [])[-8:],
        "evaluation_indicators": memory.get("evaluation_indicators", [])[-8:],
        "terminology": memory.get("terminology", {}),
        "outline_summary": memory.get("outline_summary", []),
        "nearby_chapter_summary": memory.get("chapter_summaries", {}).get(
            str(chapter.get("number", "")), ""
        ),
        "style_preferences": memory.get("style_preferences", {}),
    }
    section_number = str(section.get("number", chapter.get("number", "")))
    commitment_brief = build_commitment_brief(memory)
    unresolved_warning = build_unresolved_warning(memory, section_number)
    selected_methods = payload.get("methods", [])
    thesis_domain = (
        memory.get("research_context", {}).get("direction", "")
        or payload.get("direction", "")
        or "quality_management"
    )
    method_context = build_method_context(
        chapter=chapter,
        selected_methods=selected_methods if isinstance(selected_methods, list) else [],
        discipline="mem",
        domain=thesis_domain,
        stage="chapter_generation",
        max_cards=4,
    )
    query = f"{topic} {title}"

    raw_refs = search_knowledge_base(query, limit=8).get("results", [])
    references = _filter_references_by_topic(topic, raw_refs, limit=4)
    reference_text = "\n\n".join(
        f"[资料{idx}] {item.get('title')} / {item.get('path')}\n{item.get('content', '')[:700]}"
        for idx, item in enumerate(references, 1)
    )
    citations = payload.get("citations", None)
    if citations is None:
        citations = load_workspace_value("citations", []) or []
    full_citations = citations
    citation_indices = payload.get("citation_indices", None)
    if citation_indices is not None:
        if not citation_indices:
            citations = []
            numbered = []
        else:
            filtered = []
            numbered = []
            for idx in citation_indices:
                if idx < len(citations):
                    filtered.append(citations[idx])
                    numbered.append(f"[{idx+1}] {citations[idx].get('formatted', '')}")
            citations = filtered
    else:
        numbered = [f"[{i+1}] {c.get('formatted', '')}" for i, c in enumerate(citations[:30])]
    citation_text = "\n".join(numbered)

    if citations:
        if citation_indices is not None and citation_indices:
            citation_rule = f"必须引用下方引用库中的全部 {len(citations)} 篇文献，每篇至少出现一次。使用右上角标 [N] 格式标注，序号必须对应引用库中的编号。不要使用\"（作者，年份）\"格式。同一段落引用不超过3篇文献。"
        else:
            citation_rule = "引用文献时使用右上角标 [N] 格式标注（如 [1][2][3]），序号必须对应下方引用库中文献的编号。不要使用\"（作者，年份）\"格式。同一段落引用不超过3篇文献。"
        citation_section = f"\n当前引用库（请在正文中使用 [1][2] 右上角标引用）：\n{citation_text}\n"
    else:
        citation_rule = "本小节不需要引用文献，严禁添加任何引用角标或文献标注。"
        citation_section = ""

    config = load_llm_config()
    if not config.api_key:
        return {"status": "error", "content": "未配置 API Key，无法调用大模型扩写。"}

    client, provider = _build_llm_client_and_provider(config)

    task = "重写并优化" if rewrite else "扩写"

    # Build full chapter summary for context (helps with chapter-arrangement sections)
    outline = load_workspace_value("outline")
    chapter_summary = ""
    chapter_arrangement_hint = ""
    if outline:
        ch_lines = []
        for ch in outline.get("chapters", []):
            ch_num = ch.get("number", "")
            ch_title = ch.get("title", "")
            ch_lines.append(f"第{ch_num}章 {ch_title}")
        if ch_lines:
            chapter_summary = "论文完整章节结构：\n" + "\n".join(ch_lines)
    if chapter_summary and ("章节安排" in title or "论文结构" in title):
        chapter_arrangement_hint = (
            "重要：本小节是章节安排总览，必须逐一概述以下全部章节（每章1-2句话），"
            "不得跳过任何一章，不得只写其中几章。\n"
        )

    prompt = f"""请为工程管理硕士论文生成章节内容。

论文主题：{topic}
项目背景与论文思路：
{project_context[:3200] or "用户未填写。"}

{chapter_summary}

当前章节：{chapter.get("title", "")}
当前小节：{section.get("number", "")} {title}
目标字数：约 {words} 字
选用方法：{methods}
任务：{task}该小节内容。
用户对本小节的补充要求：
{section_prompt or "用户未填写。"}

论文长期记忆：
{json.dumps(memory_brief, ensure_ascii=False)[:2600]}

{commitment_brief}

{unresolved_warning}

{method_context}

{chapter_arrangement_hint}
要求：
1. 使用正式、克制、学术化的中文表达。
2. 围绕本小节展开，不要写成整章综述。
3. 可以参考本地知识库资料，但不要虚构文献出处。
4. 内容应包含问题边界、分析逻辑、方法应用和与论文主题的衔接。
5. 必须与长期记忆中的研究对象、问题、方案和指标保持一致。
6. 必须优先满足"用户对本小节的补充要求"，但不能违背论文长期记忆和已确认大纲。
7. {citation_rule}
8. 严禁使用 Markdown 语法，禁止出现 **加粗**、*斜体*、# 标题、列表符号、代码块、反引号。
9. 不要输出"以下是"等寒暄，直接给正文。
{citation_section}

本地知识库参考：
{reference_text or "未检索到高相关资料。"}
"""
    try:
        response = _llm_create_message(
            client, provider, config,
            system="你是专业的工程管理硕士论文写作助手，擅长质量管理、风险管理、流程优化和项目管理论文写作。",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=_token_budget(config, "default"),
        )
        content = clean_generated_content(_llm_response_text(response, provider), section)
        consistency = verify_commitments(content, memory, section_number)
        citation_check = verify_citations(
            content,
            citation_indices if citation_indices is not None else [],
            full_citations if full_citations else [],
        )
        return {
            "status": "ok",
            "content": content,
            "consistency": consistency,
            "citation_check": citation_check,
            "references": [
                {
                    "title": item.get("title"),
                    "path": item.get("path"),
                    "score": item.get("score"),
                }
                for item in references
            ],
        }
    except Exception as exc:
        return {"status": "error", "content": f"大模型调用失败：{exc}"}


def compile_thesis() -> Dict[str, Any]:
    """编译全部草稿为完整论文，并运行跨章节一致性校验。"""
    drafts = load_drafts()
    outline = load_workspace_value("outline")
    memory = merge_memory_schema(load_workspace_value("thesis_memory", {}) or {})

    if not outline:
        return {"status": "error", "message": "未找到大纲，请先生成章节大纲。"}

    # 按章节/小节顺序组装全文
    chapters_output: List[Dict[str, Any]] = []
    full_text_parts: List[str] = []
    total_words = 0
    word_count_by_chapter: Dict[str, int] = {}

    for ch in outline.get("chapters", []):
        ch_num = str(ch.get("number", ""))
        ch_title = ch.get("title", "")
        ch_text_parts = [f"第{ch_num}章 {ch_title}", ""]
        ch_word_count = 0

        for sec in ch.get("sections", []):
            sec_num = sec.get("number", "")
            sec_title = sec.get("title", "")
            draft_key = f"{ch_num}.{sec_num}"
            content = drafts.get(draft_key, "").strip()

            if content:
                ch_text_parts.append(f"{sec_num} {sec_title}")
                ch_text_parts.append(content)
                ch_text_parts.append("")
                ch_word_count += len(content)

            # Subsections
            for sub_idx, sub in enumerate(sec.get("subsections", []) or []):
                sub_draft_key = f"{ch_num}.{sec_num}.{sub_idx}"
                sub_content = drafts.get(sub_draft_key, "").strip()
                if sub_content:
                    sub_title = sub.get("title", "")
                    ch_text_parts.append(f"{sub_title}")
                    ch_text_parts.append(sub_content)
                    ch_text_parts.append("")
                    ch_word_count += len(sub_content)

        ch_text = "\n".join(ch_text_parts).strip()
        full_text_parts.append(ch_text)
        chapters_output.append({
            "chapter": ch_num,
            "title": ch_title,
            "word_count": ch_word_count,
            "has_content": ch_word_count > 0,
        })
        word_count_by_chapter[ch_num] = ch_word_count
        total_words += ch_word_count

    full_text = "\n\n".join(full_text_parts)

    # 跨章节一致性校验：每章验证前文章节的承诺
    chapter_consistency: Dict[str, Any] = {}
    total_hard_unresolved = 0
    for ch in outline.get("chapters", []):
        ch_num = str(ch.get("number", ""))
        # 收集该章全部内容
        ch_content_parts = []
        for sec in ch.get("sections", []):
            dk = f"{ch_num}.{sec.get('number', '')}"
            c = drafts.get(dk, "")
            if c:
                ch_content_parts.append(c)
            for sub_idx in range(len(sec.get("subsections", []) or [])):
                sc = drafts.get(f"{ch_num}.{sec.get('number', '')}.{sub_idx}", "")
                if sc:
                    ch_content_parts.append(sc)
        ch_content = "\n".join(ch_content_parts)
        result = verify_commitments(ch_content, memory, ch_num)
        chapter_consistency[ch_num] = result
        total_hard_unresolved += result.get("hard_unresolved", 0)

    # 检查大纲覆盖度（哪些章节/小节还没有草稿）
    missing_sections: List[Dict[str, str]] = []
    for ch in outline.get("chapters", []):
        ch_num = str(ch.get("number", ""))
        for sec in ch.get("sections", []):
            sec_num = sec.get("number", "")
            dk = f"{ch_num}.{sec_num}"
            if not drafts.get(dk, "").strip():
                missing_sections.append({
                    "chapter": ch_num,
                    "section": sec_num,
                    "title": sec.get("title", ""),
                })

    result = {
        "status": "ok",
        "total_words": total_words,
        "word_count_by_chapter": word_count_by_chapter,
        "chapters": chapters_output,
        "missing_sections": missing_sections,
        "missing_count": len(missing_sections),
        "chapter_consistency": chapter_consistency,
        "total_hard_unresolved": total_hard_unresolved,
        "full_text_preview": full_text[:2000],
    }

    # 缓存编译结果
    save_workspace_value("thesis_compiled", {
        "full_text": full_text,
        "compiled_at": time.time(),
        "total_words": total_words,
    })
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    (OUTPUT_ROOT / "thesis_compiled.md").write_text(full_text, encoding="utf-8")

    return result


def check_blind_review_risks(payload: Dict[str, Any]) -> Dict[str, Any]:
    """
    盲审风险检查。支持单章或全文多章扫描。
    payload: {content, chapter_title, chapter_number, methods} 或 {chapters: [{title, number, content}]}
    """
    methods = payload.get("methods", [])
    severity_filter = payload.get("severity_filter")
    category_filter = payload.get("category_filter")

    chapters = payload.get("chapters")
    if chapters:
        # 多章扫描：逐章检查并汇总
        all_results = []
        total_triggered = 0
        chapter_results = []
        for ch in chapters:
            content = ch.get("content", "")
            if not content.strip():
                continue
            if methods:
                scan = run_method_risk_scan(content, methods, chapter_title=ch.get("title", ""))
            else:
                scan = run_risk_scan(
                    content,
                    chapter_title=ch.get("title", ""),
                    chapter_number=ch.get("number", ""),
                    severity_filter=severity_filter,
                    category_filter=category_filter,
                )
            triggered = [r for r in scan.get("results", []) if r["triggered"]]
            all_results.extend(scan.get("results", []))
            total_triggered += len(triggered)
            chapter_results.append({
                "chapter_title": ch.get("title", ""),
                "chapter_number": ch.get("number", ""),
                "total_risks": scan.get("total_risks", 0),
                "triggered": len(triggered),
                "results": scan.get("results", []),
            })

        # 去重汇总
        seen_ids = set()
        unique_results = []
        for r in all_results:
            if r["risk_id"] not in seen_ids:
                seen_ids.add(r["risk_id"])
                unique_results.append(r)

        critical_count = len([r for r in unique_results if r.get("triggered") and r.get("severity") == "critical"])
        high_count = len([r for r in unique_results if r.get("triggered") and r.get("severity") == "high"])

        return {
            "status": "ok",
            "total_risks": len(unique_results),
            "triggered": total_triggered,
            "critical_count": critical_count,
            "high_count": high_count,
            "results": unique_results,
            "chapter_results": chapter_results,
            "summary": f"全文扫描触发 {total_triggered} 项风险（致命 {critical_count}，高 {high_count}）",
        }

    # 单章扫描
    content = payload.get("content", "")
    if not content:
        return {"status": "error", "message": "请提供待检查的内容"}

    chapter_title = payload.get("chapter_title", "")
    chapter_number = payload.get("chapter_number", "")

    if methods:
        scan = run_method_risk_scan(content, methods, chapter_title=chapter_title)
    else:
        scan = run_risk_scan(
            content,
            chapter_title=chapter_title,
            chapter_number=chapter_number,
            severity_filter=severity_filter,
            category_filter=category_filter,
        )

    return {
        **scan,
        "formatted_report": format_risk_report(scan),
    }


def check_aigc_rate(payload: Dict[str, Any]) -> Dict[str, Any]:
    """AIGC率评估：扫描章节草稿，检测AI生成特征并评估风险等级。"""
    import re
    chapters = payload.get("chapters", []) or []
    if not chapters:
        return {"status": "error", "message": "暂无已写内容可供检测"}

    # AI-generated text indicators — based on research from ImBD (AAAI 2025),
    # MPU (ICLR 2024), StyleDecipher, and Chinese academic AIGC detection literature.
    # Categories: lexical, syntactic, discourse, stylistic.
    ai_patterns = [
        # === Lexical: overused connectors & fillers (AI多用连词/严谨指示词) ===
        (r"值得注意的是[,，]", "AI高频转折"),
        (r"总的来说[,，]", "AI总结句式"),
        (r"综上所述[,，]", "AI总结句式"),
        (r"此外[,，]", "AI过渡词密度"),
        (r"另外[,，]", "AI过渡词密度"),
        (r"与此同时[,，]", "AI并列过渡"),
        (r"不可忽视的是", "AI强调句式"),
        (r"不容忽视的是", "AI强调句式"),
        (r"需要指出的是", "AI严谨指示词"),
        (r"需要强调的是", "AI严谨指示词"),
        (r"具有重要的意义", "AI空洞表达"),
        (r"具有重要的(理论|现实|实践)意义", "AI空洞表达"),
        (r"具有十分重要的意义", "AI空洞表达"),
        (r"发挥着(至关)?重要的作用", "AI空洞表达"),
        (r"在当今(社会|时代|背景下)", "AI开篇模板"),
        (r"随着[^。]{1,40}的(不断|持续|深入)发展", "AI背景套话"),
        (r"在(新|当前)(时代|形势|阶段)下", "AI时代套话"),
        (r"可以这样说[,，]", "AI口语句式"),
        (r"换句话说[,，]", "AI释义句式"),
        (r"简而言之[,，]", "AI释义句式"),
        (r"毋庸置疑[,，]", "AI断言句式"),
        (r"显而易见[,，]", "AI断言句式"),
        # === Syntactic: template sentence structures ===
        (r"不仅[^。]{1,30}而且", "AI递进句式"),
        (r"既[^。]{1,20}又[^。]{1,20}", "AI并列句式"),
        (r"一方面[^。]{1,30}另一方面", "AI正反句式"),
        (r"首先[^。]{1,20}其次[^。]{1,20}最后", "AI列举模板"),
        (r"第一[,，][^。]{1,20}第二[,，][^。]{1,20}第三", "AI列举模板"),
        (r"通过[^。]{1,30}可以看出", "AI分析句式"),
        (r"通过对[^。]{1,30}的分析", "AI分析句式"),
        (r"为[^。]{1,30}提供了(有力|重要|坚实)的(支撑|依据|保障)", "AI结论套话"),
        # === Discourse: paragraph-level patterns ===
        (r"^本文[^。]{1,50}(旨在|通过|基于|从|以)", "AI论文开篇"),
        (r"^随着[^。]{1,40}，[^。]{1,30}问题(日益|越来越|逐渐)", "AI论文开篇递进"),
        (r"^在[^。]{1,40}的(大背景|背景|语境)下", "AI语境铺垫"),
        # === Stylistic: impersonal / overly formal ===
        (r"从某种(意义|程度)上(说|讲)", "AI模糊限定"),
        (r"在一定(程度|意义)上", "AI模糊限定"),
        (r"需要(我们)?(清醒地)?认识到", "AI说教句式"),
        (r"必须(清醒地)?(认识|意识)到", "AI说教句式"),
    ]
    # Human-like indicators — specific references, concrete details, personal voice
    human_patterns = [
        (r"例如[^，,]{1,30}(项目|案例|企业|公司|单位|部门)", "具体案例引用"),
        (r"(比如|譬如|像)[^，,]{1,30}(项目|案例|企业|公司)", "具体举例"),
        (r"\d{4}年", "时间具体化"),
        (r"\d+年\d+月", "精确时间"),
        (r"(第|该|本)[^，]{1,10}(项目|案例|企业|公司|单位|工程)", "实体指代"),
        (r"(笔者|本人|我们|我)(在|参与|负责|主持)", "亲身参与"),
        (r"(调查|访谈|实地|问卷|走访)了?\d+", "实证数据痕迹"),
        (r"以[^，]{1,20}(项目|工程|企业)为例", "案例详述"),
        (r"(数据|资料|信息)(来源|取自|来自|源于)", "数据来源说明"),
        (r"(如表|见图|如图|见表)[\d一二三四五六七八九十]", "图表引用"),
        (r"(根据|依据|按照)[^，]{1,30}(标准|规范|规定|要求)", "标准规范引用"),
    ]

    total_score = 0
    total_chars = 0
    chapter_results = []

    for ch in chapters:
        content = ch.get("content", "").strip()
        if not content:
            continue
        chars = len(content)
        total_chars += chars

        chapter_score = 0
        results = []
        for pattern, label in ai_patterns:
            matches = list(re.finditer(pattern, content))
            count = len(matches)
            if count > 0:
                density = count / max(chars, 1) * 1000
                triggered = density > 0.15 or count >= 2
                chapter_score += count * 3
                snippets = []
                for m in matches[:3]:
                    start = max(0, m.start() - 20)
                    end = min(len(content), m.end() + 30)
                    snippets.append(content[start:end].replace("\n", " "))
                results.append({
                    "risk_id": f"aigc_{label}",
                    "label": label,
                    "triggered": triggered,
                    "severity": "medium" if triggered else "low",
                    "detail": f"出现 {count} 次（密度 {density:.2f}/千字）",
                    "snippet": " ... ".join(snippets[:2]) if snippets else "",
                })

        # Human indicator bonus (negative score = good)
        for pattern, label in human_patterns:
            matches = list(re.finditer(pattern, content))
            count = len(matches)
            if count > 0:
                chapter_score -= count * 2
                results.append({
                    "risk_id": f"human_{label}",
                    "label": f"✓ {label}",
                    "triggered": False,
                    "severity": "none",
                    "detail": f"发现 {count} 处具体指代（人工特征）",
                    "snippet": "",
                })

        # Sentence length variance check
        sentences = re.split(r'[。！？；\n]', content)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 5]
        if sentences:
            lengths = [len(s) for s in sentences]
            avg_len = sum(lengths) / len(lengths)
            if len(lengths) > 5:
                variance = sum((l - avg_len) ** 2 for l in lengths) / len(lengths)
                cv = (variance ** 0.5) / avg_len if avg_len > 0 else 0
                if cv < 0.3 and avg_len > 60:
                    chapter_score += 10
                    results.append({
                        "risk_id": "aigc_uniform_sentences",
                        "label": "句式单一",
                        "triggered": True,
                        "severity": "high",
                        "detail": f"平均句长 {avg_len:.0f} 字，变异系数仅 {cv:.2f}（AI特征：句长过于均匀）",
                        "snippet": "",
                    })

        # Paragraph start repetition check
        para_starts = []
        for para in content.split("\n"):
            para = para.strip()
            if len(para) > 10:
                para_starts.append(para[:8])
        if len(para_starts) > 2:
            unique_starts = len(set(para_starts))
            if unique_starts < len(para_starts) * 0.5:
                chapter_score += 8
                results.append({
                    "risk_id": "aigc_repetitive_starts",
                    "label": "段落开头重复",
                    "triggered": True,
                    "severity": "medium",
                    "detail": f"{len(para_starts)} 段落中仅 {unique_starts} 种不同开头（AI特征：模板化开头）",
                    "snippet": ", ".join(list(set(para_starts))[:3]),
                })

        # Lexical diversity check (type-token ratio)
        # Research shows AI text has lower lexical diversity than human text (1.72x less).
        words = re.findall(r"[一-鿿]+|[a-zA-Z]+", content)
        if len(words) > 20:
            types = len(set(words))
            ttr = types / len(words)
            if ttr < 0.35:
                chapter_score += 10
                results.append({
                    "risk_id": "aigc_low_lexical_diversity",
                    "label": "词汇多样性低",
                    "triggered": True,
                    "severity": "high",
                    "detail": f"词汇多样性 TTR 仅 {ttr:.2f}（{len(words)} 词中仅 {types} 种），AI文本常用词重复率高",
                    "snippet": "",
                })
            elif ttr < 0.50:
                chapter_score += 4
                results.append({
                    "risk_id": "aigc_moderate_lexical_diversity",
                    "label": "词汇多样性偏低",
                    "triggered": True,
                    "severity": "medium",
                    "detail": f"词汇多样性 TTR {ttr:.2f}（{len(words)} 词中 {types} 种），低于人类写作平均水平",
                    "snippet": "",
                })

        # Paragraph length uniformity
        para_lengths = [len(p.strip()) for p in content.split("\n") if len(p.strip()) > 20]
        if len(para_lengths) >= 3:
            avg_pl = sum(para_lengths) / len(para_lengths)
            para_var = sum((l - avg_pl) ** 2 for l in para_lengths) / len(para_lengths)
            para_cv = (para_var ** 0.5) / avg_pl if avg_pl > 0 else 0
            if para_cv < 0.35:
                chapter_score += 6
                results.append({
                    "risk_id": "aigc_uniform_paragraphs",
                    "label": "段落长度均匀",
                    "triggered": True,
                    "severity": "medium",
                    "detail": f"{len(para_lengths)} 段落长度变异系数仅 {para_cv:.2f}（AI倾向产出整齐划一的段落）",
                    "snippet": "",
                })

        # Determine risk level (floor at 0 — human indicators can't make score negative)
        chapter_score = max(0, chapter_score)
        density = chapter_score / max(chars, 1) * 1000
        if density < 15:
            risk_level = "低"
        elif density < 35:
            risk_level = "中"
        else:
            risk_level = "高"

        total_score += chapter_score
        # Build marked spans for frontend highlighting
        spans = []
        for r in results:
            if not r.get("triggered"):
                continue
            # Re-scan to get exact positions
            for pattern, label in ai_patterns:
                if label == r["label"]:
                    for m in re.finditer(pattern, content):
                        spans.append({
                            "start": m.start(),
                            "end": m.end(),
                            "label": label,
                            "reason": r["detail"],
                            "severity": r["severity"],
                        })
                    break
        # Sentence uniformity: mark all sentences
        if any(r["risk_id"] == "aigc_uniform_sentences" and r["triggered"] for r in results):
            # Mark the whole paragraph as suspicious (no specific span)
            pass

        # Sort spans by position
        spans.sort(key=lambda s: s["start"])
        # Merge overlapping spans
        merged = []
        for s in spans:
            if merged and s["start"] <= merged[-1]["end"] + 5:
                merged[-1]["end"] = max(merged[-1]["end"], s["end"])
                merged[-1]["reason"] += "；" + s["reason"]
            else:
                merged.append(s.copy())

        chapter_results.append({
            "chapter_title": ch.get("title", ""),
            "chapter_number": ch.get("number", ""),
            "score": round(density, 1),
            "risk_level": risk_level,
            "full_text": content,
            "spans": merged,
            "results": results,
        })

    # Overall assessment (floor at 0)
    total_score = max(0, total_score)
    overall_density = total_score / max(total_chars, 1) * 1000
    if overall_density < 15:
        overall_risk = "低"
        interpretation = "文本AI生成特征较少，人工撰写痕迹明显，盲审风险较低"
    elif overall_density < 35:
        overall_risk = "中"
        interpretation = "部分段落存在AI生成特征，建议对标注段落进行人工润色"
    else:
        overall_risk = "高"
        interpretation = "文本AI生成特征显著，强烈建议使用AIGC降重功能进行改写"

    return {
        "status": "ok",
        "overall_score": round(overall_density, 1),
        "risk_level": overall_risk,
        "interpretation": interpretation,
        "chapter_results": chapter_results,
        "summary": f"AIGC率得分 {overall_density:.1f}（{overall_risk}风险）：{interpretation}",
    }


def reduce_aigc_rate(payload: Dict[str, Any]) -> Dict[str, Any]:
    """AIGC降重：对高AIGC率段落调用LLM进行改写。"""
    import re
    chapters = payload.get("chapters", []) or []
    if not chapters:
        return {"status": "error", "message": "暂无已写内容可供降重"}

    config = load_llm_config()
    if not config.api_key:
        return {"status": "error", "message": "请先在基本配置中设置API Key后再进行AIGC降重"}

    # First run AIGC check to identify risky paragraphs
    check_result = check_aigc_rate({"chapters": chapters})
    if check_result.get("status") != "ok":
        return check_result

    # Collect high-risk paragraphs
    risky_paragraphs = []
    for ch_result in check_result.get("chapter_results", []):
        if ch_result.get("risk_level") != "高":
            continue
        risky_paragraphs.append({
            "chapter_title": ch_result.get("chapter_title", ""),
            "chapter_number": ch_result.get("chapter_number", ""),
            "score": ch_result.get("score", 0),
        })

    if not risky_paragraphs:
        return {
            "status": "ok",
            "summary": "未检测到高风险段落，无需降重",
            "results": [],
        }

    # Re-read full chapter content for risky chapters and rewrite with LLM
    results = []
    for i, ch in enumerate(chapters):
        content = ch.get("content", "").strip()
        if not content:
            continue
        # Check if this chapter had high risk
        matching = [r for r in risky_paragraphs if r.get("chapter_number") == ch.get("number", "")]
        if not matching:
            continue

        # Split into paragraphs and rewrite each
        paras = [p.strip() for p in content.split("\n") if len(p.strip()) > 50]
        if not paras:
            continue

        rewritten_paras = []
        for pi, para in enumerate(paras[:5]):  # Limit to 5 paragraphs per chapter
            ai_indicators = [
                "值得注意的是", "总的来说", "在当今", "具有重要的意义",
                "不可忽视的是", "首先", "其次", "最后", "此外",
            ]
            if not any(ind in para for ind in ai_indicators):
                continue

            try:
                prompt = f"""请改写以下学术段落，降低AI生成痕迹，使其更自然、更像人类撰写。要求：
1. 保持原意和学术性不变
2. 使用更多样化的句式和表达
3. 避免模板化过渡词（如"值得注意的是""总的来说"等）
4. 加入具体细节或例子的感觉（如有数据可保留）
5. 句长要有变化，长短交错
6. 只输出改写后的段落，不要加任何解释

原文段落：
{para[:800]}

改写后："""
                client, provider = _build_llm_client_and_provider(config)
                response = _llm_create_message(
                    client, provider, config,
                    system="",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=_token_budget(config, "default"),
                )
                rewritten = _llm_response_text(response, provider).strip()

                results.append({
                    "chapter_title": ch.get("title", ""),
                    "section_label": f"段落{pi + 1}",
                    "original": para[:500],
                    "rewritten": rewritten[:1000],
                })
            except Exception as e:
                results.append({
                    "chapter_title": ch.get("title", ""),
                    "section_label": f"段落{pi + 1}",
                    "original": para[:300],
                    "rewritten": f"[改写失败: {str(e)[:100]}]",
                })

    return {
        "status": "ok",
        "summary": f"已完成 {len(results)} 段改写" if results else "未找到需要改写的段落",
        "results": results,
    }


def generate_subsections(payload: Dict[str, Any]) -> Dict[str, Any]:
    topic = payload.get("topic", "论文主题")
    project_context = persist_project_context(payload)
    chapter = payload.get("chapter", {})
    methods = "、".join(payload.get("methods", [])[:8]) or "未指定"
    config = load_llm_config()
    if not config.api_key:
        raise RuntimeError("未配置 API Key")

    prompt = f"""请为论文单章生成三级目录。

论文主题：{topic}
项目背景与论文思路：
{project_context[:2200] or "用户未填写。"}

章节：{json.dumps(chapter, ensure_ascii=False)}
已选方法：{methods}

要求：
1. 只补充每个二级小节下的 1-3 个三级小节。
2. 不修改章标题和二级小节标题。
3. 严禁使用 Markdown 语法，标题中不得出现 #、*、**、反引号。
4. 只输出 JSON，结构为：
{{"sections":[{{"title":"二级标题","subsections":[{{"title":"三级标题"}}]}}]}}
"""
    fallback = False
    try:
        client, provider = _build_llm_client_and_provider(config)
        response = _llm_create_message(
            client, provider, config,
            system="你是工程管理硕士论文目录设计助手。",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=_token_budget(config, "small"),
        )
        raw = _extract_json_object(_llm_response_text(response, provider))
        sections = raw.get("sections", [])
        if not sections:
            raise ValueError("LLM returned empty sections array")
    except Exception as exc:
        logger.warning("三级目录 LLM 生成失败，启用本地模板兜底: %s", exc)
        fallback = True
        sections = [
            {
                "title": section.get("title", ""),
                "subsections": [
                    {"title": f"{section.get('title', '本节')}的现状分析"},
                    {"title": f"{section.get('title', '本节')}的关键问题"},
                ],
            }
            for section in chapter.get("sections", [])
        ]

    for section_index, section in enumerate(chapter.get("sections", []), 1):
        if "本章小结" in str(section.get("title", "")):
            section["subsections"] = []
            continue
        matched = (
            sections[section_index - 1] if section_index - 1 < len(sections) else {}
        )
        section["subsections"] = [
            {
                "level": 3,
                "number": f"{chapter.get('number', 1)}.{section_index}.{idx}",
                "title": strip_markdown(item.get("title", str(item))),
                "estimated_words": max(
                    200, int(section.get("estimated_words", 600) / 2)
                ),
            }
            for idx, item in enumerate(matched.get("subsections", [])[:3], 1)
        ]
    existing_summary = [
        section
        for section in chapter.get("sections", [])
        if "本章小结" in str(section.get("title", ""))
    ]
    if not existing_summary:
        chapter.setdefault("sections", []).append(
            {
                "level": 2,
                "number": f"{chapter.get('number', 1)}.{len(chapter.get('sections', [])) + 1}",
                "title": "本章小结",
                "estimated_words": 500,
                "subsections": [],
            }
        )
    return {"status": "ok", "chapter": chapter, "fallback": fallback}


class ThesisMindHandler(BaseHTTPRequestHandler):
    def do_GET(self) -> None:
        try:
            path = self.path.split("?", 1)[0].rstrip("/") or "/"
            if path.startswith("/api/") and not _check_license_api(self, path):
                return
            if path == "/api/outlines":
                from urllib.parse import parse_qs, urlparse
                qs = parse_qs(urlparse(self.path).query)
                direction = qs.get("direction", [""])[0]
                if direction:
                    outlines = query_outlines_by_direction(direction)
                    _json_response(self, {"direction": direction, "outlines": outlines, "count": len(outlines)})
                else:
                    catalog = load_catalog()
                    _json_response(self, catalog)
                return

            if path == "/api/config":
                config = load_llm_config()
                _json_response(
                    self,
                    {
                        "provider": config.provider,
                        "model": config.model,
                        "base_url": config.base_url,
                        "api_key_configured": bool(config.api_key),
                        "api_key_preview": _mask_secret(config.api_key),
                        "auth_mode": config.auth_mode,
                        "max_tokens": config.max_tokens,
                        "directions": DIRECTIONS,
                    },
                )
                return
            if path == "/api/license/status":
                manager = LicenseManager()
                _json_response(self, manager.get_license_status())
                return
            if path.startswith("/api/methodologies"):
                if "/generate-summaries" in self.path:
                    items = scan_methodologies(force=False)
                    cached = _load_method_summaries()
                    missing = [it["name"] for it in items if not it.get("summary") and not cached.get(it["name"])]
                    if missing:
                        new_summaries = _generate_method_summaries_via_llm(missing)
                        cached.update(new_summaries)
                        _save_method_summaries(cached)
                        _json_response(self, {"status": "ok", "generated": len(new_summaries), "total": len(cached)})
                        return
                    return
                force = "refresh=1" in self.path
                _json_response(self, {"items": scan_methodologies(force=force)})
                return
            if path == "/api/method-catalog":
                from src.method_registry import get_method_catalog
                _json_response(self, {"catalog": get_method_catalog()})
                return
            if path == "/api/projects":
                _json_response(
                    self,
                    {
                        "current_project_id": current_project_id(),
                        "projects": list_projects(),
                    },
                )
                return
            if path.startswith("/api/tasks/"):
                task_id = path.rsplit("/", 1)[-1]
                with TASK_LOCK:
                    task = TASKS.get(task_id)
                if not task:
                    _json_response(
                        self,
                        {"status": "error", "message": "task not found"},
                        status=404,
                    )
                    return
                permission_menu = task.get("permission_menu") or TASK_PERMISSION_BY_KIND.get(
                    task.get("kind", ""), "workflow"
                )
                if not _check_license_menu(self, permission_menu):
                    return
                _json_response(self, {"task_id": task_id, **task})
                return
            if path == "/api/workspace":
                _json_response(
                    self,
                    {
                        "outline": load_workspace_value("outline"),
                        "markdown": load_workspace_value("markdown", ""),
                        "project_context": load_workspace_value("project_context", ""),
                        "current_project_id": current_project_id(),
                        "current_direction": load_workspace_value("current_direction"),
                        "projects": list_projects(),
                        "citations": load_workspace_value("citations", []),
                        "paper_citations": load_workspace_value("paper_citations", []),
                        "section_citations": load_workspace_value("section_citations", {}),
                        "drafts": load_drafts(),
                        "thesis_memory": load_workspace_value("thesis_memory", {}),
                        "phase_methods": _normalize_phase_methods(load_workspace_value("phase_methods", {"discover": [], "solve": [], "validate": []})),
                        "method_pools": {
                            "discover": _normalize_method_pool(load_workspace_value("method_pool:discover", [])),
                            "solve": _normalize_method_pool(load_workspace_value("method_pool:solve", [])),
                            "validate": _normalize_method_pool(load_workspace_value("method_pool:validate", [])),
                        },
                        "proposal_content": load_workspace_value("proposal_content", ""),
                        "framework_svg": load_workspace_value("framework_svg", ""),
                        "framework_topic": load_workspace_value("framework_topic", ""),
                        "framework_direction": load_workspace_value("framework_direction", ""),
                        "framework_phase_methods": load_workspace_value("framework_phase_methods", {}),
                    },
                )
                return
            if path == "/api/workspace/value":
                from urllib.parse import parse_qs, urlparse
                qs = parse_qs(urlparse(self.path).query)
                key = (qs.get("key", [""])[0] or "").strip()
                if not key:
                    _json_response(self, {"status": "error", "message": "missing key"}, status=400)
                    return
                _json_response(self, {"status": "ok", "key": key, "value": load_workspace_value(key)})
                return
            if path == "/api/export/docx":
                outline = load_workspace_value("outline")
                if not outline:
                    _json_response(self, {"error": "没有可导出的大纲"}, status=400)
                    return
                citations = load_workspace_value("citations", []) or []
                file_path = export_docx(outline, load_drafts(), citations)
                _file_response(
                    self,
                    file_path,
                    "thesis_export.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                return
            if path == "/api/export/pdf":
                outline = load_workspace_value("outline")
                if not outline:
                    _json_response(self, {"error": "没有可导出的大纲"}, status=400)
                    return
                citations = load_workspace_value("citations", []) or []
                file_path = export_pdf(outline, load_drafts(), citations)
                _file_response(self, file_path, "thesis_export.pdf", "application/pdf")
                return
            if path == "/api/citation-cards/stats":
                from src.paper_store import PAPER_DB_PATH
                conn = sqlite3.connect(str(PAPER_DB_PATH))
                total = conn.execute("SELECT COUNT(*) FROM citation_cards").fetchone()[0]
                verified = conn.execute("SELECT COUNT(*) FROM citation_cards WHERE verified = 1").fetchone()[0]
                rows = conn.execute(
                    "SELECT direction_label, COUNT(*) as cnt FROM citation_cards WHERE direction_label != '' GROUP BY direction_label ORDER BY cnt DESC"
                ).fetchall()
                dir_dist = {row[0]: row[1] for row in rows}
                rows = conn.execute(
                    "SELECT ref_type, COUNT(*) as cnt FROM citation_cards GROUP BY ref_type ORDER BY cnt DESC"
                ).fetchall()
                type_dist = {row[0]: row[1] for row in rows}
                conn.close()
                _json_response(self, {
                    "total": total,
                    "verified": verified,
                    "direction_distribution": dir_dist,
                    "type_distribution": type_dist,
                })
                return

            if path.startswith("/api/citation-cards"):
                from urllib.parse import parse_qs, urlparse
                from src.paper_store import PAPER_DB_PATH, _card_row_to_dict, _expand_method_terms
                qs = parse_qs(urlparse(self.path).query)
                direction = qs.get("direction", [""])[0]
                method = qs.get("method", [""])[0]
                keyword = qs.get("keyword", [""])[0]
                verified_str = qs.get("verified", [""])[0]
                ref_type = qs.get("ref_type", [""])[0]
                year = qs.get("year", [""])[0]
                min_quality = float(qs.get("min_quality", [0])[0])
                offset = int(qs.get("offset", [0])[0])
                limit = min(int(qs.get("limit", [50])[0]), 200)

                conn = sqlite3.connect(str(PAPER_DB_PATH))
                conn.row_factory = sqlite3.Row

                conditions = []
                params: List[Any] = []

                if direction:
                    dir_map = {v: k for k, v in {
                        "quality_management": "质量管理",
                        "risk_management": "风险管理",
                        "schedule_management": "进度管理",
                        "requirements_management": "需求管理",
                        "process_optimization": "流程优化",
                        "cost_management": "成本管理",
                        "supply_chain_logistics": "供应链物流",
                    }.items()}
                    dir_id = dir_map.get(direction, direction)
                    conditions.append("direction_id = ?")
                    params.append(dir_id)
                if method:
                    method_clauses = []
                    for term in _expand_method_terms(method):
                        method_clauses.append("methods_json LIKE ?")
                        params.append(f"%{term}%")
                    conditions.append("(" + " OR ".join(method_clauses) + ")")
                if keyword:
                    conditions.append("(title LIKE ? OR formatted LIKE ?)")
                    params.extend([f"%{keyword}%", f"%{keyword}%"])
                if verified_str:
                    conditions.append("verified = ?")
                    params.append(int(verified_str))
                if ref_type:
                    conditions.append("ref_type = ?")
                    params.append(ref_type)
                if year:
                    conditions.append("year = ?")
                    params.append(year)
                if min_quality > 0:
                    conditions.append("quality_score >= ?")
                    params.append(min_quality)

                where = (" WHERE " + " AND ".join(conditions)) if conditions else ""
                count_sql = f"SELECT COUNT(*) FROM citation_cards{where}"
                total = conn.execute(count_sql, params).fetchone()[0]

                query_sql = f"SELECT * FROM citation_cards{where} ORDER BY quality_score DESC LIMIT ? OFFSET ?"
                rows = conn.execute(query_sql, params + [limit, offset]).fetchall()
                conn.close()
                cards = [_card_row_to_dict(r) for r in rows]
                _json_response(self, {"cards": cards, "total": total, "offset": offset, "limit": limit})
                return

            if path == "/api/best-practices":
                tag = None
                try:
                    from urllib.parse import parse_qs, urlparse
                    qs = parse_qs(urlparse(self.path).query)
                    tag = qs.get("tag", [None])[0]
                except Exception:
                    pass
                practices = _load_best_practices()
                if tag:
                    practices = [p for p in practices if tag in (p.get("tags") or [])]
                _json_response(self, {"practices": practices, "count": len(practices)})
                return

            if path == "/api/paper-library":
                citations = _load_workspace_raw("paper_citations", [])
                _json_response(self, {"citations": citations})
                return

            # ── Paper management (Pro+) ──
            if path == "/api/papers/stats":
                if not _check_license_api(self, "/api/papers/stats"):
                    return
                from src.paper_store import get_paper_stats
                _json_response(self, {"status": "ok", **get_paper_stats()})
                return

            if path == "/api/papers/list":
                if not _check_license_api(self, "/api/papers/list"):
                    return
                from urllib.parse import parse_qs, urlparse
                from src.paper_store import list_all_papers
                qs = parse_qs(urlparse(self.path).query)
                direction = qs.get("direction", [""])[0]
                keyword = qs.get("keyword", [""])[0]
                limit = int(qs.get("limit", ["50"])[0])
                offset = int(qs.get("offset", ["0"])[0])
                result = list_all_papers(direction_id=direction, keyword=keyword, limit=limit, offset=offset)
                _json_response(self, {"status": "ok", **result})
                return

            if path == "/api/papers/detail":
                if not _check_license_api(self, "/api/papers/detail"):
                    return
                from urllib.parse import parse_qs, urlparse
                from src.paper_store import get_paper_by_id
                qs = parse_qs(urlparse(self.path).query)
                doc_id = qs.get("doc_id", [""])[0]
                if not doc_id:
                    _json_response(self, {"status": "error", "message": "缺少 doc_id"}, status=400)
                    return
                paper = get_paper_by_id(doc_id)
                if paper:
                    _json_response(self, {"status": "ok", "paper": paper})
                else:
                    _json_response(self, {"status": "error", "message": "论文不存在"}, status=404)
                return

            if path == "/api/license/history":
                if not _check_license_api(self, "/api/license/history"):
                    return
                manager = LicenseManager()
                _json_response(self, {"history": manager.load_license_history()})
                return

            # ── PPT Engine API (read-only) ──
            if path == "/api/ppt/templates":
                from ppt_engine import list_defense_templates
                _json_response(self, {"templates": list_defense_templates()})
                return

            if path.startswith("/api/ppt/download/"):
                filename = path.rsplit("/", 1)[-1]
                if not filename or ".." in filename:
                    self.send_error(400)
                    return
                filepath = OUTPUT_DIR / "ppt" / filename
                if not filepath.exists():
                    self.send_error(404)
                    return
                body = filepath.read_bytes()
                self.send_response(200)
                self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.presentationml.presentation")
                self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
                self.send_header("Content-Length", str(len(body)))
                self.end_headers()
                self.wfile.write(body)
                return

            path = _safe_static_path(path)
            if not path.exists():
                self.send_error(404)
                return
            content_type = (
                mimetypes.guess_type(str(path))[0] or "application/octet-stream"
            )
            body = path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
        except Exception as exc:
            _json_response(self, {"error": str(exc)}, status=500)

    def do_POST(self) -> None:
        try:
            path = self.path.split("?", 1)[0].rstrip("/") or "/"
            if path.startswith("/api/") and not _check_license_api(self, path):
                return

            # Multipart upload handling (paper upload)
            content_type = self.headers.get("Content-Type", "")
            if content_type.startswith("multipart/form-data"):
                payload = _read_multipart(self)
                if path == "/api/papers/upload":
                    if not _check_license_api(self, "/api/papers/upload"):
                        return
                    _json_response(self, _handle_paper_upload(payload))
                    return
                if path == "/api/table/generate":
                    _handle_table_generate(self, payload)
                    return
                _json_response(self, {"status": "error", "message": "未知的 multipart 端点"}, status=400)
                return

            payload = _read_json(self)
            if path == "/api/config":
                api_key = str(payload.get("api_key", "")).strip()
                max_tokens = str(payload.get("max_tokens", "4000")).strip()
                values = {
                    "LLM_PROVIDER": str(payload.get("provider", "deepseek")).strip()
                    or "deepseek",
                    "ANTHROPIC_BASE_URL": str(payload.get("base_url", "")).strip(),
                    "ANTHROPIC_MODEL": str(payload.get("model", "")).strip()
                    or "deepseek-v4-pro",
                    "MODEL_ID": str(payload.get("model", "")).strip()
                    or "deepseek-v4-pro",
                    "ANTHROPIC_AUTH_MODE": "api_key",
                    "LLM_MAX_TOKENS": max_tokens,
                }
                if api_key:
                    values["ANTHROPIC_API_KEY"] = api_key
                _write_env_values(values)
                from dotenv import load_dotenv as _reload_dotenv
                _reload_dotenv(ENV_PATH, override=True)
                _json_response(
                    self,
                    {
                        "status": "ok",
                        "provider": values["LLM_PROVIDER"],
                        "model": values["ANTHROPIC_MODEL"],
                        "base_url": values["ANTHROPIC_BASE_URL"],
                        "api_key_configured": bool(api_key)
                        or load_llm_config().api_key is not None,
                        "max_tokens": int(max_tokens),
                    },
                )
                return

            if path == "/api/license/activate":
                code = str(payload.get("code", "")).strip()
                if not code:
                    _json_response(self, {"error": "请输入许可证激活码"}, status=400)
                    return
                manager = LicenseManager()
                ok, msg = manager.save_license(code, validate=True)
                if ok:
                    _json_response(self, {"status": "ok", "message": msg, "license": manager.get_license_status()})
                else:
                    _json_response(self, {"error": msg}, status=400)
                return

            if path == "/api/license/trial":
                ok, msg = TrialLicense.start()
                if ok:
                    manager = LicenseManager()
                    _json_response(self, {"status": "ok", "message": msg, "license": manager.get_license_status()})
                else:
                    _json_response(self, {"error": msg}, status=400)
                return

            if path == "/api/projects/create":
                project = create_project(str(payload.get("topic", "") or ""))
                _json_response(
                    self,
                    {"status": "ok", "project": project, "projects": list_projects()},
                )
                return

            if path == "/api/projects/switch":
                set_current_project(str(payload.get("project_id", "")))
                _json_response(
                    self,
                    {
                        "status": "ok",
                        "current_project_id": current_project_id(),
                        "projects": list_projects(),
                    },
                )
                return

            if path == "/api/project-context":
                persist_project_context(payload)
                _json_response(self, {"status": "ok"})
                return

            if path == "/api/workspace/save":
                key = str(payload.get("key", "")).strip()
                if not key or key.startswith("__"):
                    _json_response(self, {"status": "error", "message": "invalid workspace key"}, status=400)
                    return
                save_workspace_value(key, payload.get("value"))
                _json_response(self, {"status": "ok", "key": key})
                return

            if path == "/api/framework":
                persist_project_context(payload)
                touch_current_project(payload.get("topic", ""))
                topic = payload.get("topic", "")
                direction = payload.get("direction_name", "")
                raw_phase_methods = payload.get("phase_methods") or {
                    "discover": payload.get("methods", []),
                    "solve": [],
                    "validate": [],
                }
                # Normalize to canonical IDs for storage; keep raw names for SVG display
                phase_methods = _filter_stale_methods(raw_phase_methods)
                svg = make_framework_svg(topic, direction, raw_phase_methods)
                OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
                (OUTPUT_ROOT / "web_research_framework.svg").write_text(
                    svg, encoding="utf-8"
                )
                # Auto-persist to workspace for revisit auto-fill
                save_workspace_value("framework_svg", svg)
                save_workspace_value("framework_topic", topic)
                save_workspace_value("framework_direction", direction)
                save_workspace_value("framework_phase_methods", phase_methods)
                _json_response(
                    self,
                    {
                        "svg": svg,
                        "output_path": "output/web_research_framework.svg",
                    },
                )
                return

            if path == "/api/framework/save":
                phase_methods = _filter_stale_methods(payload.get("phase_methods", {}) or {})
                save_workspace_value("framework_svg", payload.get("svg", ""))

                save_workspace_value("framework_topic", payload.get("topic", ""))
                save_workspace_value("framework_direction", payload.get("direction", ""))
                save_workspace_value("framework_phase_methods", phase_methods)
                # Merge framework into thesis_memory for consistency tracking
                flat_methods: list = []
                for phase_list in phase_methods.values():
                    if isinstance(phase_list, list):
                        flat_methods.extend(phase_list)
                build_thesis_memory({
                    "topic": payload.get("topic", ""),
                    "direction": payload.get("direction", ""),
                    "methods": flat_methods,
                    "phase_methods": phase_methods,
                })
                stale = _mark_all_drafts_stale("研究框架已更新，本章可能需要重新生成以保持一致性")
                _json_response(self, {"status": "ok", "stale_chapters": stale})
                return

            if path == "/api/table/generate":
                _json_response(
                    self,
                    {"status": "error", "message": "请使用 multipart/form-data 上传 Excel 文件"},
                    status=400,
                )
                return

            if path == "/api/table/generate-from-text":
                _handle_table_generate_from_text(self, payload)
                return

            if path == "/api/outline":
                persist_project_context(payload)
                touch_current_project(payload.get("topic", ""))
                task_id = start_outline_task(payload)
                _json_response(self, {"task_id": task_id, "status": "queued"})
                return

            if path == "/api/domain-templates/build":
                _json_response(
                    self, {"status": "ok", "result": build_domain_templates()}
                )
                return

            if path == "/api/citation-index/build":
                _json_response(self, {"status": "ok", "result": build_citation_index()})
                return

            if path in {"/api/knowledge-base/init", "/api/knowledge_base/init"}:
                if not _check_license_api(self, "/api/knowledge-base/init"):
                    return
                task_id = start_knowledge_base_init_task()
                _json_response(self, {"status": "ok", "task_id": task_id})
                return

            # ── Paper management ──
            if path == "/api/papers/pipeline/start":
                if not _check_license_api(self, "/api/papers/pipeline/start"):
                    return
                doc_id = payload.get("doc_id", "") if payload else ""
                task_id = start_paper_pipeline_task(doc_id=doc_id)
                _json_response(self, {"status": "ok", "task_id": task_id})
                return

            if path == "/api/papers/delete":
                if not _check_license_api(self, "/api/papers/delete"):
                    return
                doc_id = (payload or {}).get("doc_id", "")
                if not doc_id:
                    _json_response(self, {"status": "error", "message": "缺少 doc_id"}, status=400)
                    return
                from src.paper_store import delete_paper
                ok = delete_paper(doc_id)
                _json_response(self, {"status": "ok" if ok else "error", "deleted": ok})
                return

            # ── License management ──
            if path == "/api/license/generate":
                if not _check_license_api(self, "/api/license/generate"):
                    return
                license_type = (payload or {}).get("license_type", "basic")
                user_email = (payload or {}).get("user_email", "")
                if license_type not in {"basic", "pro", "vip", "admin"}:
                    _json_response(self, {"status": "error", "message": "无效的许可证类型"}, status=400)
                    return
                manager = LicenseManager()
                try:
                    code = manager.generate_license(license_type=license_type, user_email=user_email)
                    _json_response(self, {"status": "ok", "license_code": code})
                except Exception as e:
                    _json_response(self, {"status": "error", "message": str(e)}, status=500)
                return

            if path == "/api/outlines":
                from urllib.parse import parse_qs, urlparse
                qs = parse_qs(urlparse(self.path).query)
                direction = qs.get("direction", [""])[0]
                if direction:
                    outlines = query_outlines_by_direction(direction)
                    _json_response(self, {"direction": direction, "outlines": outlines, "count": len(outlines)})
                else:
                    catalog = load_catalog()
                    _json_response(self, catalog)
                return

            if path == "/api/paper-library/stats":
                from src.paper_store import get_paper_stats, get_papers_by_direction
                stats = get_paper_stats()
                _json_response(self, {"status": "ok", **stats})
                return

            if path == "/api/paper-library/papers":
                from src.paper_store import get_papers_by_direction
                dir_id = payload.get("direction", "")
                papers = get_papers_by_direction(dir_id, limit=50) if dir_id else []
                _json_response(self, {"status": "ok", "papers": papers})
                return

            if path == "/api/citations/generate":
                task_id = start_citation_generate_task(payload)
                _json_response(self, {"status": "queued", "task_id": task_id})
                return

            if path == "/api/citations/save":
                citations = _normalize_citations(payload.get("citations", []))
                save_workspace_value("citations", citations)
                _json_response(self, {"status": "ok", "citations": citations})
                return

            if path == "/api/outline/save":
                persist_project_context(payload)
                outline = payload.get("outline")
                markdown = outline_to_markdown(outline) if outline else ""
                save_workspace_value("outline", outline)
                save_workspace_value("markdown", markdown)
                if outline:
                    build_thesis_memory(payload, outline)
                if payload.get("skip_stale"):
                    stale = []
                else:
                    stale = _mark_all_drafts_stale("章节大纲已更新，本章可能需要重新生成以保持一致性")
                _json_response(self, {"status": "ok", "markdown": markdown, "stale_chapters": stale})
                return

            if path == "/api/subsections":
                _json_response(self, generate_subsections(payload))
                return

            if path == "/api/expand":
                persist_project_context(payload)
                _json_response(self, local_expand(payload, rewrite=False))
                return

            if path == "/api/rewrite":
                persist_project_context(payload)
                _json_response(self, local_expand(payload, rewrite=True))
                return

            if path == "/api/drafts/save":
                result = save_draft(
                    str(payload.get("draft_key", "")), str(payload.get("content", ""))
                )
                _json_response(self, result)
                return

            if path == "/api/chat/test":
                _json_response(self, _test_llm_connection())
                return

            if path == "/api/chat":
                _json_response(self, _chat_with_llm(payload))
                return

            if path == "/api/consistency/validate":
                memory = merge_memory_schema(load_workspace_value("thesis_memory", {}) or {})
                drafts = load_drafts()
                results = {}
                total_unresolved = 0
                for draft_key, draft_content in drafts.items():
                    ch_result = verify_commitments(draft_content, memory, draft_key)
                    results[draft_key] = ch_result
                    total_unresolved += ch_result.get("unresolved", 0)
                _json_response(self, {
                    "status": "ok",
                    "total_unresolved": total_unresolved,
                    "by_section": results,
                })
                return

            if path == "/api/thesis/compile":
                _json_response(self, compile_thesis())
                return

            if path == "/api/method-supplement":
                _json_response(self, _supplement_method(payload))
                return

            if path == "/api/method-assignments/save":
                phase_methods = payload.get("phase_methods") or {}
                save_workspace_value("phase_methods", _normalize_phase_methods({
                    "discover": phase_methods.get("discover", []),
                    "solve": phase_methods.get("solve", []),
                    "validate": phase_methods.get("validate", []),
                }))
                _json_response(self, {"status": "ok"})
                return

            if path == "/api/method-pool/save":
                method_pool = payload.get("method_pool") or []
                phase = str(payload.get("phase", "")).strip()
                if phase in ("discover", "solve", "validate"):
                    save_workspace_value(f"method_pool:{phase}", _normalize_method_pool(method_pool))
                else:
                    # 兼容旧版：无 phase 时保存到全局
                    save_workspace_value("method_pool", _normalize_method_pool(method_pool))
                _json_response(self, {"status": "ok"})
                return

            if path == "/api/methods/create-card":
                name = str(payload.get("name", "")).strip()
                phase = str(payload.get("phase", "discover")).strip()
                direction = str(payload.get("direction", "")).strip()
                if not name:
                    _json_response(self, {"status": "error", "message": "方法名不能为空"}, status=400)
                    return
                result = _create_method_card(name, phase, direction)
                _json_response(self, result)
                return

            if path == "/api/methods/save":
                methods = payload.get("methods", [])
                if isinstance(methods, list):
                    save_workspace_value("selected_methods", methods)
                phase_methods = payload.get("phase_methods") or {}
                if isinstance(phase_methods, dict):
                    save_workspace_value("phase_methods", _normalize_phase_methods({
                        "discover": phase_methods.get("discover", []),
                        "solve": phase_methods.get("solve", []),
                        "validate": phase_methods.get("validate", []),
                    }))
                _json_response(self, {"status": "ok"})
                return


            if path == "/api/proposal":
                if not _check_license_api(self, path):
                    return
                task_id = start_proposal_task(payload)
                _json_response(self, {"status": "queued", "task_id": task_id})
                return

            if path == "/api/proposal/save":
                content = str(payload.get("content", "")).strip()
                save_workspace_value("proposal_content", content)
                _json_response(self, {"status": "ok"})
                return

            if path == "/api/proposal/export":
                if not _check_license_api(self, path):
                    return
                content = str(payload.get("content", "")).strip()
                if not content:
                    _json_response(self, {"status": "error", "message": "内容为空"}, status=400)
                    return
                docx_bytes = _build_proposal_docx(content)
                self.send_response(200)
                self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                self.send_header("Content-Disposition", 'attachment; filename="proposal.docx"')
                self.send_header("Content-Length", str(len(docx_bytes)))
                self.end_headers()
                self.wfile.write(docx_bytes)
                return

            if path == "/api/ppt/generate":
                if not _check_license_api(self, path):
                    return
                ppt_type = str(payload.get("ppt_type", "defense")).strip()
                pptx_bytes = _generate_ppt(ppt_type, payload)
                filename_map = {"proposal": "proposal.pptx", "midterm": "midterm.pptx", "defense": "defense.pptx"}
                filename = filename_map.get(ppt_type, "thesis.pptx")
                self.send_response(200)
                self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.presentationml.presentation")
                self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
                self.send_header("Content-Length", str(len(pptx_bytes)))
                self.end_headers()
                self.wfile.write(pptx_bytes)
                return

            if path == "/api/blind-review-check":
                if not _check_license_api(self, path):
                    return
                _json_response(self, check_blind_review_risks(payload))
                return

            if path == "/api/aigc/check":
                if not _check_license_api(self, path):
                    return
                _json_response(self, check_aigc_rate(payload))
                return

            if path == "/api/aigc/reduce":
                if not _check_license_api(self, path):
                    return
                _json_response(self, reduce_aigc_rate(payload))
                return

            if path == "/api/citation-cards/update":
                from src.paper_store import PAPER_DB_PATH
                card_id = str(payload.get("card_id", "")).strip()
                if not card_id:
                    _json_response(self, {"status": "error", "message": "card_id required"}, status=400)
                    return
                conn = sqlite3.connect(str(PAPER_DB_PATH))
                allowed = ["formatted", "title", "authors", "year", "ref_type", "verified", "direction_id", "direction_label", "theory_tags_json"]
                for field, value in payload.items():
                    if field in allowed and field != "card_id":
                        conn.execute(f"UPDATE citation_cards SET {field} = ? WHERE card_id = ?", (value, card_id))
                conn.commit()
                conn.close()
                _json_response(self, {"status": "ok", "card_id": card_id})
                return

            if path == "/api/citation-cards/batch":
                from src.paper_store import PAPER_DB_PATH, _card_row_to_dict
                card_ids = payload.get("card_ids", []) or []
                cards = []
                if card_ids:
                    conn = sqlite3.connect(str(PAPER_DB_PATH))
                    conn.row_factory = sqlite3.Row
                    placeholders = ",".join(["?" for _ in card_ids])
                    rows = conn.execute(
                        f"SELECT * FROM citation_cards WHERE card_id IN ({placeholders})",
                        card_ids,
                    ).fetchall()
                    conn.close()
                    cards = [_card_row_to_dict(r) for r in rows]
                _json_response(self, {"cards": cards, "total": len(cards)})
                return

            if path == "/api/citation-cards/classify":
                task_id = start_classify_citations_task(payload)
                _json_response(self, {"status": "queued", "task_id": task_id})
                return

            if path == "/api/citation-cards/scan-verify":
                task_id = _start_scan_verify_task(payload)
                _json_response(self, {"status": "queued", "task_id": task_id})
                return

            if path == "/api/citation-cards/toggle-verify":
                from src.paper_store import PAPER_DB_PATH
                card_ids = payload.get("card_ids", [])
                verified = 1 if payload.get("verified", True) else 0
                if not card_ids:
                    _json_response(self, {"status": "error", "message": "card_ids required"}, status=400)
                    return
                conn = sqlite3.connect(str(PAPER_DB_PATH))
                placeholders = ",".join("?" for _ in card_ids)
                conn.execute(f"UPDATE citation_cards SET verified = ? WHERE card_id IN ({placeholders})", [verified] + card_ids)
                conn.commit()
                conn.close()
                _json_response(self, {"status": "ok", "count": len(card_ids), "verified": bool(verified)})
                return

            if path == "/api/citation-cards/delete":
                from src.paper_store import PAPER_DB_PATH
                card_ids = payload.get("card_ids", [])
                if not card_ids:
                    _json_response(self, {"status": "error", "message": "card_ids required"}, status=400)
                    return
                conn = sqlite3.connect(str(PAPER_DB_PATH))
                placeholders = ",".join("?" for _ in card_ids)
                conn.execute(f"DELETE FROM citation_cards WHERE card_id IN ({placeholders})", card_ids)
                conn.commit()
                conn.close()
                _json_response(self, {"status": "ok", "deleted": len(card_ids)})
                return

            if path == "/api/best-practices":
                action = str(payload.get("action", "save")).strip()
                if action == "delete":
                    ok = _delete_best_practice(str(payload.get("id", "")))
                    _json_response(self, {"status": "ok" if ok else "error", "message": "已删除" if ok else "未找到该条目"})
                else:
                    entry = _save_best_practice({
                        "title": str(payload.get("title", "")).strip()[:120] or "未命名",
                        "content": str(payload.get("content", "")).strip(),
                        "tags": payload.get("tags", []) if isinstance(payload.get("tags"), list) else [],
                        "source_message": payload.get("source_message", ""),
                    })
                    _json_response(self, {"status": "ok", "entry": entry})
                return

            # --- Citation checklist persistence ---

            if path == "/api/workspace/save-checklists":
                key = str(payload.get("key", "paper_citations"))
                value = payload.get("value", [])
                save_workspace_value(key, value)
                _json_response(self, {"status": "ok"})
                return

            # --- Paper citation library APIs (merged citation workspace) ---

            if path == "/api/paper-library/sync":
                citations = payload.get("citations", [])
                _save_workspace_raw("paper_citations", citations)
                _json_response(self, {"status": "ok", "count": len(citations)})
                return

            if path == "/api/paper-library/add":
                new_citations = payload.get("citations", [])
                existing = _load_workspace_raw("paper_citations", [])
                existing_formatted = {c.get("formatted", "") for c in existing}
                added = 0
                for c in new_citations:
                    if c.get("formatted", "") not in existing_formatted:
                        existing.append(c)
                        existing_formatted.add(c.get("formatted", ""))
                        added += 1
                _save_workspace_raw("paper_citations", existing)
                _json_response(self, {"status": "ok", "added": added, "total": len(existing)})
                return

            if path == "/api/paper-library/remove":
                remove_formatted = set(payload.get("formatted_texts", []))
                existing = _load_workspace_raw("paper_citations", [])
                kept = [c for c in existing if c.get("formatted", "") not in remove_formatted]
                _save_workspace_raw("paper_citations", kept)
                _json_response(self, {"status": "ok", "removed": len(existing) - len(kept), "total": len(kept)})
                return

            if path == "/api/paper-library/update":
                old_formatted = payload.get("old_formatted", "")
                updated = payload.get("citation", {})
                existing = _load_workspace_raw("paper_citations", [])
                for i, c in enumerate(existing):
                    if c.get("formatted", "") == old_formatted:
                        existing[i] = updated
                        break
                _save_workspace_raw("paper_citations", existing)
                _json_response(self, {"status": "ok"})
                return

            if path == "/api/citations/dedup":
                texts = payload.get("texts", [])
                if not texts:
                    _json_response(self, {"duplicates": []})
                    return
                import sqlite3 as _sqlite3
                from src.paper_store import PAPER_DB_PATH as _PAPER_DB
                _conn = _sqlite3.connect(str(_PAPER_DB))
                dupes = set()
                for text in texts:
                    key = text.strip()[:200]
                    if key:
                        row = _conn.execute(
                            "SELECT 1 FROM citation_cards WHERE formatted LIKE ? LIMIT 1",
                            (f"%{key[:80]}%",),
                        ).fetchone()
                        if row:
                            dupes.add(text)
                _conn.close()
                # Also check against paper library
                library = _load_workspace_raw("paper_citations", [])
                lib_texts = {c.get("formatted", "") for c in library}
                dupes |= {t for t in texts if t in lib_texts}
                _json_response(self, {"duplicates": list(dupes)})
                return

            if path == "/api/citation-cards/insert":
                from src.paper_store import PAPER_DB_PATH as _PAPER_DB2, upsert_citation_card as _upsert
                citations = payload.get("citations", [])
                if not citations:
                    _json_response(self, {"status": "ok", "inserted": 0})
                    return
                inserted = 0
                for c in citations:
                    card = {
                        "formatted": c.get("formatted", ""),
                        "title": c.get("title", c.get("formatted", "")[:120]),
                        "authors": c.get("authors", ""),
                        "year": c.get("year", ""),
                        "ref_type": c.get("ref_type", "期刊文章"),
                        "verified": c.get("verified", 0),
                        "quality_score": c.get("quality_score", 0.0),
                        "paper_id": c.get("paper_id", ""),
                        "direction_id": c.get("direction_id", ""),
                        "direction_label": c.get("direction_label", ""),
                        "methods": c.get("methods", []),
                        "source_paper_title": c.get("source_paper_title", ""),
                        "source_section": c.get("source_section", "用户添加"),
                    }
                    _upsert(card, db_path=_PAPER_DB2)
                    inserted += 1
                _json_response(self, {"status": "ok", "inserted": inserted})
                return

            if path == "/api/citation-cards/parse-and-add":
                raw_text = str(payload.get("raw_text", "")).strip()
                if not raw_text:
                    _json_response(self, {"status": "error", "message": "引用文本不能为空"}, status=400)
                    return
                # Split on blank lines for multi-citation paste
                blocks = [b.strip() for b in re.split(r"\n\s*\n", raw_text) if b.strip()]
                if len(blocks) > 20:
                    _json_response(self, {"status": "error", "message": "最多支持20条引用"}, status=400)
                    return
                config = load_llm_config()
                config.model = "deepseek-v4-flash"
                if not config.api_key:
                    _json_response(self, {"status": "error", "message": "未配置 API Key"}, status=500)
                    return
                client, provider = _build_llm_client_and_provider(config)

                numbered = "\n\n".join(
                    f"[{i+1}] {b}" for i, b in enumerate(blocks)
                )
                parse_prompt = (
                    f"你是学术引用解析专家。请解析以下参考文献，提取结构化字段。\n\n"
                    f"每条引用的字段：\n"
                    f"- formatted: 完整的格式化引用文本（保留原始格式）\n"
                    f"- title: 文献题名\n"
                    f"- authors: 作者（多个用分号分隔）\n"
                    f"- year: 出版年份（如 \"2024\"）\n"
                    f"- ref_type: 期刊文章/学位论文/会议论文/图书/标准/报告/专利/报纸/电子资源/其他\n"
                    f"- language: zh（中文）或 en（英文）\n\n"
                    f"规则：\n"
                    f"1. 支持 GB/T 7714、APA、MLA、Chicago 等任何引用格式\n"
                    f"2. 无法确定的字段留空字符串\n"
                    f"3. 如果文本不是参考文献，所有字段留空\n\n"
                    f"参考文献文本：\n{numbered}\n\n"
                    f"只输出JSON数组（不要markdown代码块）：\n"
                    f'[{{"formatted": "...", "title": "...", "authors": "...", "year": "...", "ref_type": "期刊文章", "language": "zh"}}]'
                )
                try:
                    response = _llm_create_message(
                        client, provider, config,
                        system="You are a citation parser. Extract structured fields from citation text. Return ONLY JSON.",
                        messages=[{"role": "user", "content": parse_prompt}],
                        max_tokens=min(config.max_tokens, 200 + len(blocks) * 300),
                        disable_thinking=True,
                    )
                    raw_resp = _llm_response_text(response, provider).strip()
                    raw_resp = re.sub(r"^```(?:json)?\s*", "", raw_resp)
                    raw_resp = re.sub(r"\s*```$", "", raw_resp)
                    parsed = json.loads(raw_resp)
                    if isinstance(parsed, dict):
                        parsed = [parsed]
                except Exception:
                    _json_response(self, {"status": "error", "message": "LLM解析失败，请检查引用格式"}, status=500)
                    return

                from src.paper_store import upsert_citation_card as _upsert, PAPER_DB_PATH as _PAPER_DB2
                inserted = 0
                for item in parsed:
                    text = (item.get("formatted") or item.get("raw") or "").strip()
                    if not text:
                        continue
                    # Skip non-citation text (too short, URLs, chapter headers)
                    if len(text) < 15 or text.startswith("http"):
                        continue
                    card = {
                        "formatted": text,
                        "title": item.get("title", "") or text[:120],
                        "authors": item.get("authors", ""),
                        "year": item.get("year", ""),
                        "ref_type": item.get("ref_type", "期刊文章"),
                        "language": item.get("language", "zh"),
                        "verified": 0,
                        "quality_score": 0.0,
                        "paper_id": "",
                        "direction_id": "",
                        "direction_label": "",
                        "methods": [],
                        "source_section": "用户添加",
                    }
                    _upsert(card, db_path=_PAPER_DB2)
                    inserted += 1
                _json_response(self, {"status": "ok", "inserted": inserted})
                return

            self.send_error(404)
        except PayloadTooLarge as exc:
            _json_response(self, {"error": str(exc)}, status=413)
        except Exception as exc:
            _json_response(self, {"error": str(exc)}, status=500)

    def log_message(self, format: str, *args: Any) -> None:
        return


def _rebuild_cards_from_source() -> None:
    """Rebuild cards.sqlite3 from the permanent cards/methods/ directory.

    Only rebuilds when one or more .md source files are newer than
    cards.sqlite3 (or cards.sqlite3 doesn't exist yet).
    """
    import src.card_builder as card_builder
    methods_dir = PROJECT_ROOT / "cards" / "methods"
    db_path = PROJECT_ROOT / "knowledge_base" / "cards.sqlite3"

    newest_source = 0.0
    for md_file in methods_dir.glob("method_*.md"):
        mtime = md_file.stat().st_mtime
        if mtime > newest_source:
            newest_source = mtime

    db_mtime = db_path.stat().st_mtime if db_path.exists() else 0.0

    if newest_source <= db_mtime:
        logger.info("卡片库无需重建（%s 张卡片已是最新）",
                     _count_cards_in_db(db_path))
        return

    saved_methods_root = card_builder.METHODS_ROOT
    try:
        card_builder.METHODS_ROOT = methods_dir
        result = card_builder.build_cards()
        logger.info("卡片库重建: %s 张卡片", result.get('cards_processed', 0))
    except Exception as exc:
        logger.error("卡片库重建失败: %s", exc)
    finally:
        card_builder.METHODS_ROOT = saved_methods_root


def _count_cards_in_db(db_path: Path) -> int:
    try:
        conn = sqlite3.connect(str(db_path))
        count = conn.execute("SELECT COUNT(*) FROM cards").fetchone()[0]
        conn.close()
        return count
    except Exception:
        return 0


def main() -> None:
    parser = argparse.ArgumentParser(description="Run ThesisMind local web app")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8765)
    args = parser.parse_args()

    # 加载 .env 到环境变量
    from dotenv import load_dotenv
    if ENV_PATH.exists():
        load_dotenv(ENV_PATH)

    # 从永久 cards/methods/ 目录重建卡片库
    _rebuild_cards_from_source()

    ThreadingHTTPServer.allow_reuse_address = True
    server = ThreadingHTTPServer((args.host, args.port), ThesisMindHandler)
    logger.info("ThesisMind web app: http://%s:%s", args.host, args.port)

    try:
        server.serve_forever()
    except KeyboardInterrupt:
        logger.info("服务关闭")


if __name__ == "__main__":
    main()
