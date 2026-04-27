#!/usr/bin/env python3
"""
docx2md.py - 高质量 Word → Markdown 转换

相比 pandoc：
- 表格转为简洁的 | 列 | 格式，不会生成 grid table
- 跳过文档信息表、参考文档等非核心内容（可配置）
- 合并单元格正确处理
- 代码块自动识别
- 支持批量转换

用法:
    python docx2md.py 需求文档.docx                   # 转换单个文件
    python docx2md.py 需求文档.docx -o output.md      # 指定输出路径
    python docx2md.py *.docx                           # 批量转换
    python docx2md.py 需求文档.docx --skip-meta        # 跳过元信息表格
"""

import re
import sys
from pathlib import Path


def cell_text(cell) -> str:
    """提取单元格纯文本，合并多段落。"""
    parts = []
    for para in cell.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    return " ".join(parts).replace("\n", " ")


def _all_same(items: list) -> bool:
    """判断列表中所有非空值是否相同（用于检测合并单元格展开）。"""
    vals = [i.strip() for i in items if i and i.strip()]
    return len(vals) > 1 and len(set(vals)) == 1


def table_to_md(table) -> str:
    """
    把 Word 表格转为 Markdown pipe table。
    处理合并单元格（取第一个非空值）。
    """
    if not table.rows:
        return ""

    # 提取所有行数据
    rows_data = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cells.append(cell_text(cell))
        rows_data.append(cells)

    if not rows_data:
        return ""

    # 列数统一（防止参差不齐）
    max_cols = max(len(r) for r in rows_data)
    rows_data = [r + [""] * (max_cols - len(r)) for r in rows_data]

    # 去掉完全空的列
    non_empty_cols = [
        i for i in range(max_cols)
        if any(rows_data[r][i].strip() for r in range(len(rows_data)))
    ]
    if not non_empty_cols:
        return ""
    rows_data = [[row[i] for i in non_empty_cols] for row in rows_data]
    col_count = len(non_empty_cols)

    # 如果只有一列且内容简单，输出为列表
    if col_count == 1:
        lines = []
        for row in rows_data:
            t = row[0].strip()
            if t:
                lines.append(f"- {t}")
        return "\n".join(lines)

    # 检查是否是键值表（2列，左列是标签）
    is_kv = (
        col_count == 2 and
        len(rows_data) <= 15 and
        all(len(r[0]) < 20 for r in rows_data if r[0])
    )
    if is_kv:
        lines = []
        for row in rows_data:
            k, v = row[0].strip(), row[1].strip()
            if k or v:
                lines.append(f"**{k}**：{v}" if k else v)
        return "\n\n".join(lines)

    # 标准表格
    lines = []
    header = rows_data[0]
    # 合并单元格去重：如果一行所有列值相同，压缩为单列
    if _all_same(header):
        header = [header[0]]
        rows_data = [[r[0]] for r in rows_data]
    lines.append("| " + " | ".join(h or " " for h in header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows_data[1:]:
        # 跳过全空行
        if not any(c.strip() for c in row):
            continue
        # 跳过内容与表头全相同的行（合并单元格残留）
        if _all_same(row) and row[0].strip() == header[0].strip():
            continue
        lines.append("| " + " | ".join(c.replace("|", "\\|") or " " for c in row) + " |")

    return "\n".join(lines)


def para_to_md(para) -> str:
    """把段落转为 Markdown，处理标题/列表/加粗等。"""
    text = para.text.strip()
    if not text:
        return ""

    style_name = ""
    try:
        if para.style and para.style.name:
            style_name = para.style.name
    except Exception:
        pass

    # 跳过目录标题行
    if text.strip() in ("目 录", "目录", "目  录", "目 次"):
        return ""

    # 标题
    heading_map = {
        "Heading 1": "#", "heading 1": "#", "标题 1": "#",
        "Heading 2": "##", "heading 2": "##", "标题 2": "##",
        "Heading 3": "###", "heading 3": "###", "标题 3": "###",
        "Heading 4": "####", "heading 4": "####", "标题 4": "####",
    }
    for style, prefix in heading_map.items():
        if style.lower() in style_name.lower():
            # 清理标题里的编号（如 "1. 需求背景" → "需求背景"）
            clean = re.sub(r'^[\d\.]+\s*', '', text).strip()
            return f"{prefix} {clean or text}"

    # 列表项
    if style_name in ("List Paragraph", "列表段落") or style_name.startswith("List"):
        return f"- {text}"

    # 纯数字编号开头的行（如 "1. 需求背景"）转为标题
    m = re.match(r'^(\d+)\.\s+(.+)$', text)
    if m and len(text) < 50:
        num = len(m.group(1))
        prefix = "#" * min(num + 1, 4)
        return f"{prefix} {m.group(2)}"

    # 处理行内加粗
    result = ""
    for run in para.runs:
        run_text = run.text
        if not run_text:
            continue
        if run.bold:
            result += f"**{run_text}**"
        elif run.italic:
            result += f"*{run_text}*"
        else:
            result += run_text

    return result.strip() or text


# 非核心内容的标题关键词（跳过这些章节）
SKIP_SECTION_KW = [
    "文档信息", "修订记录", "参考文档", "相关干系人",
    "附录", "变更历史", "版本历史",
]


# 判断为元信息表格的内容关键词
_META_TABLE_KW = [
    "文档名称", "文档编号", "修订历史", "版本号", "撰写人", "负责人",
    "文档目的", "内容描述", "修订者", "评审人员",
]


def should_skip_table(table, prev_heading: str = "") -> bool:
    """判断表格是否是非核心内容（文档信息/修订记录等）。"""
    if any(kw in prev_heading for kw in SKIP_SECTION_KW):
        return True
    # 检查表格内容：如果多行含元信息关键词，跳过
    if table.rows:
        meta_hits = 0
        for row in table.rows[:5]:
            row_text = " ".join(cell_text(c) for c in row.cells)
            if any(kw in row_text for kw in _META_TABLE_KW):
                meta_hits += 1
        if meta_hits >= 2:  # 至少2行命中，确认为元信息表
            return True
    return False


def docx_to_md(docx_path: Path, skip_meta: bool = True) -> str:
    """主转换函数。"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要 python-docx: pip install python-docx")
        sys.exit(1)

    doc = Document(str(docx_path))
    lines = []
    prev_heading = ""
    table_set = set(id(t) for t in doc.tables)
    processed_tables = set()

    # 按文档顺序遍历段落和表格
    from docx.oxml.ns import qn
    body = doc.element.body

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # 段落
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            md_line = para_to_md(para)
            if not md_line:
                if lines and lines[-1] != "":
                    lines.append("")
                continue

            # 记录当前标题，用于判断是否跳过后续表格
            if md_line.startswith("#"):
                prev_heading = md_line.lstrip("# ").strip()

            lines.append(md_line)

        elif tag == "tbl":
            # 表格
            from docx.table import Table
            table = Table(child, doc)
            table_id = id(child)

            if table_id in processed_tables:
                continue
            processed_tables.add(table_id)

            if skip_meta and should_skip_table(table, prev_heading):
                continue

            table_md = table_to_md(table)
            if table_md:
                if lines and lines[-1] != "":
                    lines.append("")
                lines.append(table_md)
                lines.append("")

    # ── 后处理：目录块过滤 ──────────────────────────────────────────────────
    # 在文档头部（第一个 # 标题之前），检测并移除连续短编号目录行块
    def _is_toc_like(line: str) -> bool:
        if not line or len(line) > 50:
            return False
        if re.search(r'\.{3,}\s*\d+$', line):  # 点划线尾
            return True
        if re.match(r'^[\d\.]+\s+\S', line) and len(line) < 25:  # 短编号行
            return True
        if re.match(r'^[一二三四五六七八九十]、\s\S', line) and len(line) < 20:
            return True
        if line.startswith("*") and len(line) < 30 and line.count("*") >= 2:
            return True
        return False

    cleaned = []
    first_heading_idx = None
    for i, line in enumerate(lines):
        if line.startswith("#"):
            first_heading_idx = i
            break

    if first_heading_idx is not None and first_heading_idx > 0:
        # 在第一个标题之前的区域内查找连续目录行聚集块
        toc_candidates = []
        for i in range(first_heading_idx):
            if _is_toc_like(lines[i]):
                toc_candidates.append(i)
        # 找连续块（>=5行）
        if toc_candidates:
            blocks = []
            block_start = toc_candidates[0]
            prev = toc_candidates[0]
            for idx in toc_candidates[1:]:
                if idx == prev + 1:
                    prev = idx
                else:
                    if prev - block_start + 1 >= 5:
                        blocks.append((block_start, prev))
                    block_start = idx
                    prev = idx
            if prev - block_start + 1 >= 5:
                blocks.append((block_start, prev))
            # 合并相邻块
            skip_idxs = set()
            if blocks:
                for s, e in blocks:
                    for x in range(s, e + 1):
                        skip_idxs.add(x)
                # 扩展跳过后面的空行
                for i in range(e + 1, min(e + 5, len(lines))):
                    if lines[i] == "":
                        skip_idxs.add(i)
                    else:
                        break
            # 重建lines（跳过目录块索引）
            if skip_idxs:
                lines = [l for i, l in enumerate(lines) if i not in skip_idxs]

    # 清理多余空行
    result = []
    prev_blank = False
    for line in lines:
        if line == "":
            if not prev_blank:
                result.append("")
            prev_blank = True
        else:
            result.append(line)
            prev_blank = False

    # 清理模板占位符：{填写...}、{xxx应当/按照/具体/详见...}
    # 注意：保留短的变量引用如 {持仓数量类因子}
    result_text = "\n".join(result).strip()
    _PLACEHOLDER_KW = ["填写", "解读", "具体", "详见", "描述", "举例", "说明", "参见"]
    _placeholder_re = re.compile(
        r'\{[^}]*?(?:' + '|'.join(_PLACEHOLDER_KW) + r')[^}]*\}'
    )
    result_text = _placeholder_re.sub('', result_text)
    # 清理由此产生的空行
    result_text = re.sub(r'\n{3,}', '\n\n', result_text)

    return result_text


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Word → Markdown 高质量转换")
    parser.add_argument("files",       nargs="+",           help="Word 文档路径（支持多个）")
    parser.add_argument("-o", "--out", type=str, default="", help="输出文件路径（单文件时有效）")
    parser.add_argument("--out-dir",   type=str, default="", help="批量输出目录")
    parser.add_argument("--skip-meta", action="store_true", default=True,
                        help="跳过文档信息/修订记录等非核心表格（默认开启）")
    parser.add_argument("--keep-meta", action="store_true",
                        help="保留所有表格（包括文档信息）")
    args = parser.parse_args()

    skip_meta = args.skip_meta and not args.keep_meta

    files = []
    for pattern in args.files:
        p = Path(pattern)
        if p.exists():
            files.append(p)
        else:
            import glob
            matched = glob.glob(pattern)
            files.extend(Path(m) for m in matched)

    if not files:
        print("错误: 没有找到文件")
        sys.exit(1)

    out_dir = Path(args.out_dir) if args.out_dir else None
    if out_dir:
        out_dir.mkdir(parents=True, exist_ok=True)

    for docx_path in files:
        if docx_path.suffix.lower() not in (".docx", ".doc"):
            print(f"跳过: {docx_path.name}（不是 Word 文档）")
            continue

        print(f"转换: {docx_path.name} ...", end=" ", flush=True)

        try:
            md_text = docx_to_md(docx_path, skip_meta=skip_meta)
        except Exception as e:
            print(f"失败: {e}")
            continue

        # 确定输出路径
        if args.out and len(files) == 1:
            out_path = Path(args.out)
        elif out_dir:
            out_path = out_dir / (docx_path.stem + ".md")
        else:
            out_path = docx_path.with_suffix(".md")

        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(md_text, encoding="utf-8")
        size_kb = out_path.stat().st_size // 1024
        lines   = len(md_text.splitlines())
        print(f"✓ {out_path.name}（{lines}行, {size_kb}KB）")


if __name__ == "__main__":
    main()
