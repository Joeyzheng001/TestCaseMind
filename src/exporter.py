"""
论文导出工具。

Word 优先复用知识库中的论文格式模板；PDF 使用本地字体渲染分页，
避免依赖外部 LaTeX/浏览器服务。
"""

from __future__ import annotations

import io
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple, Union

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.transforms import Bbox
    _HAS_MATPLOTLIB = True
except ImportError:
    _HAS_MATPLOTLIB = False

import os as _os
PROJECT_ROOT = Path(__file__).resolve().parent.parent
_user_data = _os.getenv("OUTPUT_PATH", "").strip()
if _user_data:
    _p = Path(_user_data)
    if _p.parts and _p.parts[0] == "~":
        _p = Path.home() / Path(*_p.parts[1:])
    OUTPUT_ROOT = _p.resolve()
else:
    OUTPUT_ROOT = Path.home() / ".thesismind"
TEMPLATE_CANDIDATES = [
    PROJECT_ROOT / "knowledge_base/references/3.0 论文格式_复旦MEM硕士论文参考模版_飞哥ProMax版2025✅(1).docx",
    PROJECT_ROOT / "knowledge_base/references/3.0 论文格式docx.docx",
]


_MD_TABLE_ROW_RE = re.compile(r"^\|.+\|$")


def _parse_markdown_table(lines: List[str]) -> Tuple[List[str], List[str], List[List[str]]]:
    """Parse a markdown table from lines. Returns (headers, alignments, rows)."""
    raw = [line.strip() for line in lines if line.strip()]
    if len(raw) < 2:
        return [], [], []

    def _cells(line: str) -> List[str]:
        s = line.strip()
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        return [c.strip() for c in s.split("|")]

    def _parse_align(cell: str) -> str:
        cell = cell.strip().strip(":-")
        if cell.startswith(":") and cell.endswith(":"):
            return "center"
        elif cell.endswith(":"):
            return "right"
        return "left"

    headers = _cells(raw[0])
    aligns = [_parse_align(c) for c in raw[1].split("|") if c.strip() and c.strip().replace("-", "").replace(":", "") == ""]
    if not aligns:
        aligns = ["left"] * len(headers)
    while len(aligns) < len(headers):
        aligns.append("left")
    rows = [_cells(r) for r in raw[2:]]
    return headers, aligns[:len(headers)], rows


def _split_content_segments(text: str) -> List[Dict[str, Any]]:
    """Split text into segments of type 'text' or 'table'. Each segment is {type, data}."""
    if not text:
        return []
    lines = text.splitlines()
    segments = []
    buf = []
    table_lines = []

    def _flush_text():
        nonlocal buf
        content = "\n".join(buf).strip()
        if content:
            segments.append({"type": "text", "data": content})
        buf = []

    def _flush_table():
        nonlocal table_lines
        headers, aligns, rows = _parse_markdown_table(table_lines)
        if headers:
            segments.append({"type": "table", "data": {"headers": headers, "aligns": aligns, "rows": rows}})
        table_lines = []

    for line in lines:
        if _MD_TABLE_ROW_RE.match(line.strip()):
            _flush_text()
            table_lines.append(line)
        else:
            if table_lines:
                _flush_table()
            buf.append(line)

    if table_lines:
        _flush_table()
    _flush_text()
    return segments


# ---------------------------------------------------------------------------
# LaTeX formula → image rendering
# ---------------------------------------------------------------------------

_FORMULA_DISPLAY_RE = re.compile(r'\$\$(.+?)\$\$', re.DOTALL)
_FORMULA_INLINE_RE = re.compile(r'\$(.+?)\$')

# Mathtext doesn't support \begin{env}...\end{env}. Pre-process known
# environments into mathtext-compatible markup so they still render.
_MATRIX_RE = re.compile(r'\\begin\{bmatrix\}(.+?)\\end\{bmatrix\}', re.DOTALL)
_CASES_RE = re.compile(r'\\begin\{cases\}(.+?)\\end\{cases\}', re.DOTALL)


def _convert_matrix_body(body: str) -> str:
    """Convert bmatrix body rows separated by \\\\, cols by & into \\atop stacks."""
    rows = [r.strip() for r in body.split('\\\\') if r.strip()]
    if not rows:
        return body
    ncols = max(len(r.split('&')) for r in rows)
    cols = []
    for c in range(ncols):
        col_stack = []
        for r in rows:
            cells = [cell.strip() for cell in r.split('&')]
            col_stack.append(cells[c] if c < len(cells) else '')
        cols.append('{' + r' \atop '.join(col_stack) + '}')
    return r' \; '.join(cols)


def _preprocess_latex(latex: str) -> str:
    """Convert unsupported LaTeX constructs to mathtext-compatible equivalents."""
    latex = _MATRIX_RE.sub(lambda m: r'\left[' + _convert_matrix_body(m.group(1)) + r'\right]', latex)
    latex = _CASES_RE.sub(
        lambda m: r'\left\{ ' + _convert_matrix_body(m.group(1)) + r' \right.',
        latex,
    )
    return latex


def _measure_mathtext(text: str, font_size: int, dpi: int) -> Tuple[float, float] | None:
    """Measure a mathtext string's width/height in inches. Returns None on failure."""
    try:
        fig, ax = plt.subplots(figsize=(0.01, 0.01), dpi=dpi)
        t = ax.text(0.5, 0.5, f"${text}$", fontsize=font_size, ha='center', va='center',
                    transform=ax.transAxes)
        fig.canvas.draw()
        bbox = t.get_window_extent(renderer=fig.canvas.get_renderer())
        bbox = bbox.transformed(fig.dpi_scale_trans.inverted())
        plt.close(fig)
        return bbox.width * 1.08, bbox.height * 1.12
    except (ValueError, RuntimeError):
        return None


def _savefig_vpad(fig, buf: io.BytesIO, dpi: int, top: float = 0.30, bottom: float = 0.10,
                   pad_inches: float = 0.04, **kw) -> None:
    """Save figure with extra transparent top/bottom padding for baseline alignment."""
    fig.canvas.draw()
    tb = fig.get_tightbbox(fig.canvas.get_renderer())
    h = tb.height
    padded = Bbox.from_bounds(
        tb.x0, tb.y0 - h * bottom,
        tb.width, h * (1 + top + bottom),
    )
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches=padded, pad_inches=pad_inches, **kw)


def _render_matrix_png(latex: str, font_size: int = 11, dpi: int = 150) -> io.BytesIO | None:
    """Render a LaTeX formula containing \\begin{bmatrix} or \\begin{cases} as a
    grid of individually-rendered mathtext cells with brackets.  Returns BytesIO
    or None."""
    if not _HAS_MATPLOTLIB:
        return None

    m_b = re.search(r'\\begin\{bmatrix\}(.+?)\\end\{bmatrix\}', latex, re.DOTALL)
    m_c = re.search(r'\\begin\{cases\}(.+?)\\end\{cases\}', latex, re.DOTALL)
    m = m_b or m_c
    if not m:
        return None
    is_cases = bool(m_c)

    prefix = latex[:m.start()].strip()
    suffix = latex[m.end():].strip()
    body = m.group(1)

    rows = [r.strip() for r in body.split('\\\\') if r.strip()]
    if not rows:
        return None

    cell_texts: List[List[str]] = []
    n_cols = 0
    for r in rows:
        cells = [c.strip() for c in r.split('&')]
        cell_texts.append(cells)
        n_cols = max(n_cols, len(cells))
    n_rows = len(cell_texts)
    for r in cell_texts:
        while len(r) < n_cols:
            r.append('')

    gap_x = 0.22
    gap_y = 0.06
    bracket_margin = 0.18

    col_w = [0.0] * n_cols
    row_h = [0.0] * n_rows

    for i in range(n_rows):
        for j in range(n_cols):
            txt = cell_texts[i][j]
            if not txt:
                col_w[j] = max(col_w[j], 0.25)
                row_h[i] = max(row_h[i], 0.18)
                continue
            dims = _measure_mathtext(txt, font_size, dpi)
            if dims:
                col_w[j] = max(col_w[j], dims[0] * 1.12)
                row_h[i] = max(row_h[i], dims[1] * 1.05)
            else:
                col_w[j] = max(col_w[j], 0.25)
                row_h[i] = max(row_h[i], 0.18)

    col_w = [max(w, 0.22) for w in col_w]
    row_h = [max(h, 0.16) for h in row_h]

    matrix_w = sum(col_w) + (n_cols - 1) * gap_x
    matrix_h = sum(row_h) + (n_rows - 1) * gap_y

    pad_x = 0.15
    pad_y = 0.12
    total_w = matrix_w + 2 * bracket_margin + pad_x
    total_h = matrix_h + pad_y * 2

    # Measure prefix for horizontal placement
    prefix_w = 0.0
    if prefix:
        dims = _measure_mathtext(prefix, font_size, dpi)
        if dims:
            prefix_w = dims[0] + 0.12

    total_w += prefix_w

    fig, ax = plt.subplots(figsize=(max(total_w, 0.5), max(total_h, 0.3)), dpi=dpi)
    ax.set_xlim(0, total_w)
    ax.set_ylim(0, total_h)
    ax.axis('off')
    fig.patch.set_alpha(0)

    mx = prefix_w + bracket_margin + pad_x * 0.5
    my = pad_y + matrix_h

    # Render cells
    y_cursor = my
    for i in range(n_rows):
        x_cursor = mx
        for j in range(n_cols):
            txt = cell_texts[i][j]
            if txt:
                ax.text(x_cursor + col_w[j] / 2, y_cursor - row_h[i] / 2,
                        f"${txt}$", fontsize=font_size, ha='center', va='center')
            x_cursor += col_w[j] + gap_x
        y_cursor -= row_h[i] + gap_y

    # Brackets
    bl = mx - 0.06
    br = mx + matrix_w + 0.03
    bb = pad_y - 0.03
    bt = pad_y + matrix_h + 0.03
    bw = 0.035

    if is_cases:
        ax.plot([bl + bw, bl], [bt, bt], 'k-', lw=1.0)
        ax.plot([bl, bl], [bt, bb], 'k-', lw=1.0)
        ax.plot([bl + bw, bl, bl, bl + bw], [bb, bb, bt, bt], 'k-', lw=1.0)
    else:
        ax.plot([bl + bw, bl, bl, bl + bw], [bb, bb, bt, bt], 'k-', lw=1.2)
        ax.plot([br - bw, br, br, br - bw], [bb, bb, bt, bt], 'k-', lw=1.2)

    # Prefix
    if prefix:
        ax.text(prefix_w / 2, total_h / 2, f"${prefix}$", fontsize=font_size,
                ha='center', va='center')

    buf = io.BytesIO()
    _savefig_vpad(fig, buf, dpi, top=0.05, bottom=0.05, pad_inches=0.04, transparent=True)
    buf.seek(0)
    plt.close(fig)
    return buf


def _render_latex_png(latex: str, font_size: int = 11, dpi: int = 150) -> io.BytesIO | None:
    """Render a LaTeX formula to a tight PNG.  Returns BytesIO or None on parse failure."""
    if not _HAS_MATPLOTLIB:
        return None

    # Route matrices to the grid renderer
    if '\\begin{bmatrix}' in latex or '\\begin{cases}' in latex:
        return _render_matrix_png(latex, font_size, dpi)

    safe = _preprocess_latex(latex)
    try:
        # Single-pass: render with va='baseline' at y=0.22 so baseline sits
        # ~22% from the bottom of the image.  When Word places the image bottom
        # on the text baseline, the formula baseline lands slightly above,
        # aligning with the x-height centre of the surrounding text.
        fig, ax = plt.subplots(figsize=(0.01, 0.01), dpi=dpi)
        text = ax.text(
            0.5, 0.22, f"${safe}$",
            fontsize=font_size, ha='center', va='baseline',
            transform=ax.transAxes,
        )
        fig.canvas.draw()
        bbox = text.get_window_extent(renderer=fig.canvas.get_renderer())
        bbox = bbox.transformed(fig.dpi_scale_trans.inverted())
        w_in, h_in = bbox.width * 1.08, bbox.height * 1.12
        plt.close(fig)

        fig, ax = plt.subplots(figsize=(max(w_in, 0.3), max(h_in, 0.2)), dpi=dpi)
        ax.text(0.5, 0.22, f"${safe}$", fontsize=font_size, ha='center', va='baseline',
                transform=ax.transAxes)
        ax.axis('off')
        fig.patch.set_alpha(0)

        buf = io.BytesIO()
        _savefig_vpad(fig, buf, dpi, top=0.35, bottom=0.05, pad_inches=0.03, transparent=True)
        buf.seek(0)
        plt.close(fig)
        return buf
    except (ValueError, RuntimeError):
        return None


def _split_with_formulas(text: str) -> List[Tuple[str, bool]]:
    """Split *text* into ``(content, is_formula)`` segments (supports both $ and $$)."""
    # Collect all matches from both patterns
    matches: List[Tuple[int, int, str]] = []  # (start, end, formula_text)
    for m in _FORMULA_DISPLAY_RE.finditer(text):
        if m.group(1):
            matches.append((m.start(), m.end(), m.group(1)))
    for m in _FORMULA_INLINE_RE.finditer(text):
        if m.group(1):
            matches.append((m.start(), m.end(), m.group(1)))

    # Sort by start position; for same start, longer match ($$) comes first
    matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))

    # Remove overlapping matches (display $$ takes precedence)
    filtered: List[Tuple[int, int, str]] = []
    last_end = 0
    for start, end, formula in matches:
        if start >= last_end:
            filtered.append((start, end, formula))
            last_end = end

    result: List[Tuple[str, bool]] = []
    cursor = 0
    for start, end, formula in filtered:
        if start > cursor:
            result.append((text[cursor:start], False))
        result.append((formula, True))
        cursor = end
    if cursor < len(text):
        result.append((text[cursor:], False))
    return result


def clean_heading(title: str, number: str = "") -> str:
    value = re.sub(r"^#{1,6}\s*", "", str(title or "")).strip()
    if number:
        value = re.sub(rf"^{re.escape(number)}\s+", "", value)
    value = re.sub(r"^\d+(?:\.\d+){0,3}\s+", "", value)
    return value.strip()


def clean_body_text(text: str, number: str = "", title: str = "") -> str:
    value = re.sub(r"^#{1,6}\s*", "", str(text or ""), flags=re.MULTILINE).strip()
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    title = clean_heading(title, number)
    lines = [line.strip() for line in value.splitlines()]
    output = []
    for index, line in enumerate(lines):
        bare = clean_heading(line, number)
        looks_like_heading = (
            bool(re.match(r"^第[一二三四五六七八九十\d]+章", line))
            or bool(re.match(r"^\d+(?:\.\d+){1,3}\s+\S{2,80}$", line))
            or bool(number and line.startswith(number))
            or bool(title and bare == title)
        )
        if index < 5 and looks_like_heading:
            continue
        if line and output and output[-1] == line:
            continue
        output.append(line)
    return "\n".join(output).strip()


def _iter_blocks(outline: Dict[str, Any], drafts: Dict[str, str]) -> Iterable[Tuple[int, str]]:
    yield 0, clean_heading(outline.get("title", "论文正文")).replace("论文:", "").strip()
    for chapter in outline.get("chapters", []):
        yield 1, clean_heading(chapter.get("title", ""))
        for section in chapter.get("sections", []):
            yield 2, f"{section.get('number')} {clean_heading(section.get('title', ''), section.get('number', ''))}"
            subsections = section.get("subsections") or []
            if subsections:
                for subsection in subsections:
                    number = subsection.get("number", "")
                    yield 3, f"{number} {clean_heading(subsection.get('title', ''), number)}"
                    draft = clean_body_text(drafts.get(number, ""), number, subsection.get("title", ""))
                    if draft:
                        yield 4, draft
            else:
                number = section.get("number", "")
                draft = clean_body_text(drafts.get(number, ""), number, section.get("title", ""))
                if draft:
                    yield 4, draft


def _template_path() -> Path | None:
    for path in TEMPLATE_CANDIDATES:
        if path.exists():
            return path
    return None


def _clear_document_body(document: Any) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _set_run_font(run, font_name: str, size=None, bold=False, color=None):
    """设置 run 的拉丁和东亚字体。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)
    if size:
        run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _png_dims_inches(buf: io.BytesIO, dpi: int = 150) -> Tuple[float, float]:
    """Return (width_inches, height_inches) from a PNG buffer header."""
    buf.seek(16)
    w_px = int.from_bytes(buf.read(4), 'big')
    h_px = int.from_bytes(buf.read(4), 'big')
    buf.seek(0)
    return w_px / dpi, h_px / dpi


def _add_body_paragraph(document, text: str) -> None:
    """Add a body-text paragraph, rendering any ``$...$`` formulas as inline images."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    segs = _split_with_formulas(text)
    if not segs:
        return

    # Entire paragraph is a single formula → display / centred
    if len(segs) == 1 and segs[0][1]:
        formula = segs[0][0]
        img = _render_latex_png(formula, font_size=12, dpi=150)
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.5
        if img:
            w, _h = _png_dims_inches(img)
            w = min(w, 5.0)
            run = para.add_run()
            run.add_picture(img, width=Inches(w))
        else:
            run = para.add_run(formula)
            _set_run_font(run, "Consolas", Pt(10))
        return

    # Mixed text + inline formulas
    para = document.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.line_spacing = 1.5
    for content, is_formula in segs:
        if is_formula:
            img = _render_latex_png(content, font_size=12, dpi=150)
            if img:
                w, h = _png_dims_inches(img)
                w = min(w, 2.5)
                run = para.add_run()
                run.add_picture(img, width=Inches(w), height=Inches(h))
            else:
                run = para.add_run(content)
                _set_run_font(run, "Consolas", Pt(10))
        else:
            run = para.add_run(content)
            _set_run_font(run, "宋体", Pt(12))


def export_docx(outline: Dict[str, Any], drafts: Dict[str, str], citations: List[Dict[str, Any]] = None) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn

    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_ROOT / "thesis_export.docx"
    template = _template_path()
    document = Document(str(template)) if template else Document()
    if template:
        _clear_document_body(document)

    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    # Normal style: 宋体 12pt
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for level, text in _iter_blocks(outline, drafts):
        if not text:
            continue
        if level == 0:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            _set_run_font(run, "黑体", Pt(18), bold=True, color=RGBColor(0, 0, 0))
        elif level in {1, 2, 3}:
            paragraph = document.add_heading(text, level=min(level, 3))
            for run in paragraph.runs:
                _set_run_font(run, "黑体", bold=True, color=RGBColor(0, 0, 0))
        else:
            segments = _split_content_segments(text)
            for seg in segments:
                if seg["type"] == "table":
                    tbl = seg["data"]
                    headers = tbl["headers"]
                    aligns = tbl["aligns"]
                    rows = tbl["rows"]
                    if not headers:
                        continue
                    ncols = len(headers)
                    docx_table = document.add_table(rows=1 + len(rows), cols=ncols)
                    docx_table.style = "Table Grid"
                    # Header row
                    for ci, hdr in enumerate(headers):
                        cell = docx_table.rows[0].cells[ci]
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(hdr)
                        _set_run_font(run, "宋体", Pt(10.5), bold=True)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Data rows
                    for ri, row in enumerate(rows):
                        for ci, val in enumerate(row):
                            cell = docx_table.rows[ri + 1].cells[ci]
                            cell.text = ""
                            run = cell.paragraphs[0].add_run(val)
                            _set_run_font(run, "宋体", Pt(10.5))
                            if ci < len(aligns):
                                if aligns[ci] == "center":
                                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                elif aligns[ci] == "right":
                                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # Add spacing after table
                    spacer = document.add_paragraph("")
                    spacer.paragraph_format.line_spacing = Pt(6)
                else:
                    for part in re.split(r"\n{2,}", seg["data"]):
                        part = part.strip()
                        if not part:
                            continue
                        _add_body_paragraph(document, part)

    # Append references
    if citations:
        ref_heading = document.add_heading("参考文献", level=1)
        for run in ref_heading.runs:
            _set_run_font(run, "黑体", bold=True, color=RGBColor(0, 0, 0))

        for i, c in enumerate(citations):
            formatted = c.get("formatted", "").strip()
            if not formatted:
                continue
            para = document.add_paragraph(f"[{i+1}] {formatted}")
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                _set_run_font(run, "宋体", Pt(10.5))

    document.save(path)
    return path


def _chinese_font_path() -> str:
    candidates = [
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
        "C:\\Windows\\Fonts\\simsun.ttc",
        "C:\\Windows\\Fonts\\msyh.ttc",
        "C:\\Windows\\Fonts\\simhei.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    raise RuntimeError(
        "未找到中文字体文件。请安装 Noto Sans CJK 或设置 THESISMIND_FONT 环境变量指向一个 .ttf/.ttc 文件。\n"
        "macOS: 系统自带宋体\n"
        "Linux: sudo apt install fonts-noto-cjk\n"
        "Windows: 系统自带宋体/微软雅黑"
    )


def _draw_pdf_table(pdf, headers: List[str], aligns: List[str], rows: List[List[str]], body_width: float) -> None:
    """Draw a formatted table in the PDF document."""
    ncols = len(headers)
    if ncols == 0:
        return
    col_w = body_width / ncols
    line_h = 5.5
    font_size = 8

    estimated_h = (len(rows) + 1) * (line_h + 2) + 4
    if pdf.get_y() + estimated_h > pdf.h - pdf.b_margin:
        pdf.add_page()

    x0 = pdf.l_margin

    def _draw_row(cells, bold=False):
        pdf.set_font("CJK", "B" if bold else "", size=font_size)
        y = pdf.get_y()
        for ci in range(ncols):
            cell_text = cells[ci] if ci < len(cells) else ""
            if len(cell_text) > 80:
                cell_text = cell_text[:77] + "..."
            pdf.set_xy(x0 + ci * col_w, y)
            a = aligns[ci][0].upper() if ci < len(aligns) else "L"
            pdf.cell(col_w, line_h + 2, cell_text, border=1, align=a)
        pdf.set_xy(x0, y + line_h + 2)

    _draw_row(headers, bold=True)
    for row in rows:
        _draw_row(row)

    pdf.ln(4)


def export_pdf(outline: Dict[str, Any], drafts: Dict[str, str], citations: List[Dict[str, Any]] = None) -> Path:
    from fpdf import FPDF

    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_ROOT / "thesis_export.pdf"

    font_path = _chinese_font_path()
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    font_sizes = {0: 22, 1: 16, 2: 13, 3: 11, 4: 10.5}

    try:
        pdf.add_font("CJK", "", font_path, uni=True)
        pdf.add_font("CJK", "B", font_path, uni=True)
    except Exception:
        pdf.add_font("CJK", font_path, uni=True)
        pdf.add_font("CJK", "B", font_path, uni=True)

    body_width = pdf.w - pdf.l_margin - pdf.r_margin

    for level, text in _iter_blocks(outline, drafts):
        if not text:
            continue

        size = font_sizes.get(level, 10.5)
        pdf.set_font("CJK", "B" if level <= 3 else "", size=size)

        if level == 0:
            pdf.set_font("CJK", "B", size=size)
            pdf.multi_cell(body_width, size * 1.4, text, align="C")
            pdf.ln(10)
        elif level == 1:
            pdf.ln(4)
            pdf.multi_cell(body_width, size * 1.3, text)
            pdf.ln(2)
        elif level in (2, 3):
            pdf.ln(2)
            indent = (level - 1) * 8
            pdf.set_x(pdf.l_margin + indent)
            pdf.multi_cell(body_width - indent, size * 1.2, text)
        else:
            pdf.set_font("CJK", "", size=size)
            segments = _split_content_segments(text)
            for seg in segments:
                if seg["type"] == "table":
                    tbl = seg["data"]
                    _draw_pdf_table(pdf, tbl["headers"], tbl["aligns"], tbl["rows"], body_width)
                else:
                    for paragraph in re.split(r"\n{2,}", seg["data"]):
                        para = paragraph.strip()
                        if not para:
                            continue
                        segs = _split_with_formulas(para)
                        if not segs:
                            continue
                        # Single formula → centred image
                        if len(segs) == 1 and segs[0][1]:
                            img = _render_latex_png(segs[0][0], font_size=12, dpi=150)
                            if img:
                                w, _h = _png_dims_inches(img)
                                w = min(w, 5.0) * 25.4  # inches → mm
                                x = pdf.l_margin + (body_width - w) / 2
                                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tf:
                                    tf.write(img.read())
                                    tf.flush()
                                    pdf.image(tf.name, x=x, w=w)
                                try:
                                    Path(tf.name).unlink()
                                except OSError:
                                    pass
                                pdf.ln(4)
                            else:
                                pdf.set_font("CJK", "", size=size)
                                pdf.multi_cell(body_width, size * 1.55, segs[0][0], align="C")
                                pdf.ln(1)
                            continue
                        # Mixed — strip $ delimiters, keep LaTeX as text
                        pdf.set_font("CJK", "", size=size)
                        plain = "".join(content for content, is_f in segs)
                        pdf.multi_cell(body_width, size * 1.55, plain, align="L")
                        pdf.ln(1)

    # Append references
    if citations:
        pdf.add_page()
        pdf.set_font("CJK", "B", size=16)
        pdf.multi_cell(body_width, 22, "参考文献", align="L")
        pdf.ln(6)
        pdf.set_font("CJK", "", size=9)
        for i, c in enumerate(citations):
            formatted = c.get("formatted", "").strip()
            if not formatted:
                continue
            ref_line = f"[{i+1}] {formatted}"
            pdf.multi_cell(body_width, 12, ref_line, align="L")
            pdf.ln(1)

    pdf.output(path)
    return path
