"""
本地文档转换工具

将 PDF/DOCX/文本类资料转换为 Markdown 或纯文本，供本地向量库索引。
转换过程默认只读本地文件，不调用外部服务。
"""

import re
import logging
from pathlib import Path
from typing import Any, Dict, Optional


TEXT_EXTENSIONS = {".md", ".txt", ".json", ".bib", ".yaml", ".yml"}
DOCUMENT_EXTENSIONS = TEXT_EXTENSIONS | {".docx", ".pdf"}


def convert_document(
    file_path: str,
    output_dir: Optional[str] = None,
    output_format: str = "md",
    max_content_chars: int = 4000,
) -> Dict[str, Any]:
    """
    将单个文档转换为 Markdown 或文本。

    Args:
        file_path: 输入文件路径
        output_dir: 输出目录；为空时只返回内容，不写文件
        output_format: md 或 txt
    """
    path = Path(file_path).expanduser()
    if not path.exists():
        return {"status": "error", "error": "file not found", "file_path": str(path)}

    try:
        content = extract_text(path)
    except Exception as exc:
        return {"status": "error", "error": str(exc), "file_path": str(path)}

    if output_format not in {"md", "txt"}:
        return {
            "status": "error",
            "error": "output_format must be md or txt",
            "file_path": str(path),
        }

    output_path = None
    if output_dir:
        output_base = Path(output_dir).expanduser()
        output_base.mkdir(parents=True, exist_ok=True)
        output_path = output_base / f"{path.stem}.{output_format}"
        output_path.write_text(content, encoding="utf-8")

    preview = content[:max_content_chars]
    truncated = len(content) > max_content_chars

    return {
        "status": "ok",
        "file_path": str(path),
        "output_path": str(output_path) if output_path else None,
        "format": output_format,
        "characters": len(content),
        "truncated": truncated,
        "content": preview,
    }


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    raise ValueError(f"unsupported document type: {suffix}")


def extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to convert DOCX files") from exc

    document = Document(path)
    lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = ""
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
        if "heading 1" in style_name or "标题 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name or "标题 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name or "标题 3" in style_name:
            lines.append(f"### {text}")
        else:
            lines.append(text)

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [_normalize_cell(cell.text) for cell in row.cells]
            rows.append(cells)
        lines.extend(_table_to_markdown(rows))

    return _normalize_text("\n\n".join(lines))


def extract_pdf(path: Path) -> str:
    reader_class = _load_pdf_reader()
    _quiet_pdf_loggers()
    reader = reader_class(str(path))

    pages = []
    for index, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        text = _normalize_text(text)
        if text:
            pages.append(f"## Page {index}\n\n{text}")

    if not pages:
        raise RuntimeError("no extractable text found in PDF")

    title = f"# {path.stem}"
    return _normalize_text(f"{title}\n\n" + "\n\n".join(pages))


def _load_pdf_reader():
    try:
        from pypdf import PdfReader

        return PdfReader
    except ImportError:
        pass

    try:
        from PyPDF2 import PdfReader

        return PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "pypdf is required to convert PDF files. Install dependencies from requirements.txt."
        ) from exc


def _quiet_pdf_loggers() -> None:
    for logger_name in ("pypdf", "PyPDF2"):
        logging.getLogger(logger_name).setLevel(logging.ERROR)


def _table_to_markdown(rows):
    if not rows:
        return []

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    body = normalized[1:]

    output = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    output.extend("| " + " | ".join(row) + " |" for row in body)
    return ["\n".join(output)]


def _normalize_cell(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().replace("|", "\\|")


def _normalize_text(text: str) -> str:
    text = text.encode("utf-8", errors="replace").decode("utf-8")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
