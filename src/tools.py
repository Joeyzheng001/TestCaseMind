"""
基础工具集 - 所有Agent可用的原子工具
对应 learn-claude-code s02 Tool Use 机制
"""

import os
import json
import subprocess
import shlex
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path
from datetime import datetime

from src.vector_store import build_index, search_index
from src.document_converter import convert_document


# ==================== 文件操作工具 ====================

PROJECT_ROOT = Path(__file__).resolve().parent.parent


def _resolve_project_path(file_path: str) -> Path:
    path = Path(file_path).expanduser()
    if not path.is_absolute():
        path = PROJECT_ROOT / path
    path = path.resolve()

    try:
        path.relative_to(PROJECT_ROOT)
    except ValueError as exc:
        raise ValueError(f"Path outside project is not allowed: {file_path}") from exc

    return path


def read_file(file_path: str) -> str:
    """读取文件内容"""
    try:
        safe_path = _resolve_project_path(file_path)
        with open(safe_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        return f"Error reading file: {str(e)}"


def write_file(file_path: str, content: str) -> str:
    """写入文件"""
    try:
        safe_path = _resolve_project_path(file_path)
        safe_path.parent.mkdir(parents=True, exist_ok=True)
        with open(safe_path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"File written successfully: {safe_path.relative_to(PROJECT_ROOT)}"
    except Exception as e:
        return f"Error writing file: {str(e)}"


def list_directory(dir_path: str) -> str:
    """列出目录内容"""
    try:
        safe_path = _resolve_project_path(dir_path)
        items = os.listdir(safe_path)
        return "\n".join(sorted(items))
    except Exception as e:
        return f"Error listing directory: {str(e)}"


# ==================== 论文分析工具 ====================


import re

# Patterns for detecting section headings in plain text (PDF/DOCX output)
_HEADING_PATTERNS = [
    # Numbered: "1 Introduction", "2.1 Methods", "3.1.2 Data Collection"
    (re.compile(r"^(\d+(?:\.\d+)*)\s+(.+)$", re.UNICODE), lambda m: (
        m.group(1).count(".") + 1, m.group(2).strip()
    )),
    # Chinese numbered: "一、引言", "二、方法"
    (re.compile(r"^([一二三四五六七八九十]+)[、，,]\s*(.+)$"), lambda m: (1, m.group(2).strip())),
    # Chinese chapter: "第一章 绪论"
    (re.compile(r"^第([一二三四五六七八九十百]+)章\s*(.+)$"), lambda m: (1, m.group(2).strip())),
    # Roman numeral: "I. Introduction", "II. Methods"
    (re.compile(r"^([IVX]+)\.\s+(.+)$"), lambda m: (1, m.group(2).strip())),
    # Keyword-style: "Abstract", "Introduction", "Methods", "Conclusion"
    (re.compile(
        r"^(Abstract|Introduction|Related\s+Work|Method|Experiment|"
        r"Results?|Discussion|Conclusion|References?|Acknowledgments?|"
        r"Appendix)\s*$",
        re.IGNORECASE,
    ), lambda m: (1, m.group(1))),
]


def _is_heading_line(line: str) -> "Optional[Tuple[int, str]]":
    """Detect if a line is a section heading. Returns (level, title) or None."""
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return None
    for pattern, extractor in _HEADING_PATTERNS:
        m = pattern.match(stripped)
        if m:
            return extractor(m)
    return None


def analyze_paper_structure(paper_content: str) -> Dict[str, Any]:
    """
    分析论文结构，支持 Markdown、PDF提取文本、DOCX提取文本。
    """
    try:
        lines = paper_content.split("\n")

        structure = {
            "title": lines[0].lstrip("#").strip() if lines else "",
            "total_lines": len(lines),
            "total_chars": len(paper_content),
            "sections": [],
            "has_abstract": any(
                "摘要" in line or "abstract" in line.lower() for line in lines[:50]
            ),
            "has_references": any(
                "参考文献" in line or "references" in line.lower()
                for line in lines[-50:]
            ),
        }

        for i, line in enumerate(lines):
            # Markdown headings
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                title = line.lstrip("#").strip()
                if title:
                    structure["sections"].append(
                        {"level": level, "title": title, "line_number": i}
                    )
                continue

            # Plain-text heading patterns (from PDF/DOCX extraction)
            heading = _is_heading_line(line)
            if heading:
                level, title = heading
                structure["sections"].append(
                    {"level": level, "title": title, "line_number": i}
                )

        return structure
    except Exception as e:
        return {"error": str(e)}


def _extract_metadata_pdf(file_path: Path) -> Dict[str, Any]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        info = reader.metadata or {}
        pages = len(reader.pages)
        text = "\n".join(
            page.extract_text() or "" for page in reader.pages[:5]
        )
        full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return {
            "is_pdf": True,
            "title": info.get("/Title", ""),
            "author": info.get("/Author", ""),
            "subject": info.get("/Subject", ""),
            "page_count": pages,
            "word_count": len(full_text.split()),
            "preview": text[:500],
        }
    except ImportError:
        return {"is_pdf": True, "error": "pypdf not installed"}
    except Exception as e:
        return {"is_pdf": True, "error": str(e)}


def _extract_metadata_docx(file_path: Path) -> Dict[str, Any]:
    try:
        from docx import Document

        doc = Document(str(file_path))
        props = doc.core_properties
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        return {
            "is_docx": True,
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "paragraph_count": len(paragraphs),
            "word_count": len(full_text.split()),
            "preview": full_text[:500],
        }
    except ImportError:
        return {"is_docx": True, "error": "python-docx not installed"}
    except Exception as e:
        return {"is_docx": True, "error": str(e)}


def extract_paper_metadata(paper_file: str) -> Dict[str, Any]:
    """
    从论文文件提取元数据，支持 .docx / .pdf / .md。
    """
    try:
        file_path = _resolve_project_path(paper_file)
        suffix = file_path.suffix.lower()
        metadata = {
            "file_path": str(file_path.relative_to(PROJECT_ROOT)),
            "file_type": suffix,
            "file_size": file_path.stat().st_size if file_path.exists() else 0,
        }

        if suffix == ".md":
            content = read_file(str(file_path))
            lines = content.split("\n")
            metadata.update({
                "title": lines[0].lstrip("#").strip()
                if lines and lines[0].startswith("#")
                else "",
                "is_markdown": True,
                "word_count": len(content.split()),
                "line_count": len(lines),
            })
        elif suffix == ".pdf":
            metadata.update(_extract_metadata_pdf(file_path))
        elif suffix == ".docx":
            metadata.update(_extract_metadata_docx(file_path))
        return metadata
    except Exception as e:
        return {"file_path": paper_file, "error": str(e)}


# ==================== 框架生成工具 ====================


def generate_mermaid_framework(topic: str, phases: List[str]) -> str:
    """
    生成Mermaid格式的研究框架流程图
    """
    mermaid_code = f"graph TD\n"
    mermaid_code += f'    Start["研究主题: {topic}"] --> Phase1\n'

    for i, phase in enumerate(phases, 1):
        next_phase = f"Phase{i + 1}" if i < len(phases) else "End"
        mermaid_code += f'    Phase{i}["第{i}阶段: {phase}"] --> {next_phase}\n'

    mermaid_code += f'    End["研究完成"] --> Output["论文输出"]\n'

    return mermaid_code


def generate_research_framework(
    topic: str, discipline: str = "general"
) -> Dict[str, Any]:
    """
    生成完整的研究框架
    """
    frameworks = {
        "general": {
            "phases": [
                "文献回顾",
                "问题定义",
                "方法设计",
                "实验实施",
                "结果分析",
                "结论讨论",
            ],
            "components": [
                "背景",
                "理论基础",
                "关键问题",
                "研究方法",
                "实验设计",
                "预期结果",
            ],
        },
        "computer_science": {
            "phases": [
                "技术调研",
                "算法设计",
                "系统实现",
                "性能测试",
                "优化改进",
                "论文总结",
            ],
            "components": [
                "问题陈述",
                "相关工作",
                "核心算法",
                "系统架构",
                "实验评估",
                "应用前景",
            ],
        },
        "medical": {
            "phases": [
                "临床背景",
                "文献综述",
                "研究方案",
                "临床试验",
                "数据分析",
                "临床意义",
            ],
            "components": [
                "医学基础",
                "诊疗现状",
                "研究假设",
                "方法论",
                "统计分析",
                "临床应用",
            ],
        },
    }

    framework_template = frameworks.get(discipline, frameworks["general"])

    return {
        "topic": topic,
        "discipline": discipline,
        "research_phases": framework_template["phases"],
        "key_components": framework_template["components"],
        "mermaid_graph": generate_mermaid_framework(
            topic, framework_template["phases"]
        ),
        "created_at": datetime.now().isoformat(),
    }


# ==================== 大纲生成工具 ====================


def _framework_text(framework: Dict[str, Any]) -> str:
    """把框架关键信息拼成可匹配的文本。"""
    values = [
        framework.get("topic", ""),
        framework.get("discipline", ""),
        " ".join(framework.get("research_phases", []) or []),
        " ".join(framework.get("key_components", []) or []),
    ]
    return " ".join(str(value) for value in values if value)


def _is_quality_management_topic(framework: Dict[str, Any]) -> bool:
    text = _framework_text(framework)
    keywords = [
        "质量管理",
        "质量改进",
        "CMMI",
        "PDCA",
        "车载",
        "新能源",
        "研发质量",
        "软件质量",
    ]
    return any(keyword.lower() in text.lower() for keyword in keywords)


def _section(number: str, title: str, subsections: List[str], depth: int) -> Dict[str, Any]:
    item: Dict[str, Any] = {"level": 2, "number": number, "title": title}
    if depth >= 3 and subsections:
        item["subsections"] = [
            {"level": 3, "number": f"{number}.{idx}", "title": subtitle}
            for idx, subtitle in enumerate(subsections, 1)
        ]
    return item


def _build_chapter(
    number: int, title: str, sections: List[Dict[str, Any]], depth: int
) -> Dict[str, Any]:
    return {
        "level": 1,
        "number": number,
        "title": f"第{number}章 {title}",
        "sections": [
            _section(f"{number}.{idx}", section["title"], section.get("subsections", []), depth)
            for idx, section in enumerate(sections, 1)
        ],
    }


def _quality_management_outline(topic: str, depth: int) -> List[Dict[str, Any]]:
    """质量管理方向的论文大纲（与通用模板章标题一致，但节/小节更贴合质量主题）。"""
    short = _shorten_topic(topic)
    chapters = [
        (
            "绪论",
            [
                {"title": "研究背景", "subsections": [f"{short}相关产业发展趋势", f"{short}管理挑战"]},
                {"title": "研究目的与意义", "subsections": ["理论意义", "实践意义"]},
                {"title": "国内外研究现状", "subsections": ["质量管理研究现状", f"{short}管理研究现状"]},
                {"title": "研究内容与技术路线", "subsections": ["研究内容", "研究方法", "技术路线"]},
                {"title": "研究创新点", "subsections": ["方法应用创新", "管理改进创新"]},
            ],
        ),
        (
            "理论基础与文献综述",
            [
                {"title": "质量管理相关理论", "subsections": ["全面质量管理", "过程质量管理"]},
                {"title": "CMMI过程改进模型", "subsections": ["CMMI模型结构", "CMMI在研发过程中的适用性"]},
                {"title": "PDCA质量改进循环", "subsections": ["PDCA循环内涵", "PDCA与持续改进机制"]},
                {"title": "问题识别与评价方法", "subsections": ["访谈问卷法", "鱼骨图与5Why分析", "AHP/德尔菲权重确定"]},
            ],
        ),
        (
            f"{short}现状与问题分析",
            [
                {"title": "研究对象与业务流程", "subsections": [f"{short}业务流程", "质量管理边界界定"]},
                {"title": "质量管理现状调研", "subsections": ["流程制度现状", "质量数据统计", "访谈问卷结果"]},
                {"title": "核心问题发现过程", "subsections": ["问题池构建", "问题归类", "关键问题筛选"]},
                {"title": "问题成因分析", "subsections": ["人员因素", "流程因素", "工具与数据因素", "组织保障因素"]},
            ],
        ),
        (
            f"{short}优化方案设计",
            [
                {"title": "优化目标与原则", "subsections": ["优化目标", "方案设计原则"]},
                {"title": "基于CMMI的研发过程改进", "subsections": ["过程域映射", "流程标准化改进"]},
                {"title": "基于PDCA的质量闭环机制", "subsections": ["计划阶段", "执行阶段", "检查阶段", "改进阶段"]},
                {"title": "关键问题专项改进措施", "subsections": ["需求变更管理", "测试验证管理", "质量数据管理"]},
                {"title": "实施保障机制", "subsections": ["组织保障", "制度保障", "工具与人员保障"]},
            ],
        ),
        (
            "方案实施与效果评价",
            [
                {"title": "实施方案设计", "subsections": ["实施范围", "实施步骤", "实施周期"]},
                {"title": "效果评价指标体系", "subsections": ["缺陷率指标", "返工率指标", "交付稳定性指标"]},
                {"title": "实施前后对比分析", "subsections": ["质量指标对比", "流程效率对比", "满意度对比"]},
                {"title": "优化效果评价", "subsections": ["改进效果总结", "适用性与局限性分析"]},
            ],
        ),
        (
            "结论与展望",
            [
                {"title": "研究结论", "subsections": ["核心问题结论", "优化方案结论", "验证结果结论"]},
                {"title": "管理启示", "subsections": ["对研发质量管理的启示", f"对{short}项目管理的启示"]},
                {"title": "研究不足与展望", "subsections": ["研究不足", "后续研究方向"]},
            ],
        ),
    ]
    return [_build_chapter(idx, title, sections, depth) for idx, (title, sections) in enumerate(chapters, 1)]


def _universal_outline(topic: str, methods: List[str], direction: str, depth: int, sections_hint: Optional[Dict[int, List[str]]] = None) -> List[Dict[str, Any]]:
    """通用工程管理硕士论文大纲模板。

    6 章框架是锁死的，不会因主题不同而改变章标题。
    节标题根据主题、方法论和研究方向动态填充；如果 sections_hint
    为 None 则使用内置默认节标题（确保 LLM 失败时也有可用大纲）。
    """
    # 从主题中提取核心短语用于章 3/4 标题（去掉过长修饰）
    short_topic = _shorten_topic(topic)

    # ── 默认节标题（LLM 失败时的兜底） ──
    default_sections = {
        1: ["研究背景", "研究目的与意义", "研究内容与方法", "创新点与论文结构"],
        2: ["相关概念与理论基础", "国内外研究现状", "文献述评与研究切入点"],
        3: ["研究对象与现状调研", "核心问题识别与分析", "问题成因深度剖析"],
        4: ["优化目标与设计原则", "核心优化方案设计", "实施路径与保障机制"],
        5: ["实施过程与数据收集", "效果评价指标体系", "实施效果分析与讨论"],
        6: ["研究结论", "管理启示", "研究不足与展望"],
    }
    # 从用户选定的方法中生成更具体的节标题
    method_names = ", ".join(methods[:4]) if methods else ""

    # ── 使用 sections_hint（LLM 生成）或默认值 ──
    sec = sections_hint if sections_hint else default_sections

    chapters = [
        (1, "绪论", sec.get(1, default_sections[1])),
        (2, "理论基础与文献综述", sec.get(2, default_sections[2])),
        (3, f"{short_topic}现状与问题分析", sec.get(3, default_sections[3])),
        (4, f"{short_topic}优化方案设计", sec.get(4, default_sections[4])),
        (5, "方案实施与效果评价", sec.get(5, default_sections[5])),
        (6, "结论与展望", sec.get(6, default_sections[6])),
    ]

    return [_build_chapter(num, title, [{"title": t, "subsections": []} for t in sections], depth)
            for num, title, sections in chapters]


def _shorten_topic(topic: str) -> str:
    """将过长主题压缩为适合嵌入章标题的简短形式。"""
    t = topic.strip()
    # 去掉「研究」「以…为例」「——…」等后缀
    for delim in ("——", "——", "—", "：", ":"):
        if delim in t:
            # 保留前半部分
            t = t.split(delim, 1)[0].strip()
    # 去掉常见前缀
    for prefix in ("基于",):
        if t.startswith(prefix) and len(t) > 8:
            pass  # keep it
    if len(t) > 20:
        t = t[:20]
    return t


def _general_outline(framework: Dict[str, Any], depth: int) -> List[Dict[str, Any]]:
    """重定向到通用模板。保留此函数以兼容旧调用方。"""
    topic = framework.get("topic", "研究主题")
    direction = framework.get("direction", "")
    methods = framework.get("methods", []) or []
    return _universal_outline(topic, methods, direction, depth)


def generate_outline(framework: Dict[str, Any], depth: int = 3) -> Dict[str, Any]:
    """
    基于研究框架生成论文大纲

    Args:
        framework: generate_research_framework 返回的框架，或包含 topic/research_phases/key_components 的字典。
        depth: 大纲层级深度。2 表示章-节，3 表示章-节-小节。
    """
    topic = framework.get("topic", "未命名")
    normalized_depth = max(2, min(int(depth or 3), 3))
    chapters = (
        _quality_management_outline(topic, normalized_depth)
        if _is_quality_management_topic(framework)
        else _general_outline(framework, normalized_depth)
    )
    total_sections = sum(len(chapter["sections"]) for chapter in chapters)
    total_subsections = sum(
        len(section.get("subsections", []))
        for chapter in chapters
        for section in chapter["sections"]
    )

    return {
        "title": f"论文: {topic}",
        "depth": normalized_depth,
        "chapters": chapters,
        "chapter_count": len(chapters),
        "metadata": {
            "total_chapters": len(chapters),
            "total_sections": total_sections,
            "total_subsections": total_subsections,
            "outline_type": "quality_management"
            if _is_quality_management_topic(framework)
            else "general",
            "created_at": datetime.now().isoformat(),
        },
    }


# ==================== 引用管理工具 ====================


def format_citation(author: str, year: str, title: str, style: str = "apa") -> str:
    """
    生成格式化的引用文献
    支持 APA, Chicago, Harvard, GB/T 7714 格式
    """
    citations = {
        "apa": f"{author} ({year}). {title}.",
        "chicago": f'{author}. "{title}." {year}.',
        "harvard": f"{author}, {year}. {title}.",
        "gb7714": f"{author}. {title}[J]. {year}.",
    }

    return citations.get(style, citations["apa"])


def search_citations(keyword: str, limit: int = 10) -> List[Dict[str, str]]:
    """
    搜索相关引用文献
    优先从本地知识库检索，避免返回占位数据。
    """
    try:
        result = search_index(query=keyword, limit=limit)
        citations = []
        seen = set()
        for item in result.get("results", []):
            metadata = item.get("metadata", {})
            title = item.get("title") or Path(item.get("path", "")).stem
            source = item.get("path", "local_knowledge_base")
            dedupe_key = (title, source)
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)
            citations.append(
                {
                    "title": title,
                    "author": metadata.get("author", "未知"),
                    "year": metadata.get("year", "未知"),
                    "source": source,
                    "score": str(item.get("score", "")),
                }
            )
            if len(citations) >= limit:
                break
        return citations
    except Exception as exc:
        return [
            {
                "title": "本地文献检索失败",
                "author": "system",
                "year": "未知",
                "source": str(exc),
            }
        ]


# ==================== 本地知识库工具 ====================


def build_knowledge_index(
    source_dirs: List[str] = None, db_path: str = None, reset: bool = False
) -> Dict[str, Any]:
    """
    构建本地向量知识库索引
    默认索引 knowledge_base/ 和 skills/
    """
    return build_index(source_dirs=source_dirs, db_path=db_path, reset=reset)


def search_knowledge_base(
    query: str, limit: int = 5, db_path: str = None
) -> Dict[str, Any]:
    """
    检索本地向量知识库
    """
    return search_index(query=query, limit=limit, db_path=db_path)


def convert_local_document(
    file_path: str, output_dir: str = None, output_format: str = "md"
) -> Dict[str, Any]:
    """
    将本地 PDF/DOCX/文本资料转换为 Markdown 或纯文本
    """
    safe_file_path = str(_resolve_project_path(file_path))
    safe_output_dir = str(_resolve_project_path(output_dir)) if output_dir else None
    return convert_document(
        file_path=safe_file_path,
        output_dir=safe_output_dir,
        output_format=output_format,
    )


# ==================== 格式检查工具 ====================


def _check_docx_format(file_path: Path) -> Dict[str, Any]:
    """实际检查 DOCX 文档的格式问题。"""
    from docx import Document

    doc = Document(str(file_path))
    issues = {
        "font_issues": [],
        "spacing_issues": [],
        "heading_issues": [],
        "page_issues": [],
        "suggestions": [],
    }

    # 检查字体
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            font_name = run.font.name
            if font_name and font_name.lower() not in (
                "times new roman", "simsun", "宋体", "arial", "calibri",
                "黑体", "simhei", "楷体", "kaiti", "fangsong", "仿宋",
                "cambria", "georgia",
            ):
                issues["font_issues"].append(
                    f"段落 {i + 1}: 非学术标准字体 '{font_name}'"
                )
                break  # 每个段落只报一次

    # 检查段落间距
    for i, para in enumerate(doc.paragraphs):
        pf = para.paragraph_format
        line_spacing = pf.line_spacing
        if line_spacing and line_spacing != 1.5:
            issues["spacing_issues"].append(
                f"段落 {i + 1}: 行距 {line_spacing} (建议 1.5 倍)"
            )

    # 检查标题层级
    heading_levels = set()
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split()[-1])
                heading_levels.add(level)
            except ValueError:
                pass
    if heading_levels and min(heading_levels) > 1:
        issues["heading_issues"].append(
            "缺少一级标题 (Heading 1)，标题层级从 H2 开始"
        )

    # 检查页边距 (默认值 2.54 cm = 1 inch)
    for section in doc.sections:
        margin_cm = round(section.left_margin / 360000 * 2.54, 1)
        if margin_cm < 2.0:
            issues["page_issues"].append(
                f"左边距 {margin_cm}cm (建议 ≥ 2.5cm)"
            )
        break  # 只检查第一个节

    if not any(issues[k] for k in issues if k != "suggestions"):
        issues["suggestions"].append("文档格式无明显问题")
    return issues


def _check_md_format(file_path: Path) -> Dict[str, Any]:
    """检查 Markdown 文档的格式问题。"""
    content = read_file(str(file_path))
    lines = content.split("\n")
    issues = {
        "font_issues": [],
        "spacing_issues": [],
        "heading_issues": [],
        "page_issues": [],
        "suggestions": [],
    }

    seen_levels = set()
    prev_level = 0
    for i, line in enumerate(lines):
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            if level - prev_level > 1:
                issues["heading_issues"].append(
                    f"行 {i + 1}: 标题层级跳跃 (H{prev_level} → H{level})"
                )
            prev_level = level
            seen_levels.add(level)

    if 1 not in seen_levels:
        issues["heading_issues"].append("缺少一级标题 (H1)")
    if not any(line.strip() and not line.startswith("#") for line in lines):
        issues["suggestions"].append("文档缺少正文内容")

    if not any(issues[k] for k in issues if k != "suggestions"):
        issues["suggestions"].append("Markdown 格式无明显问题")
    return issues


def check_document_format(file_path: str) -> Dict[str, Any]:
    """
    检查文档格式规范，实际读取文件内容进行分析。
    支持 .docx, .md 格式。
    """
    try:
        path = _resolve_project_path(file_path)
    except Exception as e:
        return {"error": str(e)}

    if not path.exists():
        return {"error": "文件不存在"}

    if path.suffix == ".docx":
        issues = _check_docx_format(path)
    elif path.suffix == ".md":
        issues = _check_md_format(path)
    else:
        return {
            "file": str(path),
            "status": "skipped",
            "issues": {"suggestions": [f"暂不支持 {path.suffix} 格式的格式检查"]},
            "passed": True,
        }

    return {
        "file": str(path),
        "status": "checked",
        "issues": issues,
        "passed": all(
            not issues[k] for k in issues if k != "suggestions"
        ),
    }


# ==================== 命令行工具 ====================


def run_command(command: str, cwd: str = None) -> str:
    """
    执行系统命令
    """
    if os.getenv("ENABLE_RUN_COMMAND", "false").lower() != "true":
        return "run_command is disabled. Set ENABLE_RUN_COMMAND=true to enable it."

    allowed_commands = {
        "ls",
        "pwd",
        "rsvg-convert",
    }
    python_commands = {"python", "python3"}

    try:
        args = shlex.split(command)
        if not args:
            return "No command provided"
        if args[0] not in allowed_commands and args[0] not in python_commands:
            return f"Command not allowed: {args[0]}"

        safe_cwd = str(_resolve_project_path(cwd or "."))
        if args[0] in python_commands:
            if len(args) < 2 or args[1].startswith("-"):
                return "Python commands must run a project script; inline code and module execution are disabled."
            args[1] = str(_resolve_project_path(args[1]))
        result = subprocess.run(
            args, cwd=safe_cwd, capture_output=True, text=True, timeout=30
        )
        output = result.stdout if result.returncode == 0 else result.stderr
        return output[:8000]
    except subprocess.TimeoutExpired:
        return "Command timed out"
    except Exception as e:
        return f"Error: {str(e)}"


# ==================== 工具注册表 ====================

TOOLS = {
    # 文件操作
    "read_file": {"handler": read_file, "description": "读取文件内容"},
    "write_file": {"handler": write_file, "description": "写入文件"},
    "list_directory": {"handler": list_directory, "description": "列出目录内容"},
    # 论文分析
    "analyze_paper_structure": {
        "handler": analyze_paper_structure,
        "description": "分析论文结构",
    },
    "extract_paper_metadata": {
        "handler": extract_paper_metadata,
        "description": "提取论文元数据",
    },
    # 框架生成
    "generate_research_framework": {
        "handler": generate_research_framework,
        "description": "生成研究框架",
    },
    # 大纲生成
    "generate_outline": {"handler": generate_outline, "description": "生成论文大纲"},
    # 引用管理
    "format_citation": {
        "handler": format_citation,
        "description": "生成格式化引用文献",
    },
    "search_citations": {
        "handler": search_citations,
        "description": "搜索相关引用文献",
    },
    # 本地知识库
    "build_knowledge_index": {
        "handler": build_knowledge_index,
        "description": "构建本地向量知识库索引",
    },
    "search_knowledge_base": {
        "handler": search_knowledge_base,
        "description": "检索本地向量知识库",
    },
    "convert_local_document": {
        "handler": convert_local_document,
        "description": "将本地PDF/DOCX/文本资料转换为Markdown或纯文本",
    },
    # 格式检查
    "check_document_format": {
        "handler": check_document_format,
        "description": "检查文档格式",
    },
    # 命令执行
    "run_command": {"handler": run_command, "description": "执行系统命令"},
}


def get_tools_definitions() -> List[Dict[str, Any]]:
    """
    获取工具定义列表 (用于Claude API)
    """
    return [
        {
            "name": name,
            "description": tool["description"],
            "input_schema": {"type": "object", "properties": {}, "required": []},
        }
        for name, tool in TOOLS.items()
    ]
